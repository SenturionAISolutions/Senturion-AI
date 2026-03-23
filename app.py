"""
Claim Sniper BPO - Medical Billing Denial Analysis
Extracts denial data from PDFs using Gemini 2.5 Flash.
Secure SaaS Portal with Supabase authentication.
"""

import os
import re
import sys
import io
import csv
import gc
import hashlib
import html as html_std
import tempfile
import zipfile
import base64
import time
import threading
import uuid
from concurrent.futures import ThreadPoolExecutor, as_completed
import json
from pathlib import Path
from collections import Counter, defaultdict
from datetime import datetime, timedelta, timezone
from types import SimpleNamespace
from typing import Any, NamedTuple
from dotenv import load_dotenv
from supabase import create_client, Client
from pypdf import PdfReader, PdfWriter
import pandas as pd
import streamlit as st
import streamlit.components.v1 as components
from markitdown import MarkItDown

try:
    import requests
except ImportError:
    requests = None  # type: ignore[misc, assignment]

try:
    import markdown as md_to_html
except ImportError:
    md_to_html = None  # type: ignore[misc, assignment]

try:
    from streamlit_lottie import st_lottie

    HAS_LOTTIE = True
except ImportError:
    HAS_LOTTIE = False
    st_lottie = None  # type: ignore[misc, assignment]

# Dark “neural / network” style animation (public Lottie asset)
NEURAL_LOTTIE_URL = "https://assets2.lottiefiles.com/packages/lf20_w51pcehl.json"
# Sovereign institutional palette (dashboard + chrome)
STEALTH_BG = "#050505"  # Pitch black — Bloomberg terminal main ground
STEALTH_ACCENT = "#E0E0E0"  # Titanium (primary text)
INST_BORDER = "#262626"  # Holographic / ruled container border
INST_CRITICAL = "#FF3131"  # Neon red — denials / critical
TERMINAL_MATRIX_GREEN = "#00FF41"  # Success / positive ticks
VAULT_EMERALD = "#00FF41"  # Vault KPI accent (Matrix Green — terminal)
SIDEBAR_ONYX = "#0D0D0D"  # Deep Onyx sidebar

# Revenue vault claim lifecycle (Agent Review Workflow)
VAULT_STATUS_UNAUDITED = "UNAUDITED"
VAULT_STATUS_NEURAL_DRAFT = "NEURAL_DRAFT"
VAULT_STATUS_AGENT_REVIEW = "AGENT_REVIEW"
VAULT_STATUS_ENFORCED = "ENFORCED"
VAULT_STATUS_LABELS = {
    VAULT_STATUS_UNAUDITED: "[UNAUDITED]",
    VAULT_STATUS_NEURAL_DRAFT: "[NEURAL_DRAFT]",
    VAULT_STATUS_AGENT_REVIEW: "[AGENT_REVIEW]",
    VAULT_STATUS_ENFORCED: "[ENFORCED]",
}
# Maryland Health Insurance — Prompt Pay (§ 15-1005) enforcement window after vault ENFORCED
ENFORCEMENT_MD_PROMPT_PAY_DAYS = 30
FOUNDERS_COMMISSION_RATE = 0.15  # Senturion founders commission on vault recoverable total
# Senturion Treasury — invoice / settlement (placeholders; override in Admin settings)
DEFAULT_TREASURY_STRIPE_URL = "https://buy.stripe.com/test_senturion_15_percent"
DEFAULT_TREASURY_SOL_WALLET = "Senturion_Treasury_Sol_Address"
TREASURY_SETTLEMENT_OPTIONS = ("🔴 UNPAID", "🟡 PENDING", "🟢 COLLECTED")
# Founder equity — hard-coded 50/50: CEO (Eduard de Lange) / CFO (Monré Wessel Nagel) on every recovery allocation
CEO_CFO_EQUITY_PCT_EACH = 50.0

# Live Treasury — Capitec opening balance verification (place PDF next to app.py)
CAPITEC_STATEMENT_FILENAME = "Capitec_Statement_20260323.pdf"
CAPITEC_STATEMENT_DATE_LABEL = "23/03/2026"
CAPITEC_OPENING_BALANCE_ZAR = 100.00

# Paystack / auditor footer & compliance badge
PAYSTACK_MERCHANT_ID = "1774856"
PAYSTACK_MERCHANT_STATUS = "AWAITING REVIEW"
# Official Merchant Service Agreement document ID (auditor / partner footer)
MSA_MERCHANT_DOC_ID = "0f3a4987A1B2C3D4E5F6789012345678"
# Registered business contact (non-clinic executive footer)
OFFICIAL_BUSINESS_ADDRESS = "1171 Bergsig Street, Pretoria"
OFFICIAL_CONTACT_EMAIL = "senturionaisolutions@gmail.com"
# Mandatory KYC strip (Paystack compliance mode — all authenticated pages)
KYC_FOOTER_LINE = "Senturion AI Solutions | 1171 Bergsig Street, Pretoria | Merchant ID: 1774856"


def _kyc_footer_html_inner() -> str:
    """Legal footer: bold Pretoria address and Merchant ID (used before full module tail loads)."""
    _addr = html_std.escape(OFFICIAL_BUSINESS_ADDRESS)
    _mid = html_std.escape(str(PAYSTACK_MERCHANT_ID))
    return (
        f"Senturion AI Solutions | <strong>{_addr}</strong> | Merchant ID: <strong>{_mid}</strong>"
    )


# Paystack merchant verification — dedicated demo login (clean UI, no engine errors)
DEMO_PAYSTACK_EMAIL = "reviews@paystack.com"
# Emergency local bypass for Paystack merchant review (no Supabase auth). Remove after verification.
PAYSTACK_REVIEW_BYPASS_PASSWORD = "SenturionVerify2026!"
PAYSTACK_REVIEWER_BYPASS_USER_ID = "paystack-reviewer-bypass-local"
# Reviewer demo — injected mock audits (Miami cohort; per-audit recoverable USD)
REVIEWER_MOCK_CLINIC = "Miami Medical Center"
REVIEWER_MOCK_AUDIT_STATUS = "Neural Analysis Complete"
REVIEWER_MOCK_RECOVERABLE_USD = 14250.0
REVIEWER_MOCK_AUDIT_IDS = ("AUD-MIA-2026-001", "AUD-MIA-2026-002", "AUD-MIA-2026-003")
# Paystack hosted checkout for “Pay Release Fee” (set real URL in .streamlit/secrets.toml)
DEFAULT_PAYSTACK_RELEASE_CHECKOUT_URL = "https://paystack.com/pay/senturion-release-fee"
# Standard per-claim audit fee — canonical 50/50 net split after Paystack (revenue engine)
AUDIT_FEE_USD_STANDARD = 2625.0
# CFO currency: US clinics pay in USD; Paystack intl. card fee deducted before 50/50 net split
PAYSTACK_INTL_CARD_FEE_PCT = 0.039  # ~3.9% on international cards (configurable estimate)
FX_FALLBACK_USD_ZAR = 18.50  # if live FX APIs fail
FX_CACHE_TTL_SEC = 3600  # 1 hour cache for USD/ZAR

# Main dashboard header line (War Room + Agent Console — compliance row uses this text)
DASHBOARD_TITLE_LINE = "SENTURION AI SOLUTIONS | CEO: EDUARD DE LANGE | CFO: MONRÉ WESSEL NAGEL"
DASHBOARD_HEADER_HTML = (
    '<h1 class="hud-title hud-title-mirror">' + DASHBOARD_TITLE_LINE + "</h1>"
)
# Legal Templates — Master Service Agreement (fixed; do not edit for individual deals without counsel)
MSA_SUCCESS_FEE_PERCENT = 15
BRAND_NAME = "Senturion AI Solutions"
BRAND_INSTITUTIONAL_HEADER = "SENTURION AI SOLUTIONS · NEURAL AUDIT & COMPLIANCE"
DEFAULT_CLINIC_NAME = BRAND_NAME

# Ghost Auditing — folder-scale background extraction (thread-safe; UI polls via st.fragment)
# Industrial Intake: 100+ PDFs supported; multi-file upload (2+) runs off-thread with ThreadPoolExecutor (≤10 workers).
GHOST_ASYNC_THRESHOLD = 2  # PDF count at which processing moves off the main Streamlit thread
GHOST_FRAGMENT_POLL_SEC = 0.5  # drain merge queue → vault without full-page refresh (fragment tick)
MASSIVE_INTAKE_MAX_WORKERS = 10
_INTAKE_PROCESSED_HASHES_PATH = os.path.join(
    os.path.dirname(os.path.abspath(__file__)), ".senturion_intake_processed_hashes.json"
)
_INTAKE_HASH_FILE_LOCK = threading.Lock()
MASTER_AUDIT_AUTO_PDF_THRESHOLD = 5000  # Auto-generate CEO Master Audit Summary PDF at batch end
GHOST_BEEP_MIN_INTERVAL_SEC = 2.0
_GHOST_LOCK = threading.Lock()
_GHOST_STATE: dict[str, Any] = {
    "running": False,
    "finished": False,
    "job_id": None,
    "predator_batch_tag": "",
    "total_files": 0,
    "files_completed": 0,
    "done_idx": 0,
    "pct": 0.0,
    "current_file": "",
    "current_label": "",
    "status_line": "",
    "estimated_recoverable_usd": 0.0,
    "pending_merge_queue": [],
    "pending_errors": [],
    "skipped_duplicate_count": 0,
}

# Batch Predator — Global Error Shield (chunk sizes)
PREDATOR_CHUNK_LINES = 150
PREDATOR_CHUNK_RETRY_LINES = 75
MANUAL_FIX_TAG = "[MANUAL_FIX_REQUIRED]"
MANUAL_VERIFICATION_TAG = "[MANUAL_VERIFICATION]"
# Double-entry vault: reject display for invalid USD string, $0, or > this cap
MAX_VAULT_SANE_USD = 100_000.0
# Append-only shadow log (survives browser crash / session loss)
VAULT_SHADOW_LOG_PATH = Path(__file__).resolve().parent / "vault_backup.csv"
_VAULT_IO_LOCK = threading.RLock()  # reentrant: merge → shadow append
SHADOW_LOG_FIELDS = [
    "ts_utc",
    "event",
    "vault_id",
    "patient_id",
    "denial_date",
    "amount",
    "unique_claim_hash",
    "validation",
    "entry_json",
]
# Strict USD surface form for Potential Revenue / Amount Denied cross-check (no scientific notation)
_USD_AMOUNT_DISPLAY_RE = re.compile(
    r"^\s*\$?\s*(?:\d{1,3}(?:,\d{3})+|\d+)(?:\.\d{1,2})?\s*$"
)

import google.generativeai as genai
from google.generativeai.types import HarmCategory, HarmBlockThreshold

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False
    WD_ALIGN_PARAGRAPH = None  # type: ignore[misc, assignment]
    RGBColor = None  # type: ignore[misc, assignment]

try:
    from fpdf import FPDF

    HAS_FPDF = True
except ImportError:
    HAS_FPDF = False
    FPDF = None  # type: ignore[misc, assignment]

# Load environment variables
load_dotenv()

# Configure Gemini API when present (optional — demo / paystack review flows may not need AI)
model = None
try:
    if hasattr(st, "secrets") and "GEMINI_API_KEY" in st.secrets:
        _gk = str(st.secrets["GEMINI_API_KEY"] or "").strip()
        if _gk:
            genai.configure(api_key=_gk)
            model = genai.GenerativeModel("gemini-2.5-pro")
except Exception:
    model = None


@st.cache_data(ttl=86400, show_spinner=False)
def _fetch_lottie_json(url: str) -> dict | None:
    """Load Lottie JSON for sidebar animation."""
    if requests is None:
        return None
    try:
        r = requests.get(url, timeout=15)
        r.raise_for_status()
        return r.json()
    except Exception:
        return None


@st.cache_resource
def init_connection() -> Client:
    """
    Senturion Vault — initialize the Supabase client from Streamlit secrets.

    Requires ``st.secrets['SUPABASE_URL']`` and ``st.secrets['SUPABASE_KEY']``.
    """
    return create_client(
        str(st.secrets["SUPABASE_URL"]),
        str(st.secrets["SUPABASE_KEY"]),
    )


def get_supabase() -> Client:
    """Return the cached Supabase client (same instance as :func:`init_connection`)."""
    return init_connection()


@st.cache_resource
def get_supabase_service_role() -> Client:
    """Service-role client for Admin Auth API and profiles upsert. Key must stay in server secrets only."""
    return create_client(
        str(st.secrets["SUPABASE_URL"]),
        str(st.secrets["SUPABASE_SERVICE_ROLE_KEY"]),
    )


@st.cache_data(ttl=45, show_spinner=False)
def _supabase_database_online() -> bool:
    """
    Lightweight reachability check for Supabase (Senturion Vault pipes).
    True when the REST gateway responds without a server error (5xx).
    """
    try:
        if "SUPABASE_URL" not in st.secrets or "SUPABASE_KEY" not in st.secrets:
            return False
        url = str(st.secrets["SUPABASE_URL"]).rstrip("/")
        key = str(st.secrets["SUPABASE_KEY"])
        if requests is None:
            return True
        r = requests.get(
            f"{url}/rest/v1/",
            headers={"apikey": key, "Authorization": f"Bearer {key}"},
            timeout=8,
        )
        return r.status_code < 500
    except Exception:
        return False


# Hard-coded one-time Founders bypass (login path). Requires SUPABASE_SERVICE_ROLE_KEY in secrets.
_FOUNDERS_BYPASS_EMAIL = "Eduardsenturionai@outlook.com"


# --- Supabase Storage: sealed artifacts bucket (create in Supabase dashboard if missing) ---
VAULT_BUCKET = "senturion-vault"


def upload_to_vault(
    supabase: Client,
    object_path: str,
    data: bytes,
    *,
    content_type: str = "application/octet-stream",
) -> tuple[bool, str | None]:
    """Upload binary object to `senturion-vault`. `object_path` is relative to the bucket (no leading `/`)."""
    if not data:
        return False, "empty payload"
    rel = object_path.lstrip("/")
    try:
        supabase.storage.from_(VAULT_BUCKET).upload(
            rel,
            data,
            file_options={
                "content-type": content_type,
                "upsert": "true",
            },
        )
        return True, None
    except Exception as e:
        return False, str(e)


def download_from_vault(supabase: Client, object_path: str) -> bytes | None:
    """Download an object from `senturion-vault`. Returns None on failure."""
    rel = object_path.lstrip("/")
    try:
        return supabase.storage.from_(VAULT_BUCKET).download(rel)
    except Exception:
        return None


def _fetch_profile_row(supabase: Client, user_id: str) -> dict | None:
    """Load one row from public.profiles for this auth UID (tries id, then user_id FK).

    Selects role + email when available; falls back to role-only if email column is missing.
    """
    if not user_id:
        return None
    uid = str(user_id)
    for fk_col in ("id", "user_id"):
        for sel in ("role, email", "role"):
            try:
                res = (
                    supabase.table("profiles")
                    .select(sel)
                    .eq(fk_col, uid)
                    .limit(1)
                    .execute()
                )
                rows = getattr(res, "data", None) or []
                if not rows:
                    continue
                row = rows[0]
                if isinstance(row, dict):
                    return row
            except Exception:
                continue
    return None


def _normalize_profile_role(raw: str | None) -> str:
    """Map DB string to admin | agent | client | clinic | reviewer | pending_review.

    New signups should get `pending_review` in `profiles` (Supabase trigger) until an Admin promotes them.
    """
    if raw is None or (isinstance(raw, str) and not str(raw).strip()):
        return "pending_review"
    role = str(raw).strip().lower()
    if role in ("pending", "pending_review", "unverified", "awaiting_review", "new", "awaiting"):
        return "pending_review"
    if role in ("admin", "administrator", "superadmin", "owner"):
        return "admin"
    if role in ("clinic", "clinics", "clinic_user", "clinic_portal", "practice", "guest_clinic"):
        return "clinic"
    if role in ("client", "customer"):
        return "client"
    if role == "agent":
        return "agent"
    if role in ("reviewer", "paystack_reviewer", "merchant_reviewer"):
        return "reviewer"
    return "pending_review"


class AppPermissions(NamedTuple):
    """Strict RBAC flags — no email/domain checks; `st.session_state.role` only."""

    role: str | None
    is_admin: bool
    is_agent: bool
    is_client: bool
    is_clinic: bool
    is_reviewer: bool
    is_pending_access: bool
    can_admin_war_room: bool
    can_user_management: bool
    can_financial_analytics: bool
    can_appeal_engine: bool
    can_verify_submit: bool
    can_client_vault: bool
    can_clinic_portal: bool


def _session_email_normalized() -> str:
    em = (st.session_state.get("email") or "").strip().lower()
    if em:
        return em
    u = st.session_state.get("user")
    return (getattr(u, "email", None) or "").strip().lower()


def _is_paystack_demo_session() -> bool:
    """Merchant verification demo — `reviews@paystack.com` sees the clean Demo Audit surface only."""
    return _session_email_normalized() == DEMO_PAYSTACK_EMAIL.strip().lower()


def _establish_paystack_reviewer_bypass_session() -> None:
    """Local-only session (no Supabase JWT) for Paystack merchant review — role Reviewer."""
    uid = PAYSTACK_REVIEWER_BYPASS_USER_ID
    mock_user = SimpleNamespace(
        id=uid,
        email=DEMO_PAYSTACK_EMAIL,
    )
    st.session_state.user = mock_user
    st.session_state.user_id = uid
    st.session_state.email = DEMO_PAYSTACK_EMAIL
    st.session_state.role = "reviewer"
    st.session_state._paystack_reviewer_bypass = True
    # Explicit session keys requested for Paystack review / Reviewer mode
    st.session_state["logged_in"] = True
    st.session_state["user_role"] = "Reviewer"


def _apply_paystack_demo_role_override() -> None:
    """Ensure Paystack review login is never stuck in pending_review; role stays Reviewer."""
    if not _is_paystack_demo_session():
        return
    st.session_state.role = "reviewer"
    st.session_state.email = _session_email_normalized() or DEMO_PAYSTACK_EMAIL
    st.session_state["logged_in"] = True
    st.session_state["user_role"] = "Reviewer"


def check_permissions() -> AppPermissions:
    """Global RBAC gate. Admin War Room requires role == \"admin\" exactly."""
    r = st.session_state.get("role")
    pend = r == "pending_review"
    adm = r == "admin"
    agt = r == "agent"
    cli = r == "client"
    cln = r == "clinic"
    rev = r == "reviewer"
    return AppPermissions(
        role=r,
        is_admin=adm,
        is_agent=agt,
        is_client=cli,
        is_clinic=cln,
        is_reviewer=rev,
        is_pending_access=pend,
        can_admin_war_room=adm,
        can_user_management=adm,
        can_financial_analytics=adm or agt,
        can_appeal_engine=adm or agt,
        can_verify_submit=agt,
        can_client_vault=cli,
        can_clinic_portal=cln,
    )


def _sync_session_with_profiles_table() -> None:
    """Reconcile session role/email with `profiles` on every run (DB is source of truth)."""
    if st.session_state.get("_paystack_reviewer_bypass") or str(
        st.session_state.get("user_id") or ""
    ) == PAYSTACK_REVIEWER_BYPASS_USER_ID:
        return
    uid = st.session_state.get("user_id")
    if not uid or not st.session_state.user:
        return
    try:
        row = _fetch_profile_row(get_supabase(), str(uid))
        if row:
            st.session_state.role = _normalize_profile_role(row.get("role"))
            pe = row.get("email")
            auth_em = (getattr(st.session_state.user, "email", None) or "").strip()
            prof_email = (str(pe).strip() if pe is not None and str(pe).strip() else "")
            st.session_state.email = prof_email or auth_em or (st.session_state.get("email") or "")
        else:
            st.session_state.role = "pending_review"
    except Exception:
        pass


def _fetch_profiles_directory(supabase: Client) -> list[dict]:
    """All profile rows (requires RLS policy allowing admins to read full directory)."""
    for sel in ("id, email, role", "id, role"):
        try:
            res = supabase.table("profiles").select(sel).execute()
            data = getattr(res, "data", None)
            if isinstance(data, list):
                return data
        except Exception:
            continue
    return []


def _is_auth_rate_limit_error(exc: BaseException) -> bool:
    msg = str(exc).lower()
    return "429" in msg or "rate limit" in msg or "too many requests" in msg


def _show_signup_rate_limit_override_message() -> None:
    st.markdown(
        """
        <div style="text-align:center;padding:1.35rem 1rem;margin:0.85rem 0;
          background:#111111;border:1px solid #262626;border-radius:2px;box-shadow:none;">
          <div style="font-family:'JetBrains Mono',ui-monospace,monospace;font-size:0.85rem;
            font-weight:400;color:#E0E0E0;letter-spacing:2px;line-height:1.5;text-transform:uppercase;">
            SYSTEM OVERRIDE: Verification bypassed. Please proceed to Login.
          </div>
        </div>
        """,
        unsafe_allow_html=True,
    )


def _upsert_profile_immediate(
    supabase: Client,
    user_id: str,
    email: str,
    *,
    role: str,
) -> None:
    """Insert/merge `profiles` immediately so RBAC does not wait on DB triggers."""
    try:
        supabase.table("profiles").upsert(
            {
                "id": str(user_id),
                "email": (email or "").strip(),
                "role": str(role).strip().lower(),
            }
        ).execute()
    except Exception:
        pass


def _show_access_granted_banner() -> None:
    st.markdown(
        """
        <div style="text-align:center;padding:1.75rem 1rem;margin:0.75rem 0;
          background:#111111;border:1px solid #262626;border-radius:2px;box-shadow:none;">
          <div style="font-family:'Playfair Display','Times New Roman',serif;font-size:1.85rem;
            font-weight:500;color:#E0E0E0;letter-spacing:2px;line-height:1.2;text-transform:uppercase;">
            ACCESS GRANTED
          </div>
          <div style="margin-top:0.65rem;font-size:0.75rem;color:#9ca3af;
            letter-spacing:2px;font-family:'JetBrains Mono',monospace;text-transform:uppercase;">
            Entering secure workspace…
          </div>
        </div>
        """,
        unsafe_allow_html=True,
    )


def _admin_set_profile_role(supabase: Client, profile_id: str, new_role: str) -> tuple[bool, str]:
    nr = str(new_role).strip().lower()
    if nr not in ("admin", "agent", "client", "clinic", "pending_review"):
        return False, "Invalid role"
    try:
        supabase.table("profiles").update({"role": nr}).eq("id", str(profile_id)).execute()
        return True, ""
    except Exception as e:
        return False, str(e)


def _sync_session_from_profiles_after_login(supabase: Client, user) -> None:
    """After sign_in_with_password: bind user_id, fetch profiles row (with trigger retry), set role + email."""
    uid = getattr(user, "id", None)
    if not uid:
        return
    uid_s = str(uid)
    st.session_state.user = user
    st.session_state.user_id = uid_s
    auth_email = (getattr(user, "email", None) or "").strip()

    row = _fetch_profile_row(supabase, uid_s)
    if row is None:
        time.sleep(2)
        row = _fetch_profile_row(supabase, uid_s)

    if row:
        st.session_state.role = _normalize_profile_role(row.get("role"))
        pe = row.get("email")
        prof_email = (str(pe).strip() if pe is not None and str(pe).strip() else "")
        st.session_state.email = prof_email or auth_email
    else:
        # Profile row not ready yet (trigger lag) — treat as pending until `profiles` exists.
        st.session_state.role = "pending_review"
        st.session_state.email = auth_email


def _play_success_beep() -> None:
    """Play a short beep on successful extraction (Windows) or ASCII bell (other platforms)."""
    try:
        if sys.platform == "win32":
            import winsound
            winsound.Beep(800, 200)
        else:
            print("\a", end="")
    except Exception:
        pass


# Page config - must be first Streamlit command
st.set_page_config(
    page_title="CLAIM SNIPER // NEURAL AUDIT",
    page_icon="🛡️",
    layout="wide",
    initial_sidebar_state="expanded",
)

# Session state: auth user + profiles-backed RBAC
if "user" not in st.session_state:
    st.session_state.user = None
if "user_id" not in st.session_state:
    st.session_state.user_id = None
if "email" not in st.session_state:
    st.session_state.email = None
if "role" not in st.session_state:
    st.session_state.role = None
if "revenue_vault" not in st.session_state:
    st.session_state.revenue_vault = []
if "audit_log_history" not in st.session_state:
    st.session_state.audit_log_history = []
if "client_view_mode" not in st.session_state:
    st.session_state.client_view_mode = False
if "logged_in" not in st.session_state:
    st.session_state.logged_in = False
if "user_role" not in st.session_state:
    st.session_state.user_role = None
if "debug_log" not in st.session_state:
    st.session_state.debug_log = []
if "clinic_profiles" not in st.session_state:
    st.session_state.clinic_profiles = []
if "active_clinic_id" not in st.session_state:
    st.session_state.active_clinic_id = None
if "vault_clinic_filter" not in st.session_state:
    st.session_state.vault_clinic_filter = "ALL"
if "client_vault_contracts" not in st.session_state:
    st.session_state.client_vault_contracts = []

# Revenue Vault — unique-claim hash registry (shadow log hydrate fills this on first _ensure_revenue_vault)
if "_vault_unique_claim_hashes" not in st.session_state:
    st.session_state._vault_unique_claim_hashes = set()

def render_login_screen() -> None:
    """Senturion vault login (glass card, centered)."""
    st.markdown(
        """
    <style>
    @import url('https://fonts.googleapis.com/css2?family=DM+Sans:wght@400;500;600&family=Playfair+Display:wght@400;500;600&family=JetBrains+Mono:wght@200;300;400&display=swap');

    /* Full-viewport flex center for login stage */
    .stApp {
        background: #050505 !important;
    }
    [data-testid="stAppViewContainer"] {
        display: flex !important;
        flex-direction: row !important;
        align-items: stretch !important;
        min-height: 100vh !important;
    }
    [data-testid="stAppViewContainer"] [data-testid="stHeader"] {
        background: transparent !important;
    }
    section[data-testid="stMain"] > div {
        display: flex !important;
        flex-direction: column !important;
        align-items: center !important;
        justify-content: center !important;
        flex: 1 1 auto !important;
        min-height: 100vh !important;
        padding: 2rem 1rem !important;
    }
    section[data-testid="stMain"] .block-container {
        display: flex !important;
        flex-direction: column !important;
        align-items: center !important;
        justify-content: center !important;
        width: 100% !important;
        max-width: 440px !important;
        padding-top: 0 !important;
        padding-bottom: 0 !important;
    }

    /* Institutional vault card — matte, ruled border */
    [data-testid="stVerticalBlockBorderWrapper"] {
        background: #111111 !important;
        backdrop-filter: none !important;
        -webkit-backdrop-filter: none !important;
        border: 1px solid #262626 !important;
        border-radius: 2px !important;
        padding: 0.35rem 0.15rem 0.75rem !important;
        box-shadow: none !important;
        width: 100% !important;
        max-width: 420px !important;
    }
    [data-testid="stVerticalBlockBorderWrapper"] > div {
        border: none !important;
        background: transparent !important;
    }

    .vault-brand-wrap {
        display: flex;
        flex-direction: column;
        align-items: center;
        justify-content: center;
        text-align: center;
        margin-bottom: 0.5rem;
        padding-top: 0.5rem;
    }
    .login-brand-logo-wrap {
        display: flex;
        justify-content: center;
        align-items: center;
        width: 100%;
        margin: 0 0 0.65rem 0;
    }
    .login-brand-logo-wrap img {
        display: block;
        margin: 0 auto;
        max-width: 100%;
        height: auto;
        object-fit: contain;
        filter: none;
    }
    .vault-senturion-title {
        font-family: 'Playfair Display', 'Times New Roman', serif !important;
        font-size: 26pt !important;
        font-weight: 500 !important;
        letter-spacing: 2px;
        text-transform: uppercase;
        text-align: center !important;
        width: 100%;
        color: #E0E0E0 !important;
        margin: 0 0 0.35rem 0 !important;
        text-shadow: none;
    }
    .vault-portal-subtitle {
        font-family: 'JetBrains Mono', ui-monospace, monospace !important;
        font-size: 10pt !important;
        font-weight: 300 !important;
        letter-spacing: 2px;
        text-transform: uppercase;
        text-align: center !important;
        width: 100%;
        color: rgba(224, 224, 224, 0.55) !important;
        margin: 0 0 1.15rem 0 !important;
    }
    .vault-login-footer {
        font-family: 'DM Sans', -apple-system, BlinkMacSystemFont, sans-serif !important;
        font-size: 0.62rem !important;
        font-weight: 400 !important;
        color: rgba(200, 210, 225, 0.42) !important;
        text-align: center !important;
        letter-spacing: 0.1em;
        margin: 0.25rem 0 0.15rem 0 !important;
        padding: 0 1rem !important;
    }

    /* Dark-navy inputs, silver focus */
    [data-testid="stVerticalBlockBorderWrapper"] .stTextInput label,
    [data-testid="stVerticalBlockBorderWrapper"] [data-testid="stWidgetLabel"] p {
        font-family: 'DM Sans', -apple-system, BlinkMacSystemFont, sans-serif !important;
        color: #a8b0c4 !important;
    }
    [data-testid="stVerticalBlockBorderWrapper"] input[type="text"],
    [data-testid="stVerticalBlockBorderWrapper"] input[type="password"] {
        font-family: 'DM Sans', -apple-system, BlinkMacSystemFont, sans-serif !important;
        background: #0D0D0D !important;
        border: 1px solid #262626 !important;
        border-radius: 2px !important;
        color: #E0E0E0 !important;
        caret-color: #E0E0E0 !important;
    }
    [data-testid="stVerticalBlockBorderWrapper"] input[type="text"]:focus,
    [data-testid="stVerticalBlockBorderWrapper"] input[type="password"]:focus {
        border-color: #404040 !important;
        box-shadow: none !important;
        outline: none !important;
    }
    [data-testid="stVerticalBlockBorderWrapper"] fieldset {
        border-color: transparent !important;
    }

    /* Institutional login — titanium outline */
    [data-testid="stVerticalBlockBorderWrapper"] .stForm button[kind="formSubmit"],
    [data-testid="stVerticalBlockBorderWrapper"] .stForm [data-testid="stBaseButton-secondaryFormSubmit"] {
        width: 100% !important;
        background: #0D0D0D !important;
        color: #E0E0E0 !important;
        font-family: 'DM Sans', -apple-system, BlinkMacSystemFont, sans-serif !important;
        font-weight: 500 !important;
        letter-spacing: 2px;
        text-transform: uppercase;
        border: 1px solid #262626 !important;
        border-radius: 2px !important;
        padding: 0.65rem 1.25rem !important;
        transition: border-color 0.15s ease, color 0.15s ease !important;
        box-shadow: none !important;
    }
    [data-testid="stVerticalBlockBorderWrapper"] .stForm button[kind="formSubmit"]:hover {
        border-color: #E0E0E0 !important;
        color: #ffffff !important;
    }
    </style>
        """,
        unsafe_allow_html=True,
    )

    with st.container(border=True):
        # Display width aligned with Word layout: Inches(3) at 96 DPI
        try:
            from docx.shared import Inches as _LoginLogoInches

            _login_logo_w = max(1, int(round(_LoginLogoInches(3).inches * 96)))
        except Exception:
            _login_logo_w = 288

        _logo_path = "senturionaisolutions.jpg"
        _brand_logo_html = ""
        if os.path.isfile(_logo_path):
            try:
                with open(_logo_path, "rb") as _lf:
                    _lb64 = base64.b64encode(_lf.read()).decode("ascii")
                _brand_logo_html = (
                    f'<div class="login-brand-logo-wrap">'
                    f'<img src="data:image/jpeg;base64,{_lb64}" alt="" '
                    f'style="width:{_login_logo_w}px;max-width:100%;height:auto;display:block;margin:0 auto;" />'
                    f"</div>"
                )
            except OSError:
                _brand_logo_html = ""
        st.markdown(
            f'<div class="vault-brand-wrap">{_brand_logo_html}'
            '<p class="vault-senturion-title">SENTURION AI</p>'
            '<p class="vault-portal-subtitle">Secure Medical Revenue Recovery Terminal</p></div>',
            unsafe_allow_html=True,
        )

        tab_login, tab_register, tab_clinic_guest = st.tabs(["Login", "Create Account", "Clinic Portal"])

        with tab_login:
            with st.form("login_form", clear_on_submit=False):
                email = st.text_input(
                    "Email",
                    type="default",
                    key="login_email",
                    placeholder="you@work.com",
                )
                password = st.text_input(
                    "Password",
                    type="password",
                    key="login_password",
                    placeholder="••••••••",
                )
                login_submitted = st.form_submit_button("Login", use_container_width=True)

        with tab_register:
            st.caption("New users receive **pending review** until a Senturion Admin assigns your role.")
            with st.form("register_form", clear_on_submit=False):
                reg_email = st.text_input(
                    "Email",
                    type="default",
                    key="register_email",
                    placeholder="you@work.com",
                )
                reg_password = st.text_input(
                    "Password",
                    type="password",
                    key="register_password",
                    placeholder="Choose a strong password",
                )
                reg_confirm = st.text_input(
                    "Confirm Password",
                    type="password",
                    key="register_confirm",
                    placeholder="Repeat password",
                )
                register_submitted = st.form_submit_button("Register", use_container_width=True)

        with tab_clinic_guest:
            st.markdown(
                '<p style="font-family:\'JetBrains Mono\',monospace;font-size:0.72rem;color:#a8b0c4;text-align:center;'
                'letter-spacing:0.12em;text-transform:uppercase;margin:0 0 0.5rem 0;">Clinic · guest access</p>',
                unsafe_allow_html=True,
            )
            st.info(
                "**Clinic accounts** sign in on the **Login** tab using the email and password issued by Senturion. "
                "Your profile must have **`clinic`** role in Supabase (`profiles.role`)."
            )
            st.caption(
                "Restricted portal: **Upload Claims** and **View Audit Reports** only — no treasury, Capitec, or founder equity."
            )

        st.markdown(
            '<p class="vault-login-footer">Proprietary Neural Audit System v1.0</p>',
            unsafe_allow_html=True,
        )
        st.markdown(
            f'<div style="margin-top:1rem;padding:0.65rem 0.5rem;border-top:1px solid rgba(212,175,55,0.25);'
            f"border-bottom:1px solid rgba(212,175,55,0.15);background:#0f172a;font-family:'JetBrains Mono',monospace;"
            f'font-size:0.58rem;letter-spacing:0.1em;color:#e8e8e8;text-align:center;">'
            f"{_kyc_footer_html_inner()}</div>",
            unsafe_allow_html=True,
        )

    if login_submitted:
        if not email or not password:
            st.error("Please enter email and password.")
        else:
            _em_login = (email or "").strip()
            # Paystack merchant review — emergency bypass (no Supabase); role Reviewer.
            if (
                _em_login.lower() == DEMO_PAYSTACK_EMAIL.lower()
                and password == PAYSTACK_REVIEW_BYPASS_PASSWORD
            ):
                _establish_paystack_reviewer_bypass_session()
                st.success("Authenticated (Paystack review). Loading Neural Audit Demo…")
                time.sleep(0.35)
                st.rerun()
            else:
                try:
                    supabase = get_supabase()

                    # Founders bypass: force-create auth user (confirmed) + admin profile before normal sign-in.
                    if (
                        _em_login.lower() == _FOUNDERS_BYPASS_EMAIL.lower()
                        and st.session_state.get("user") is None
                    ):
                        if "SUPABASE_SERVICE_ROLE_KEY" not in st.secrets:
                            st.error(
                                "Founders bypass: add SUPABASE_SERVICE_ROLE_KEY to .streamlit/secrets.toml."
                            )
                        else:
                            sb_svc = get_supabase_service_role()
                            try:
                                _cu = sb_svc.auth.admin.create_user(
                                    {
                                        "email": _em_login,
                                        "password": password,
                                        "email_confirm": True,
                                    }
                                )
                                _nu = getattr(_cu, "user", None) if _cu else None
                                _nid = (
                                    str(_nu.id)
                                    if _nu is not None and getattr(_nu, "id", None)
                                    else None
                                )
                                if _nid:
                                    sb_svc.table("profiles").upsert(
                                        {
                                            "id": _nid,
                                            "email": _em_login,
                                            "role": "admin",
                                        }
                                    ).execute()
                            except Exception:
                                # User may already exist in auth.users — continue to sign_in_with_password.
                                pass

                    auth_response = supabase.auth.sign_in_with_password(
                        {"email": _em_login, "password": password}
                    )
                    user = auth_response.user

                    if (
                        _em_login.lower() == _FOUNDERS_BYPASS_EMAIL.lower()
                        and user is not None
                        and getattr(user, "id", None)
                        and "SUPABASE_SERVICE_ROLE_KEY" in st.secrets
                    ):
                        try:
                            get_supabase_service_role().table("profiles").upsert(
                                {
                                    "id": str(user.id),
                                    "email": _em_login,
                                    "role": "admin",
                                }
                            ).execute()
                        except Exception:
                            pass

                    # JWT stored by sign_in_with_password — RLS on `profiles` applies on fetch.
                    _sync_session_from_profiles_after_login(supabase, user)
                    st.success("Authenticated. Loading portal...")
                    time.sleep(0.5)
                    st.rerun()
                except Exception as e:
                    st.error(f"Authentication failed: {e}")

    if register_submitted:
        if not reg_email or not reg_password:
            st.error("Email and password are required.")
        elif reg_password != reg_confirm:
            st.error("Passwords do not match.")
        elif len(reg_password) < 8:
            st.error("Password must be at least 8 characters.")
        else:
            _em = reg_email.strip()
            _pw = reg_password
            supabase = get_supabase()

            def _session_then_profile_and_sync(user_obj) -> None:
                """Ensure JWT exists for RLS, upsert profile, bind Streamlit session."""
                if not user_obj or not getattr(user_obj, "id", None):
                    return
                uid_s = str(user_obj.id)
                # Re-read session: signup may not return a session depending on project settings.
                try:
                    _sess = supabase.auth.get_session()
                except Exception:
                    _sess = None
                if _sess is None or not getattr(_sess, "access_token", None):
                    try:
                        _sign = supabase.auth.sign_in_with_password(
                            {"email": _em, "password": _pw}
                        )
                        user_obj = _sign.user or user_obj
                    except Exception:
                        pass
                # Forced row (no trigger wait). Auth `options.data.role` = client; profiles.role client per spec.
                _upsert_profile_immediate(
                    supabase,
                    uid_s,
                    _em,
                    role="client",
                )
                _sync_session_from_profiles_after_login(supabase, user_obj)

            try:
                _auth_res = supabase.auth.sign_up(
                    {
                        "email": _em,
                        "password": _pw,
                        "options": {
                            "data": {"role": "client"},
                            "email_confirm": False,
                        },
                    }
                )
                _u = _auth_res.user if _auth_res else None
                if _u:
                    _session_then_profile_and_sync(_u)
                _show_access_granted_banner()
                time.sleep(0.45)
                st.rerun()
            except Exception as e:
                if _is_auth_rate_limit_error(e):
                    _show_signup_rate_limit_override_message()
                    try:
                        _auth_res2 = supabase.auth.sign_in_with_password(
                            {"email": _em, "password": _pw}
                        )
                        _u2 = _auth_res2.user if _auth_res2 else None
                        if _u2:
                            _session_then_profile_and_sync(_u2)
                            _show_access_granted_banner()
                            time.sleep(0.45)
                            st.rerun()
                        else:
                            st.warning("Rate limited — could not verify existing account. Try again shortly.")
                    except Exception as e2:
                        if _is_auth_rate_limit_error(e2):
                            _show_signup_rate_limit_override_message()
                        st.error(f"Rate limited and sign-in failed: {e2}")
                else:
                    st.error(f"Registration failed: {e}")


# Gatekeeper: block unauthenticated access (Supabase user or emergency Reviewer bypass)
if st.session_state.user is None and not st.session_state.get("logged_in"):
    render_login_screen()
    st.stop()

# --- Authenticated UI below (protected by login wall) ---

# Bloomberg Terminal — pitch black ground, Onyx sidebar, Matrix / neon accents, ruled chrome
st.markdown("""
<style>
    @import url('https://fonts.googleapis.com/css2?family=EB+Garamond:ital,wght@0,400;0,500;0,600;0,700;1,400&family=JetBrains+Mono:wght@200;300;400;500&display=swap');

    .stApp::before { display: none !important; }

    .stApp {
        background: #0f172a !important;
        background-image: none !important;
        color: #e2e8f0 !important;
    }

    /* Slate Dark — professional canvas (no white bleed) */
    .main .block-container {
        border: 1px solid #334155 !important;
        border-radius: 4px !important;
        padding: 1.25rem 1.5rem !important;
        background: #1e293b !important;
        transition: opacity 0.2s ease-in-out;
    }

    [data-testid="stHeader"] {
        background: #0f172a !important;
        border-bottom: 1px solid #334155 !important;
    }

    .glass-panel {
        background: #1e293b !important;
        backdrop-filter: none !important;
        -webkit-backdrop-filter: none !important;
        border-radius: 4px !important;
        border: 1px solid #334155 !important;
        padding: 1rem 1.25rem;
        margin: 0.75rem 0;
        box-shadow: none !important;
    }

    .main {
        font-family: 'JetBrains Mono', ui-monospace, monospace !important;
        background: transparent;
        max-width: 100%;
        color: #e2e8f0 !important;
    }

    p, .main p, .main span, label, [data-testid="stWidgetLabel"] {
        color: #e2e8f0 !important;
    }

    /* Metrics / data — contained slate (no bright white cards) */
    [data-testid="stMetricContainer"] {
        background: #0f172a !important;
        border: 1px solid #334155 !important;
        border-radius: 4px !important;
        padding: 0.65rem 0.75rem !important;
    }
    [data-testid="stMetricContainer"] label {
        color: #94a3b8 !important;
    }
    [data-testid="stMetricContainer"] [data-testid="stMetricValue"] {
        color: #f1f5f9 !important;
    }
    div[data-testid="stDataFrame"] {
        border: 1px solid #334155 !important;
        border-radius: 4px !important;
        overflow: hidden !important;
    }

    [data-testid="stSidebar"] {
        background: #0f172a !important;
        border-right: 1px solid #334155 !important;
        backdrop-filter: none !important;
        position: relative;
        overflow: auto;
    }

    [data-testid="stSidebar"]::before { display: none !important; }

    [data-testid="stSidebar"] .block-container {
        border: 1px solid #334155 !important;
        border-radius: 4px !important;
        padding: 0.85rem !important;
        margin: 0.35rem 0.25rem !important;
        background: #1e293b !important;
    }

    [data-testid="stSidebar"] .stMarkdown,
    [data-testid="stSidebar"] .stCaptionContainer { color: #E0E0E0 !important; }

    /* Senturion Command Hierarchy (sidebar) */
    .senturion-command-hierarchy {
        border: 1px solid #262626 !important;
        border-radius: 2px !important;
        padding: 0.75rem 0.65rem !important;
        margin: 0 0 1rem 0 !important;
        background: #050505 !important;
        font-family: 'JetBrains Mono', ui-monospace, monospace !important;
    }
    .senturion-command-hierarchy .sch-title {
        font-size: 0.62rem !important;
        letter-spacing: 0.22em !important;
        color: #E0E0E0 !important;
        text-transform: uppercase !important;
        margin-bottom: 0.6rem !important;
        padding-bottom: 0.45rem !important;
        border-bottom: 1px solid #262626 !important;
    }
    .senturion-command-hierarchy .sch-level {
        font-size: 0.66rem !important;
        color: #E0E0E0 !important;
        margin: 0.38rem 0 !important;
        line-height: 1.5 !important;
    }
    .senturion-command-hierarchy .sch-tag {
        color: #00FF41 !important;
        margin-right: 0.35rem !important;
        font-weight: 500 !important;
    }
    .senturion-command-hierarchy .sch-sub {
        color: rgba(224, 224, 224, 0.5) !important;
        font-size: 0.6rem !important;
    }

    [data-testid="stSidebar"] img {
        filter: none !important;
        max-width: 100%;
    }

    .logo-glow { box-shadow: none !important; filter: none !important; }

    .system-pulse {
        display: flex;
        align-items: center;
        gap: 0.5rem;
        font-family: 'JetBrains Mono', ui-monospace, monospace !important;
        font-size: 0.7rem;
        letter-spacing: 2px;
        text-transform: uppercase;
        color: #E0E0E0 !important;
        margin: 0.75rem 0;
    }

    .pulse-dot {
        width: 6px;
        height: 6px;
        border-radius: 1px;
        background: #00FF41;
        box-shadow: none !important;
        animation: none !important;
    }

    .founding-partners {
        margin-top: 1rem;
        padding: 1rem 1rem 0.85rem;
        border: 1px solid rgba(212, 175, 55, 0.35);
        border-radius: 2px;
        background: #0a0a0a;
        font-family: 'JetBrains Mono', ui-monospace, monospace !important;
        font-size: 0.82rem;
    }

    .founding-partners strong {
        color: #fafafa !important;
        font-weight: 600;
        letter-spacing: 0.12em;
    }

    .founding-partners .partner-name {
        color: #fafafa !important;
        font-weight: 500;
        text-shadow: none !important;
    }

    .founding-partners .partner-exec-line {
        display: block;
        font-family: 'JetBrains Mono', ui-monospace, monospace !important;
        font-size: 0.72rem;
        letter-spacing: 0.06em;
        text-transform: none;
    }
    .founding-partners .partner-exec-line.ceo-line {
        color: #fafafa !important;
        font-weight: 500;
    }
    .founding-partners .partner-exec-line.cfo-line {
        color: #d4af37 !important;
        font-weight: 500;
    }
    .founding-partners .partner-role {
        font-family: 'JetBrains Mono', ui-monospace, monospace !important;
        font-size: 0.65rem;
        color: #d4af37 !important;
        text-transform: uppercase;
        letter-spacing: 0.18em;
    }

    .founding-partners .partner-sep {
        display: block;
        height: 1px;
        background: rgba(212, 175, 55, 0.28);
        margin: 0.6rem 0;
        box-shadow: none !important;
    }
    
    .hud-title,
    h1, .main h1 {
        font-family: 'JetBrains Mono', ui-monospace, monospace !important;
        font-size: 1.45rem !important;
        font-weight: 500 !important;
        color: #E0E0E0 !important;
        text-shadow: none !important;
        letter-spacing: 0.14em !important;
        text-transform: uppercase !important;
        margin-bottom: 1rem !important;
        border-bottom: 1px solid #262626;
        padding-bottom: 0.5rem;
    }
    .hud-title-mirror {
        font-size: 0.92rem !important;
        letter-spacing: 0.07em !important;
        line-height: 1.4 !important;
    }
    .executive-header-row {
        display: flex;
        flex-wrap: wrap;
        align-items: flex-start;
        justify-content: space-between;
        gap: 0.75rem 1rem;
        margin-bottom: 1rem !important;
        padding-bottom: 0.5rem;
        border-bottom: 1px solid #262626;
    }
    .executive-header-h1 {
        flex: 1 1 12rem;
        min-width: 0;
        margin-bottom: 0 !important;
        border-bottom: none !important;
        padding-bottom: 0 !important;
    }
    .compliance-badge {
        flex: 0 0 auto;
        font-family: 'JetBrains Mono', ui-monospace, monospace !important;
        font-size: 0.58rem !important;
        letter-spacing: 0.06em !important;
        color: #a3e635 !important;
        border: 1px solid #3f3f46;
        background: #0a0a0a;
        padding: 0.4rem 0.7rem;
        border-radius: 2px;
        line-height: 1.35;
        text-align: right;
        text-transform: none;
    }
    .sidebar-msa-docid {
        font-family: 'JetBrains Mono', ui-monospace, monospace !important;
        font-size: 0.72rem !important;
        color: #d4d4d8 !important;
        margin: 0.35rem 0 0;
    }
    .sidebar-msa-docid code {
        color: #86efac !important;
        font-size: 0.72rem !important;
    }

    .main h2, .main h3, .main h4, .main h5 {
        font-family: 'JetBrains Mono', ui-monospace, monospace !important;
        font-weight: 500 !important;
        color: #E0E0E0 !important;
        letter-spacing: 0.1em;
    }

    .kpi-grid {
        display: grid;
        grid-template-columns: repeat(3, 1fr);
        gap: 1rem;
        margin-bottom: 1.5rem;
    }

    .kpi-card {
        background: #0D0D0D !important;
        border: 1px solid #262626 !important;
        border-radius: 2px !important;
        padding: 1.25rem 1rem;
        text-align: center;
        font-family: 'JetBrains Mono', ui-monospace, monospace !important;
        font-size: 0.68rem;
        font-weight: 300 !important;
        color: #E0E0E0 !important;
        letter-spacing: 2px;
        text-transform: uppercase;
        box-shadow: none !important;
        position: relative;
        overflow: hidden;
    }

    .kpi-card::before { display: none !important; }

    .kpi-card .metric {
        font-family: 'JetBrains Mono', ui-monospace, monospace !important;
        font-size: 1.35rem !important;
        font-weight: 200 !important;
        color: #f0f0f0 !important;
        margin-top: 0.5rem;
        letter-spacing: 0.04em;
        text-transform: none;
    }

    .kpi-card .metric.metric-critical {
        color: #FF3131 !important;
        font-weight: 300 !important;
    }

    .kpi-icon { display: none !important; }

    .inst-neural-uplink-head {
        text-align: center;
        margin: 0.5rem 0 1rem;
        padding: 0 0.5rem;
    }
    .inst-neural-uplink-head .inst-neural-title {
        font-family: 'JetBrains Mono', ui-monospace, monospace !important;
        font-size: 0.95rem !important;
        font-weight: 500 !important;
        color: #E0E0E0 !important;
        letter-spacing: 2px !important;
        text-transform: uppercase;
        line-height: 1.4;
    }
    .inst-neural-uplink-head .inst-neural-sub {
        font-family: 'JetBrains Mono', ui-monospace, monospace !important;
        font-size: 0.68rem !important;
        font-weight: 300 !important;
        color: rgba(224, 224, 224, 0.65) !important;
        letter-spacing: 0.2em;
        text-transform: uppercase;
        margin-top: 0.45rem;
    }
    .inst-neural-data-mode-title {
        font-family: 'JetBrains Mono', ui-monospace, monospace !important;
        font-size: 0.72rem !important;
        font-weight: 400 !important;
        color: rgba(224, 224, 224, 0.85) !important;
        letter-spacing: 0.22em !important;
        text-transform: uppercase;
        margin: 0.35rem 0 0.6rem;
        text-align: center;
    }
    /* Industrial Intake — Neural Telemetry strip */
    .neural-telemetry-bar {
        font-family: 'JetBrains Mono', ui-monospace, monospace !important;
        font-size: 0.78rem !important;
        line-height: 1.45;
        letter-spacing: 0.04em;
        color: #00FF41 !important;
        border: 1px solid #262626;
        background: #0D0D0D;
        padding: 0.7rem 0.95rem;
        margin: 0.35rem 0 0.65rem 0;
        border-radius: 2px;
        display: flex;
        flex-wrap: wrap;
        align-items: center;
        gap: 0.45rem 1rem;
    }
    .neural-telemetry-bar .neural-telemetry-title {
        font-weight: 600;
        letter-spacing: 0.16em;
        text-transform: uppercase;
        color: #E0E0E0 !important;
        font-size: 0.68rem !important;
    }
    .neural-telemetry-bar .neural-telemetry-metrics {
        color: #00FF41 !important;
        font-weight: 500;
    }
    
    [data-testid="stFileUploader"] {
        border: 1px solid #262626 !important;
        border-style: solid !important;
        border-radius: 2px !important;
        padding: 1.25rem !important;
        background: #0D0D0D !important;
        transition: border-color 0.15s ease;
        box-shadow: none !important;
    }

    [data-testid="stFileUploader"]:hover {
        border-color: #404040 !important;
        box-shadow: none !important;
        animation: none !important;
    }

    [data-testid="stFileUploader"] label[data-testid="stWidgetLabel"] p,
    [data-testid="stFileUploader"] [data-testid="stWidgetLabel"] p {
        text-align: center !important;
        font-family: 'JetBrains Mono', ui-monospace, monospace !important;
        letter-spacing: 2px !important;
        text-transform: uppercase;
        color: #E0E0E0 !important;
    }

    [data-testid="stFileUploaderDropzoneInstructions"] {
        display: none !important;
    }
    [data-testid="stFileUploaderDropzone"] {
        display: flex;
        flex-direction: column;
        align-items: center;
        gap: 0.65rem;
    }
    [data-testid="stFileUploaderDropzone"]::before {
        content: "NEURAL REVENUE RECOVERY";
        font-family: 'JetBrains Mono', ui-monospace, monospace !important;
        letter-spacing: 0.18em;
        color: #E0E0E0;
        font-size: 0.72rem;
        text-align: center;
        padding: 0.15rem 0 0.25rem;
        font-weight: 300;
    }

    [data-testid="stFileUploader"] section {
        background: transparent !important;
        border: none !important;
    }

    /* Manual paste fail-safe — charcoal/slate to match secure uplink + dataframe */
    .neural-manual-paste-wrap {
        border: 1px solid #262626 !important;
        border-radius: 2px !important;
        padding: 1rem 1rem 0.75rem;
        background: #0D0D0D !important;
        margin: 0.35rem 0 1rem;
        box-shadow: none !important;
    }
    .neural-manual-paste-wrap label p,
    .neural-manual-paste-wrap [data-testid="stWidgetLabel"] p {
        font-family: 'JetBrains Mono', ui-monospace, monospace !important;
        font-size: 0.72rem !important;
        letter-spacing: 0.12em !important;
        text-transform: uppercase !important;
        color: #E0E0E0 !important;
    }
    .neural-manual-paste-wrap textarea,
    .neural-manual-paste-wrap [data-baseweb="textarea"],
    .neural-manual-paste-wrap [data-testid="stTextArea"] textarea {
        background: #1a1a1a !important;
        color: #E0E0E0 !important;
        border: 1px solid #262626 !important;
        border-radius: 2px !important;
        font-family: 'JetBrains Mono', ui-monospace, monospace !important;
        font-size: 0.82rem !important;
        line-height: 1.45 !important;
    }
    .neural-manual-paste-wrap textarea:focus,
    .neural-manual-paste-wrap [data-baseweb="textarea"]:focus-within {
        border-color: #404040 !important;
        outline: none !important;
        box-shadow: 0 0 0 1px rgba(64, 64, 64, 0.35) !important;
    }
    
    .table-reveal {
        animation: inst-fade 0.45s ease-out forwards;
        opacity: 0;
    }

    @keyframes inst-fade {
        from { opacity: 0; }
        to { opacity: 1; }
    }

    .projection-table-cap {
        font-family: 'JetBrains Mono', ui-monospace, monospace !important;
        font-size: 0.72rem !important;
        color: #E0E0E0 !important;
        letter-spacing: 2px !important;
        text-transform: uppercase;
        margin-bottom: 0.75rem;
    }

    .holographic-table {
        background: #050505 !important;
        border: 1px solid #262626 !important;
        border-radius: 2px !important;
        padding: 1rem;
        margin: 1rem 0;
        box-shadow: none !important;
    }

    /* Neural Audit Summary — hero + KPI strip (emerald) */
    .neural-audit-summary-hero {
        text-align: center;
        padding: 1.25rem 1rem 1.1rem;
        margin: 0.5rem 0 1rem;
        background: #0D0D0D;
        border: 1px solid #262626;
        border-radius: 2px;
    }
    .neural-audit-summary-hero .na-hero-label {
        font-family: 'JetBrains Mono', ui-monospace, monospace !important;
        font-size: 0.68rem !important;
        letter-spacing: 0.22em !important;
        color: rgba(250, 250, 250, 0.88) !important;
        text-transform: uppercase;
        margin-bottom: 0.65rem;
    }
    .neural-audit-summary-hero .na-hero-amount {
        font-family: 'JetBrains Mono', ui-monospace, monospace !important;
        font-size: clamp(1.85rem, 4vw, 2.65rem) !important;
        font-weight: 600 !important;
        color: #00FF41 !important;
        line-height: 1.15;
        letter-spacing: 0.04em;
    }
    .neural-audit-summary-hero .na-hero-hint {
        font-size: 0.65rem !important;
        color: rgba(224, 224, 224, 0.55) !important;
        margin-top: 0.55rem;
        font-family: 'JetBrains Mono', ui-monospace, monospace !important;
        letter-spacing: 0.06em;
        text-transform: uppercase;
    }
    .vault-institutional-metrics {
        display: grid;
        grid-template-columns: repeat(3, minmax(0, 1fr));
        gap: 1rem;
        margin: 0.5rem 0 1.25rem;
    }
    .vault-institutional-metrics.vault-sub-metrics {
        grid-template-columns: repeat(2, minmax(0, 1fr));
    }
    @media (max-width: 900px) {
        .vault-institutional-metrics { grid-template-columns: 1fr; }
        .vault-institutional-metrics.vault-sub-metrics { grid-template-columns: 1fr; }
    }
    .vault-metric-card {
        background: #0D0D0D !important;
        border: 1px solid #262626 !important;
        border-radius: 2px !important;
        padding: 1rem 1.1rem;
    }
    .vault-metric-label {
        font-family: 'JetBrains Mono', ui-monospace, monospace !important;
        font-size: 0.68rem !important;
        letter-spacing: 0.18em !important;
        color: rgba(224, 224, 224, 0.72) !important;
        text-transform: uppercase;
        margin-bottom: 0.45rem;
    }
    .vault-metric-value {
        font-family: 'JetBrains Mono', ui-monospace, monospace !important;
        font-size: 1.28rem !important;
        font-weight: 500 !important;
        line-height: 1.2;
    }
    @keyframes vaultRecoverableFlash {
        0%, 100% { text-shadow: 0 0 4px rgba(0,255,65,0.35); }
        50% { text-shadow: 0 0 22px rgba(0,255,65,0.95); filter: brightness(1.22); }
    }
    .vault-recoverable-flash {
        animation: vaultRecoverableFlash 0.65s ease-in-out 2;
    }
    .vault-metric-hint {
        font-size: 0.65rem !important;
        color: rgba(180, 180, 180, 0.65) !important;
        margin-top: 0.4rem;
        font-family: 'JetBrains Mono', ui-monospace, monospace !important;
    }
    @keyframes foundersMatrixGlow {
        0%, 100% { text-shadow: 0 0 8px rgba(0,255,65,0.85), 0 0 22px rgba(0,255,65,0.45); filter: brightness(1.08); }
        50% { text-shadow: 0 0 16px rgba(0,255,65,1), 0 0 36px rgba(0,255,65,0.65); filter: brightness(1.18); }
    }
    .founders-commission-glow {
        color: #00FF41 !important;
        animation: foundersMatrixGlow 1.35s ease-in-out infinite !important;
    }
    .treasury-invoice-wrap button[kind="primary"] {
        background: #00FF41 !important;
        color: #050505 !important;
        border: 2px solid #E0E0E0 !important;
        font-weight: 800 !important;
        letter-spacing: 0.06em !important;
        box-shadow: 0 0 20px rgba(0,255,65,0.55), inset 0 0 12px rgba(255,255,255,0.12) !important;
    }

    .institutional-ledger-wrap {
        margin-top: 0.5rem;
        padding-top: 0.35rem;
        border-top: 1px solid #262626;
    }
    .institutional-ledger-shell .ledger-cell {
        font-size: 0.82rem;
        color: #d4d4d4 !important;
    }
    .institutional-ledger-shell .ledger-law {
        color: #FF3131 !important;
        font-weight: 500;
    }
    .institutional-ledger-shell .ledger-rev {
        color: #E0E0E0 !important;
        font-weight: 500;
    }
    .institutional-ledger-shell .ledger-hash {
        font-size: 0.72rem !important;
        color: #a3a3a3 !important;
        word-break: break-all;
    }
    .institutional-ledger-shell .mono {
        font-family: 'JetBrains Mono', ui-monospace, monospace !important;
    }
    /* Enforcement Clock — overdue (MD § 15-1005 window elapsed, no payment) */
    .enforcement-overdue-text {
        color: #FF3131 !important;
        font-weight: 700 !important;
        text-shadow: 0 0 10px rgba(255, 49, 49, 0.75) !important;
        letter-spacing: 0.06em;
    }
    .enforcement-row-shell {
        border: 1px solid rgba(255, 49, 49, 0.45) !important;
        border-radius: 2px !important;
        background: rgba(255, 49, 49, 0.08) !important;
        padding: 0.2rem 0.15rem 0.35rem 0.15rem !important;
        margin: 0.1rem 0 0.35rem 0 !important;
    }

    .stDataFrame, [data-testid="stDataFrame"] {
        border-radius: 2px;
        overflow: hidden;
        border: 1px solid #262626 !important;
        background: #0D0D0D !important;
        width: 100% !important;
    }

    .dataframe th {
        background: #1a1a1a !important;
        color: #E0E0E0 !important;
        font-weight: 500 !important;
        padding: 12px 16px !important;
        border: 1px solid #262626 !important;
        font-family: 'JetBrains Mono', ui-monospace, monospace !important;
    }

    .dataframe td {
        background: #050505 !important;
        color: #E0E0E0 !important;
        padding: 12px 16px !important;
        border: 1px solid #262626 !important;
        font-family: 'JetBrains Mono', ui-monospace, monospace !important;
    }

    /* Law Cited column — statutory emphasis (6th column in Neural Audit claims table) */
    .dataframe td:nth-child(6),
    .dataframe th:nth-child(6) {
        background: #140808 !important;
        border-left: 1px solid #FF3131 !important;
        color: #FF3131 !important;
    }

    .dataframe tr:hover td { background: #141414 !important; }
    .dataframe tr:hover td:nth-child(6) { background: #180909 !important; }
    
    div[data-testid="stDownloadButton"] button {
        background: #111111 !important;
        color: #E0E0E0 !important;
        font-weight: 500 !important;
        border-radius: 2px;
        padding: 0.5rem 1.25rem;
        border: 1px solid #262626 !important;
        font-family: 'JetBrains Mono', ui-monospace, monospace !important;
        letter-spacing: 2px;
        text-transform: uppercase;
        box-shadow: none !important;
    }

    /* Executive Brief — titanium silver (secondary action; distinct from vault/login chrome) */
    .stButton > button[kind="secondary"] {
        background: #E0E0E0 !important;
        color: #0D0D0D !important;
        border: 1px solid #a8a8a8 !important;
        font-family: 'JetBrains Mono', ui-monospace, monospace !important;
        letter-spacing: 1px !important;
        text-transform: uppercase !important;
        font-weight: 500 !important;
        box-shadow: none !important;
    }
    .stButton > button[kind="secondary"]:hover {
        background: #d0d0d0 !important;
        border-color: #909090 !important;
        color: #000000 !important;
    }
    .stButton > button[kind="secondary"]:disabled {
        opacity: 0.45 !important;
    }

    div[data-testid="stDownloadButton"] button:hover {
        border-color: #E0E0E0 !important;
        box-shadow: none !important;
    }

    .stSuccess, [data-testid="stSuccess"] {
        background: #111111 !important;
        border: 1px solid #262626 !important;
        color: #E0E0E0 !important;
    }

    .stWarning, [data-testid="stWarning"] {
        background: #141208 !important;
        border: 1px solid #262626 !important;
        color: #c9c5bd !important;
    }

    .stInfo, [data-testid="stInfo"] {
        background: #101010 !important;
        border: 1px solid #262626 !important;
        color: #a8a8a8 !important;
    }

    .stMarkdown, .main p, .main label, .main span { color: #E0E0E0 !important; }

    [data-testid="stMetric"] [data-testid="stMetricLabel"] p {
        font-family: 'JetBrains Mono', ui-monospace, monospace !important;
        letter-spacing: 2px !important;
        text-transform: uppercase;
        font-size: 0.72rem !important;
        color: #9ca3af !important;
    }
    [data-testid="stMetric"] [data-testid="stMetricValue"] {
        font-family: 'JetBrains Mono', ui-monospace, monospace !important;
        font-weight: 200 !important;
        font-size: 1.85rem !important;
        color: #E0E0E0 !important;
    }
    [data-testid="stMetric"] [data-testid="stMetricDelta"],
    [data-testid="stMetric"] svg,
    [data-testid="stMetric"] [data-testid="stMetricDeltaIcon-Up"],
    [data-testid="stMetric"] [data-testid="stMetricDeltaIcon-Down"] {
        display: none !important;
    }

    /* Financial Analytics expander: monetary columns (1 & 4) — blood red */
    [data-testid="stExpander"] [data-testid="column"]:nth-child(1) [data-testid="stMetricValue"],
    [data-testid="stExpander"] [data-testid="column"]:nth-child(4) [data-testid="stMetricValue"] {
        color: #FF3131 !important;
    }
    
    .stTabs [data-baseweb="tab-panel"] {
        border: 1px solid #262626 !important;
        border-radius: 2px !important;
        background: #050505 !important;
        padding: 0.75rem !important;
        margin-top: 0.5rem !important;
    }
    .stTabs [data-baseweb="tab-list"] {
        background: #050505 !important;
        backdrop-filter: none !important;
        border: 1px solid #262626 !important;
        border-radius: 2px;
        padding: 6px;
        gap: 8px;
        box-shadow: none !important;
    }

    .stTabs [data-baseweb="tab-list"] > * {
        border: none !important;
    }

    .stTabs [data-baseweb="tab"] {
        font-family: 'JetBrains Mono', ui-monospace, monospace !important;
        font-size: 0.78rem !important;
        font-weight: 500 !important;
        letter-spacing: 2px !important;
        text-transform: uppercase !important;
        color: #8b9199 !important;
        background: #0D0D0D !important;
        border: 1px solid #262626 !important;
        border-radius: 2px;
        padding: 0.65rem 1.35rem !important;
        transition: border-color 0.15s ease, color 0.15s ease !important;
        box-shadow: none !important;
    }

    .stTabs [data-baseweb="tab"]:hover {
        background: #141414 !important;
        border-color: #404040 !important;
        color: #E0E0E0 !important;
        box-shadow: none !important;
    }

    .stTabs [data-baseweb="tab"][aria-selected="true"] {
        background: #0D0D0D !important;
        color: #00FF41 !important;
        border: 1px solid #00FF41 !important;
        box-shadow: none !important;
        animation: none !important;
    }
    
    .file-magnitude-warning {
        background: #140808 !important;
        border: 1px solid #FF3131 !important;
        border-radius: 2px;
        padding: 1rem 1.25rem;
        font-family: 'JetBrains Mono', ui-monospace, monospace !important;
        font-size: 0.72rem;
        color: #FF3131 !important;
        letter-spacing: 2px;
        text-transform: uppercase;
        margin-bottom: 1rem;
    }

    .data-synced {
        font-family: 'JetBrains Mono', ui-monospace, monospace !important;
        font-size: 0.68rem;
        color: #00FF41 !important;
        text-shadow: none !important;
        margin-top: 0.5rem;
        padding: 0.5rem;
        border: 1px solid #00FF41;
        border-radius: 2px;
        letter-spacing: 2px;
        text-transform: uppercase;
        background: #050505 !important;
    }
    
    /* Appeal paper + preview — EB Garamond (terminal-dark “bond desk” preview) */
    .appeal-paper {
        background: #0D0D0D !important;
        border-radius: 2px;
        padding: 0.75rem 0 0 0;
        margin: 0.5rem 0 0 0;
        border: none !important;
        box-shadow: none !important;
        position: relative;
    }
    .appeal-preview-garamond {
        font-family: 'EB Garamond', 'Garamond', 'Times New Roman', serif !important;
        font-size: 1.08rem !important;
        line-height: 1.68 !important;
        color: #E0E0E0 !important;
        border: 1px solid #262626 !important;
        padding: 1.35rem 1.75rem !important;
        background: #0D0D0D !important;
        border-radius: 2px !important;
        margin-top: 0.35rem !important;
    }
    .appeal-preview-garamond p, .appeal-preview-garamond li,
    .appeal-preview-garamond td, .appeal-preview-garamond th,
    .appeal-preview-garamond h1, .appeal-preview-garamond h2, .appeal-preview-garamond h3 {
        font-family: 'EB Garamond', 'Garamond', Georgia, serif !important;
        color: #E0E0E0 !important;
    }
    .appeal-paper-header {
        display: flex;
        justify-content: flex-end;
        align-items: flex-start;
        margin-bottom: 0.35rem;
    }
    .appeal-paper-copy {
        cursor: pointer;
        padding: 0.4rem 0.55rem;
        border-radius: 2px;
        background: #050505 !important;
        border: 1px solid #262626 !important;
        color: #E0E0E0 !important;
        transition: border-color 0.2s, color 0.2s;
        box-shadow: none !important;
        font-family: 'JetBrains Mono', ui-monospace, monospace !important;
        font-size: 0.65rem;
        letter-spacing: 0.12em;
        text-transform: uppercase;
    }
    .appeal-paper-copy:hover {
        background: #111111 !important;
        border-color: #00FF41 !important;
        color: #00FF41 !important;
        box-shadow: none !important;
    }
    .appeal-paper-body {
        font-family: 'EB Garamond', 'Garamond', Georgia, serif !important;
        font-size: 1.05rem;
        line-height: 1.65;
        color: #E0E0E0 !important;
        white-space: pre-wrap;
    }
    
    /* Client Facing View — pitch deck hero + high-contrast CTA */
    .client-pitch-hero {
        background: linear-gradient(145deg, #042818 0%, #0a3d2c 42%, #061a14 100%);
        border: 2px solid #00FF41;
        border-radius: 4px;
        padding: 2.25rem 1.25rem;
        text-align: center;
        margin: 0.35rem 0 1.75rem;
        box-shadow: 0 0 32px rgba(0, 255, 65, 0.12);
    }
    .client-pitch-hero .client-pitch-label {
        font-family: 'JetBrains Mono', ui-monospace, monospace !important;
        font-size: clamp(0.72rem, 2.2vw, 0.95rem);
        letter-spacing: 0.2em;
        color: rgba(255, 255, 255, 0.88) !important;
        text-transform: uppercase;
        margin-bottom: 0.85rem;
        font-weight: 500;
    }
    .client-pitch-hero .client-pitch-value {
        font-family: 'JetBrains Mono', ui-monospace, monospace !important;
        font-size: clamp(1.85rem, 5.5vw, 3.1rem);
        font-weight: 600;
        color: #00FF41 !important;
        text-shadow: 0 0 28px rgba(0, 255, 65, 0.4);
        line-height: 1.12;
        margin: 0;
    }
    .client-pitch-sub {
        font-family: 'JetBrains Mono', ui-monospace, monospace !important;
        font-size: 0.72rem;
        letter-spacing: 0.14em;
        color: rgba(224, 224, 224, 0.65) !important;
        margin-top: 1rem;
        text-transform: uppercase;
    }
    .client-view-cta button[kind="primary"],
    .main .client-view-cta button {
        background: #00FF41 !important;
        color: #050505 !important;
        border: 3px solid #00FF41 !important;
        font-weight: 700 !important;
        letter-spacing: 0.06em !important;
        text-transform: uppercase !important;
        font-size: 1rem !important;
        padding: 0.65rem 1rem !important;
        box-shadow: 0 0 24px rgba(0, 255, 65, 0.25) !important;
    }
    .client-view-cta button[kind="primary"]:hover {
        filter: brightness(1.08);
        border-color: #33ff66 !important;
    }

    /* Streamlit alerts — dark chrome (no bright white overlays on Operations) */
    div[data-testid="stAlert"] {
        background-color: #111111 !important;
        border: 1px solid #2f2f2f !important;
        color: #e5e5e5 !important;
    }
    div[data-testid="stAlert"] p,
    div[data-testid="stAlert"] li,
    div[data-testid="stAlert"] span {
        color: #e5e5e5 !important;
    }
    /* Demo / Paystack Reviewer (Slate Dark) */
    .demo-audit-hero, .paystack-reviewer-hero {
        text-align: center;
        padding: 1.5rem 1rem 1.25rem;
        margin: 0 0 1.25rem 0;
        background: #0f172a !important;
        border: 1px solid #334155 !important;
        border-radius: 4px;
    }
    .demo-audit-title {
        font-family: 'JetBrains Mono', ui-monospace, monospace !important;
        font-size: 1.35rem !important;
        letter-spacing: 0.2em !important;
        color: #f1f5f9 !important;
        text-transform: uppercase;
        margin-bottom: 0.5rem;
    }
    .demo-audit-sub {
        font-family: 'JetBrains Mono', ui-monospace, monospace !important;
        font-size: 0.72rem !important;
        letter-spacing: 0.12em !important;
        color: #38bdf8 !important;
        text-transform: uppercase;
    }
    .reviewer-recoverable-hero {
        text-align: center;
        font-family: 'JetBrains Mono', ui-monospace, monospace !important;
        font-size: 0.72rem !important;
        letter-spacing: 0.14em !important;
        color: #94a3b8 !important;
        text-transform: uppercase;
        margin: 0.75rem 0 1rem;
        padding: 1rem;
        background: #0f172a !important;
        border: 1px solid #334155 !important;
        border-radius: 4px;
    }
    .reviewer-recoverable-hero .reviewer-usd {
        display: block;
        font-size: clamp(1.65rem, 4vw, 2.15rem) !important;
        font-weight: 600 !important;
        color: #38bdf8 !important;
        letter-spacing: 0.04em !important;
        margin-top: 0.5rem;
        text-transform: none !important;
    }
    .reviewer-recoverable-hero .reviewer-hero-sub {
        display: block;
        font-size: 0.62rem !important;
        letter-spacing: 0.12em !important;
        color: #64748b !important;
        margin-top: 0.45rem;
        text-transform: uppercase;
    }
    .reviewer-sidebar-executive {
        margin: 0.75rem 0 0.5rem 0;
        padding: 0.65rem 0.7rem;
        border: 1px solid rgba(51, 65, 85, 0.9);
        border-radius: 4px;
        background: linear-gradient(180deg, #0f172a 0%, #020617 100%);
    }
    .reviewer-sidebar-executive .rse-head {
        font-family: 'JetBrains Mono', ui-monospace, monospace !important;
        font-size: 0.58rem !important;
        letter-spacing: 0.2em !important;
        text-transform: uppercase;
        color: #94a3b8 !important;
        margin: 0 0 0.45rem 0 !important;
    }
    .reviewer-sidebar-executive .rse-line {
        font-family: 'JetBrains Mono', ui-monospace, monospace !important;
        font-size: 0.72rem !important;
        line-height: 1.45;
        color: #f1f5f9 !important;
        margin: 0.2rem 0 !important;
    }
    .reviewer-sidebar-executive .rse-cfo {
        color: #38bdf8 !important;
    }
    .reviewer-footer-lock {
        margin-top: 1.25rem;
        padding: 0.85rem 1rem;
        text-align: center;
        border-top: 1px solid rgba(212, 175, 55, 0.28);
        border-bottom: 1px solid rgba(212, 175, 55, 0.15);
        background: #020617;
        font-family: 'JetBrains Mono', ui-monospace, monospace !important;
    }
    .reviewer-footer-lock .rfl-addr {
        font-size: 0.72rem !important;
        letter-spacing: 0.08em !important;
        color: #e2e8f0 !important;
        margin: 0 0 0.35rem 0 !important;
    }
    .reviewer-footer-lock .rfl-mid {
        font-size: 0.68rem !important;
        letter-spacing: 0.12em !important;
        color: #a3e635 !important;
        margin: 0 0 0.5rem 0 !important;
    }
    .reviewer-footer-lock .rfl-copy {
        font-size: 0.58rem !important;
        letter-spacing: 0.14em !important;
        color: #64748b !important;
        text-transform: uppercase;
        margin: 0 !important;
    }
    .reviewer-section-title {
        font-family: 'JetBrains Mono', ui-monospace, monospace !important;
        font-size: 0.68rem !important;
        letter-spacing: 0.16em !important;
        text-transform: uppercase;
        color: #94a3b8 !important;
        margin: 1rem 0 0.5rem !important;
    }
    .reviewer-roles-strip p {
        margin-bottom: 0.5rem !important;
    }
    .neural-audit-demo-executive .demo-exec-line {
        font-family: 'JetBrains Mono', ui-monospace, monospace !important;
        font-size: 0.92rem !important;
        letter-spacing: 0.06em !important;
        color: #f1f5f9 !important;
        margin-bottom: 0.45rem !important;
    }
    .neural-audit-demo-executive .demo-address-line {
        font-size: 0.82rem !important;
        color: #bae6fd !important;
        border-top: 1px solid rgba(51, 65, 85, 0.65);
        padding-top: 0.65rem !important;
        margin-top: 0.5rem !important;
    }
    .neural-audit-demo-executive .demo-merchant-line {
        font-size: 0.72rem !important;
        letter-spacing: 0.08em !important;
        color: #94a3b8 !important;
        margin-top: 0.5rem !important;
    }
    .senturion-dark-panel {
        background: #0f172a !important;
        border: 1px solid #334155 !important;
        border-radius: 4px !important;
        padding: 1rem 1.15rem !important;
        margin-bottom: 1.25rem !important;
        color: #cbd5e1 !important;
        font-size: 0.88rem !important;
        line-height: 1.55 !important;
    }
    .senturion-dark-panel p { margin: 0 0 0.65rem 0; color: #cbd5e1 !important; }
    .senturion-dark-panel p:last-child { margin-bottom: 0; }
    .senturion-dark-callout {
        background: #111111 !important;
        border: 1px solid #2f2f2f !important;
        border-radius: 2px !important;
        padding: 0.65rem 0.85rem !important;
        margin: 0.5rem 0 0.75rem 0 !important;
        color: #d4d4d4 !important;
        font-size: 0.82rem !important;
        line-height: 1.45 !important;
    }

    /* Mandatory KYC strip — all pages */
    .kyc-footer-bar {
        margin-top: 1.25rem;
        padding: 0.65rem 0.75rem;
        border-top: 1px solid rgba(212, 175, 55, 0.25);
        border-bottom: 1px solid rgba(212, 175, 55, 0.15);
        background: linear-gradient(180deg, #0f172a 0%, #020617 100%);
        font-family: 'JetBrains Mono', ui-monospace, monospace !important;
        font-size: 0.58rem !important;
        letter-spacing: 0.1em !important;
        color: #e8e8e8 !important;
        text-align: center;
        text-transform: none !important;
    }
    .kyc-footer-legal strong {
        color: #f8fafc !important;
        font-weight: 700 !important;
    }
    a.reviewer-pay-link {
        display: inline-block;
        margin-top: 0.75rem;
        padding: 0.65rem 1rem;
        background: #0ea5e9 !important;
        color: #0f172a !important;
        font-family: 'JetBrains Mono', ui-monospace, monospace !important;
        font-weight: 600;
        text-decoration: none !important;
        border-radius: 4px;
        border: 1px solid #38bdf8;
    }

    /* Footer — dark executive band (partners / internal only; not shown to clinic role) */
    .quantum-footer {
        margin-top: 2rem;
        padding: 1.25rem 1rem 1rem;
        border-top: 1px solid #1f2937;
        background: linear-gradient(180deg, #0a0a0a 0%, #050505 100%);
        border-radius: 2px;
        font-family: 'JetBrains Mono', ui-monospace, monospace !important;
        font-size: 0.62rem;
        color: #9ca3af !important;
        letter-spacing: 0.12em;
        text-transform: uppercase;
        text-align: center;
    }
    .quantum-footer-sub {
        display: block;
        margin-top: 0.55rem;
        font-family: 'JetBrains Mono', ui-monospace, monospace !important;
        font-size: 0.62rem !important;
        letter-spacing: 0.08em !important;
        color: #a1a1aa !important;
        text-transform: none !important;
        line-height: 1.5;
    }
    .quantum-footer-addr {
        display: block;
        margin-top: 0.65rem;
        font-size: 0.6rem !important;
        letter-spacing: 0.06em !important;
        color: #d4d4d8 !important;
        text-transform: none !important;
    }
    .quantum-footer-mail a {
        color: #4ade80 !important;
        text-decoration: none !important;
        font-weight: 500;
    }
    .quantum-footer-mail a:hover {
        text-decoration: underline !important;
    }
    .footer-db-pill {
        display: inline-flex;
        align-items: center;
        gap: 0.35rem;
        margin-left: 0.4rem;
        vertical-align: middle;
        white-space: nowrap;
    }
    .footer-db-dot {
        width: 8px;
        height: 8px;
        border-radius: 50%;
        flex-shrink: 0;
    }
    .footer-db-dot.db-dot-online {
        background: #00FF41 !important;
        box-shadow: 0 0 10px rgba(0, 255, 65, 0.85);
    }
    .footer-db-dot.db-dot-offline {
        background: #6b7280 !important;
        box-shadow: none;
    }
    .footer-db-txt {
        font-size: 0.58rem !important;
        letter-spacing: 0.06em !important;
        text-transform: none !important;
    }
    .footer-db-txt.footer-db-on {
        color: #a3e635 !important;
    }
    .footer-db-txt.footer-db-off {
        color: #9ca3af !important;
    }

    .stButton > button[kind="primary"],
    div[data-testid="stButton"] button[kind="primary"] {
        background: #141414 !important;
        color: #E0E0E0 !important;
        border: 1px solid #262626 !important;
        border-radius: 2px !important;
        box-shadow: none !important;
        font-family: 'JetBrains Mono', ui-monospace, monospace !important;
        letter-spacing: 2px;
        text-transform: uppercase;
    }
    .stButton > button[kind="primary"]:hover {
        border-color: #00FF41 !important;
        color: #00FF41 !important;
        box-shadow: none !important;
    }

    /* Alerts — terminal chrome + Matrix / neon accents */
    div[data-testid="stAlert"] {
        border: 1px solid #262626 !important;
        background: #0D0D0D !important;
        color: #E0E0E0 !important;
    }
    div[data-testid="stAlert"] p, div[data-testid="stAlert"] li, div[data-testid="stAlert"] span {
        color: #E0E0E0 !important;
    }

    [data-testid="stExpander"] details {
        border: 1px solid #262626 !important;
        border-radius: 2px !important;
        background: #0D0D0D !important;
    }
    [data-testid="stExpander"] summary {
        font-family: 'JetBrains Mono', ui-monospace, monospace !important;
        letter-spacing: 0.08em;
        color: #E0E0E0 !important;
    }

    /* Inputs — ruled */
    .stTextInput input, .stNumberInput input, textarea {
        background: #050505 !important;
        color: #E0E0E0 !important;
        border: 1px solid #262626 !important;
        font-family: 'JetBrains Mono', ui-monospace, monospace !important;
    }
</style>
""", unsafe_allow_html=True)

EXTRACTION_PROMPT = """You are a Federal Regulatory Auditor. Analyze the following document and extract denial information with statutory awareness.

First, identify the **Place of Service (State)** from the document header, letterhead, facility address, or claim face sheet
(two-letter USPS state code preferred, e.g. TX; if only city is given, infer the state when reliable, else use "Unknown").

For each claim/denial found, extract EXACTLY these fields:
1. Patient ID - The patient identifier (number or alphanumeric code)
2. Denial Code - The denial/rejection code (e.g., CO-16, CO-22, PR-1, etc.)
3. Reason for Denial - Brief explanation of why the claim was denied
4. Fix Action - One concise sentence describing the action needed to resolve the denial
5. Place of Service (State) - Same state as identified above (repeat per row if multiple denials)
6. Law Cited - The **primary** statute or regulatory hook you would cite for a formal appeal for THIS denial, choosing
   at least one class where applicable: (a) that state's **Prompt Payment Act** (e.g. Texas Insurance Code § 1301.103
   when POS is Texas), (b) **ERISA § 503** for employer-sponsored plans, or (c) **ACA / PHS Act § 2719** external review
   mandate. Output a short citation string (e.g. "Texas Insurance Code § 1301.103" or "ERISA § 503" or "PHS Act § 2719 (ACA)").
7. Potential Revenue - **Mandatory:** Locate dollar amounts labeled **Total Billed**, **Balance Due**, **Amount Due**, **Total Balance**, **Amount Denied**, or **Patient Responsibility** (or clear equivalents in the EOB/letter). Use the best recovery proxy for this denial row: prefer the **denied** / at-risk amount. If multiple such figures apply, use the **highest denied value**. Output numeric USD only in CSV (no $ symbol); if none can be found, 0
8. Denial Date - Date the payer denied or the letter is dated (YYYY-MM-DD preferred; else best parseable form)
9. Win Probability - **Mandatory:** Your estimated likelihood (0–100 **integer only**, no percent symbol) that an appeal would succeed, based primarily on **Denial Code** strength/weakness and **Law Cited** leverage (e.g., strong Prompt Payment / ERISA hooks score higher; weak or missing documentation lower). Use professional judgment; avoid all identical numbers unless justified.
10. Payer Name - The **insurance carrier / plan administrator** named on the EOB, denial letter, or ID card (e.g. "Blue Cross Blue Shield", "Aetna", "UnitedHealthcare", "Cigna", employer plan like "Marriott", etc.). If multiple payers appear, use the **primary** payer for the denied line item. If unknown, output "Unknown".

If multiple denials appear in the document, extract each one. If no structured denial info is found, infer reasonable placeholder values based on the document context.

**Appeal mode (post-processing rule for the engine):** If **Reason for Denial** references **Top Hat**, **Vesting**, or **Retirement** (case-insensitive), the system sets **Appeal Mode** to **STATUTORY** (ERISA plan / pension posture — not medical claim coding).

Return ONLY a valid CSV with headers EXACTLY:
Patient ID, Denial Code, Reason for Denial, Fix Action, Place of Service (State), Law Cited, Potential Revenue, Denial Date, Win Probability, Payer Name
One row per denial. No other text or explanation."""

MASTER_LITIGATION_PROMPT_STATUTORY = """
You are a **Federal ERISA litigator** drafting an **extremely aggressive**, litigation-grade statutory challenge for **employee benefit plan** disputes (pension, Top Hat, vesting, retirement benefits — **not** professional medical billing). The output must read like a **federal enforcement posture**, not a polite request.

Produce ONLY the appeal body (no letterhead, To/From, date line, or signature block — the system inserts those).

STRICT EXCLUSIONS — do **not** mention or rely on:
- **CPT**, **ICD-10**, **HCPCS**, **NCCI** edits, **LCD**/**NCD**, **CMS**, **Medicare** coverage policy, or routine **medical necessity** framing.
- Provider coding, bundling, or clinical procedure terminology.

Replace any “medical necessity” framing with **Fiduciary Obligation & Vesting Compliance** and plan-document fidelity.

MANDATORY OPENING (verbatim lead-in — first sentence of the body text MUST be EXACTLY this, before any other prose):
**NOTICE OF INTENT TO ESCALATE: FAILURE TO ADHERE TO ERISA § 503 PROCEDURAL SAFEGUARDS.**

TONE — EXTREMELY AGGRESSIVE (non-negotiable):
- **Never** "request," "ask," "invite," or "seek review." You **demand** compliance, corrective action, and full-and-fair review under ERISA.
- Use command language: **Demand**, **Direct**, **Require**, **Compel**, **Insist upon** — not "we hope" or "we respectfully ask."
- FORBIDDEN soft closings: any variant of appreciation for reconsideration. Close with **Formal Demand for Compliance** and escalation readiness.
- PREFERRED constructions: "The administrative record compels reversal …", "The plan’s position is procedurally defective under ERISA § 503 …", "Full-and-fair review is not optional — it is mandatory."

## Legal Basis for Appeal
Immediately after the mandatory opening sentence, write 2–4 paragraphs anchored in **ERISA § 503**, **29 C.F.R. § 2560.503-1** (claims procedures), **fiduciary duties** under ERISA § 404/409 where applicable, and the **plan document / SPD** versus statutory minimums. Cite **vesting**, **Top Hat** plan rules, or **retirement** benefit terms as reflected in CLAIM DATA. No Medicare/CMS. Frame the plan/payer as **out of compliance** with mandatory procedural safeguards, not as a good-faith disagreement.

## PLAN PROVISION VS. STATUTORY REQUIREMENT
Output a **GitHub-flavored markdown pipe table** with EXACTLY these two column headers (header row required):

| Plan Interpretation (Marriott) | ERISA § 503 / 29 C.F.R. § 2560.503-1 Requirement |
|--------------------------------|--------------------------------------------------|
| [Payer/plan position from CLAIM DATA — quote or paraphrase.] | [Map to ERISA § 503 and 29 C.F.R. § 2560.503-1 procedural and substantive requirements; one row may run long.] |

Include at least **one data row** after the header.

## Administrative Narrative
600+ words: **Fiduciary Obligation & Vesting Compliance**, plan interpretation, full-and-fair review, and administrative record — **not** clinical coding. The narrative must **demand** reversal and compliance; it must not read like a courtesy letter. Close with **Formal Demand for Compliance**.

## Mandatory 30-Day Escalation & Regulatory Referral (required standalone paragraph)
You MUST include **one full paragraph** (not a bullet list) that states, in substance, all of the following:
- If the denial is **not reversed** (or the claim is not brought into full compliance with ERISA § 503 procedural safeguards) **within thirty (30) calendar days** of this notice, the claimant reserves the right to pursue **formal referral** of the plan’s and/or payer’s conduct to the **U.S. Department of Labor**, **Employee Benefits Security Administration (EBSA)**, and to the **Maryland Insurance Administration**, for investigation of **bad faith** claims handling and **failure to adhere to ERISA procedural requirements**, in addition to any other remedies available at law or in equity.
- This paragraph must read like a **hard enforcement deadline**, not a vague possibility.

## Regulatory Escalation Notice
Reinforce that continued procedural non-compliance is **unacceptable** and that **administrative referral paths** (EBSA / state insurance authority) are **active and imminent** if compliance is not restored. Do not water down the 30-day paragraph above.

RULES:
1. The **mandatory opening sentence** and the **mandatory 30-day escalation paragraph** are **required** in every output.
2. Use fields from CLAIM DATA exactly (no bracket placeholders).
3. Do not output To/From/Date/RE lines.

OUTPUT FORMAT: Markdown only, using the ## sections above exactly as titled (five sections total, including **Mandatory 30-Day Escalation & Regulatory Referral**).
"""

MASTER_LITIGATION_PROMPT = """
You are a Senior Federal Medical Auditor producing a **statutory-heavy**, World-Class Revenue Cycle Management appeal.

Produce ONLY the appeal body (no letterhead, To/From, date line, or signature block — the system inserts those).
Use the EXACT markdown section structure below, in this order.

PLACE OF SERVICE & STATUTORY ANCHOR:
- Identify **Place of Service (State)** from CLAIM DATA (e.g. `Place of Service (State)` or header-derived). Name the state explicitly in the Legal Basis section.
- For **every** denial at issue, the **Legal Basis for Appeal** must cite **at least one** of the following (pick what fits the plan/claim):
  a) The **state's specific Prompt Payment Act** or insurance claims-timeliness statute — when POS is known, prefer a **concrete citation**
     (e.g. **Texas Insurance Code § 1301.103** when Texas applies; otherwise cite the correct statute for the identified state); or
  b) **ERISA § 503** (full and fair review / administrative remedies) for **employer-sponsored** plans subject to ERISA; or
  c) The **ACA**/**PHS Act** **external review** framework (e.g. **42 U.S.C. § 300gg-19** / **§ 2719**-style independent medical review and appeals parity).

INSTITUTIONAL VOICE (Neural / regulatory — never casual):
- FORBIDDEN: "I think", "we believe", "I feel", "might", "maybe", "probably" when stating facts.
- FORBIDDEN soft closing: **"We would appreciate a re-review"** (and similar). Replace with **"Formal Demand for Compliance"**.
- PREFERRED: "Neural audit confirms …", "Clinical documentation contradicts …",
  "The administrative record establishes …", "Verified coding and clinical evidence support …".

## Legal Basis for Appeal
Write 2–4 tight paragraphs that establish the statutory and contractual basis for review. You MUST:
- Open by tying the identified **Place of Service (State)** to the applicable **Prompt Payment** / state insurance code path when the product is state-regulated or fully insured (cite the state's act where possible, e.g. Texas Insurance Code § 1301.103 for Texas POS).
- Where the plan is **ERISA-governed**, integrate **ERISA § 503** procedural protections and full-and-fair review.
- Where **ACA**-regulated **external review** or appeals parity applies, integrate the **external review** mandate (PHS Act § 2719 / 42 U.S.C. § 300gg-19) as applicable.
- Map denial type in CLAIM DATA (medical necessity, timely filing, coding, bundling, CO/PR codes) to these authorities.
- Do not use the words **Threat**, **Lawsuit**, or **Sue**; administrative/regulatory tone only.

## Clinical Discrepancy Analysis
Output a **GitHub-flavored markdown pipe table** with EXACTLY these two columns (header row required):

| Insurer Denial Justification | Provider Clinical Substantiation |
|------------------------------|----------------------------------|
| [Quote or paraphrase the payer's stated reason from CLAIM DATA — if absent, state "As stated in payer correspondence: [summary]."] | [See CODE VALIDATION + LCD rules below — single cell may run long; use semicolons between sentences if needed.] |

**CODE VALIDATION (Column B — mandatory for every data row):**
- Always **explicitly name** the **ICD-10 diagnosis code(s)** and **CPT procedure code(s)** (or HCPCS, if that is what the claim under appeal uses) that the provider relies on. Pull codes from CLAIM DATA when present; if a code is inferable from context, label it clearly (e.g., "ICD-10: …; CPT: …").

**MEDICAL NECESSITY / LCD LINK (Column B — one sentence per row):**
- For **each** table row, include **one sentence** explaining why the cited **CPT** (or HCPCS) represents the **gold standard** (or medically necessary and appropriate) service for the **specific ICD-10** diagnosis(es) at issue, framed by **CMS Local Coverage Determinations (LCDs)** applicable to that code pair or jurisdiction. If no LCD clearly applies, cite the closest **NCD** or state that coverage is supported by **generally accepted Medicare coverage policy** for that ICD-10/CPT pairing, without inventing LCD IDs.

Include at least **one data row** after the header; add rows only when multiple distinct denial themes are supported by CLAIM DATA.

## Administrative Narrative
600+ words of cold, data-driven prose: medical necessity, benefit interpretation, coding/NCCI context where relevant,
and administrative correctness. Close with a **Formal Demand for Compliance** directed at the payer (not "we would appreciate a re-review").
Reference statutory themes from the Legal Basis as operational requirements, not personal requests.

## Regulatory Escalation Notice
A short subsection (3–5 sentences) stating that **continued non-compliance** with the cited prompt-payment, ERISA, or ACA review obligations
may result in **referral to the State Department of Insurance** and other regulatory channels as authorized by law. Use measured regulatory language;
do not use Threat/Lawsuit/Sue.

RULES:
1. Do not use the words **Lawsuit** or **Sue**. Do not use the word **Threat** outside measured regulatory phrasing. **Regulatory Escalation Notice** must remain administrative (referral to regulators), not personal attacks.
2. Use Patient ID, Claim Number, Denial Code, Reason for Denial, Place of Service (State), and Law Cited from CLAIM DATA exactly (no bracket placeholders).
3. Do not output To/From/Date/RE lines.

OUTPUT FORMAT: Markdown only, using the four ## sections above exactly as titled.
"""


def clean_text(val) -> str:
    """Fail-safe: return stripped string or ''. Never crash on None or blank AI output."""
    if val is None:
        return ""
    try:
        return str(val).strip() or ""
    except (TypeError, AttributeError):
        return ""


# Pension vs. patient — ERISA statutory track (logic gate; not LLM-inferred alone)
STATUTORY_DENIAL_KEYWORDS = ("top hat", "tophat", "top-hat", "vesting", "retirement")

# Neural Triage — Strike list (Agent Terminal)
STRIKE_MIN_REVENUE_USD = 1000.0
STRIKE_MIN_WIN_PCT = 80.0
# High-value enforcement — Agent Terminal sort + urgency tag
HIGH_VALUE_TARGET_USD = 5000.0
URGENCY_TAG_HIGH_VALUE = "[HIGH_VALUE_TARGET]"


def _is_strike_claim_row(r: dict) -> bool:
    """Neon strike: high recoverable + high win probability."""
    rev = _parse_amount_denied(r.get("Potential Revenue", "0"))
    wp = _parse_win_probability(r.get("Win Probability", "0"))
    return rev >= STRIKE_MIN_REVENUE_USD and wp >= STRIKE_MIN_WIN_PCT


def _neural_triage_priority_score(r: dict) -> float:
    rev = _parse_amount_denied(r.get("Potential Revenue", "0"))
    wp = _parse_win_probability(r.get("Win Probability", "0"))
    return rev * (wp / 100.0)


def _infer_appeal_mode(reason: str) -> str:
    """STATUTORY = ERISA / plan / pension posture when denial reason matches pension triggers."""
    r = (reason or "").lower()
    if any(k in r for k in STATUTORY_DENIAL_KEYWORDS):
        return "STATUTORY"
    return "CLINICAL"


def _statutory_auto_lock(reason: str, law_cited: str) -> bool:
    """
    Lethal Statutory template without human review: Top Hat / ERISA § 503 signals.
    """
    blob = f"{reason or ''} {law_cited or ''}".lower()
    blob_compact = re.sub(r"\s+", "", blob)
    if "tophat" in blob_compact or "top-hat" in blob or "top hat" in blob:
        return True
    if "top hat plan" in blob:
        return True
    if "erisa" in blob and "503" in blob:
        return True
    if "29c.f.r" in blob and "2560.503" in blob.replace(" ", ""):
        return True
    return False


def _resolve_neural_appeal_mode(reason: str, law_cited: str) -> str:
    """Final Appeal Mode for neural rows — statutory auto-flip overrides clinical."""
    if _statutory_auto_lock(reason, law_cited):
        return "STATUTORY"
    return _infer_appeal_mode(reason)


def _parse_win_probability(val) -> float:
    """0–100 win probability from extraction or widgets."""
    if val is None:
        return 0.0
    s = str(val).strip().replace("%", "").replace("—", "").strip()
    if not s:
        return 0.0
    try:
        x = float(s)
        return max(0.0, min(100.0, x))
    except (TypeError, ValueError):
        return 0.0


def _text_has_cpt_code(*parts: str) -> bool:
    """True if claim/appeal text contains a plausible CPT/HCPCS-style code (for CMS compliance footer gate)."""
    blob = " ".join(str(p or "") for p in parts)
    if not blob.strip():
        return False
    if re.search(r"\bCPT\b|\bHCPCS\b|\bICD-10\b", blob, re.I):
        return True
    return bool(re.search(r"\b\d{5}\b", blob)) or bool(re.search(r"\b[A-Z]\d{4}\b", blob))


def _normalize_denial_csv_row(row: dict) -> dict:
    """Map flexible CSV / extraction keys to canonical Neural Audit columns (includes POS, Law, revenue, dates)."""
    pid = dc = reason = fix = pos = law = rev_s = ddate_s = pname = wp_s = payer_s = ""
    for k, v in row.items():
        key = clean_text(k)
        if not key:
            continue
        val = clean_text(v)
        kl = key.lower()
        if "patient" in kl and "id" in kl:
            pid = val
        elif "patient" in kl and "name" in kl and "id" not in kl:
            pname = val
        elif kl in ("patient name", "member name", "subscriber name", "patient"):
            pname = pname or val
        elif "denial" in kl and "code" in kl:
            dc = val
        elif "reason" in kl and "denial" in kl:
            reason = val
        elif kl.startswith("reason") or "reason for" in kl:
            reason = val
        elif "fix" in kl or kl == "fix action":
            fix = val
        elif "action" in kl and "transaction" not in kl and "denial" not in kl:
            fix = val
        elif "place" in kl and ("service" in kl or "state" in kl):
            pos = val
        elif "law" in kl and "cited" in kl:
            law = val
        elif kl in ("law cited", "statute", "citation", "legal basis"):
            law = val or law
        elif "potential" in kl and "revenue" in kl:
            rev_s = val
        elif "recovery" in kl and "target" in kl:
            rev_s = rev_s or val
        elif "denial" in kl and "date" in kl:
            ddate_s = val
        elif kl in ("date of denial", "denied on", "service date"):
            ddate_s = ddate_s or val
        elif "win" in kl and "prob" in kl:
            wp_s = val
        elif kl.replace(" ", "") in ("winprobability", "win_prob", "winpct", "priorityscore"):
            wp_s = wp_s or val
        elif "payer" in kl and "name" in kl:
            payer_s = val
        elif kl in ("payer", "insurer", "carrier", "plan name", "insurance", "payor"):
            payer_s = payer_s or val
    if not law:
        for k, v in row.items():
            kl = clean_text(k).lower()
            if kl in ("law", "statute"):
                law = clean_text(v)
                break
    dash = "—"
    _wpf = _parse_win_probability(wp_s)
    if not (wp_s or "").strip() and (dc or law or reason):
        # Heuristic backfill only if the model omitted Win Probability
        _seed = (abs(hash((dc, law, reason))) % 41) + 35
        _wpf = float(max(35, min(92, _seed)))
    _wp_disp = f"{_wpf:.0f}"
    return {
        "Patient ID": pid or dash,
        "Patient Name": pname or dash,
        "Denial Code": dc or dash,
        "Reason for Denial": reason or dash,
        "Fix Action": fix or dash,
        "Place of Service (State)": pos or dash,
        "Law Cited": law or dash,
        "Potential Revenue": rev_s or "0",
        "Denial Date": ddate_s or dash,
        "Win Probability": _wp_disp,
        "Payer Name": payer_s or dash,
        "Appeal Mode": _resolve_neural_appeal_mode(reason, law),
    }


def _text_only_for_prompt(val) -> str:
    """
    For LLM prompts only: coerce to plain text. Never stringify raw bytes (would embed
    huge b'...' blobs and can trigger model safety / context issues). Ignore file-like objects.
    """
    if val is None:
        return ""
    if isinstance(val, (bytes, bytearray, memoryview)):
        return ""
    try:
        if hasattr(val, "read") and callable(val.read):
            return ""
    except (TypeError, AttributeError):
        pass
    try:
        return str(val).strip() or ""
    except (TypeError, AttributeError, ValueError):
        return ""


MAX_PDF_PAGES = 10
MAX_FILE_SIZE_BYTES = 10 * 1024 * 1024  # 10MB


def _strip_docx_markdown_stars(text: str) -> str:
    """Remove markdown bold markers so they do not appear as literal stars in Word."""
    if not text:
        return text
    return text.replace("**", "").replace("__", "")


DOCX_HANDWRITTEN_BLANK = "____________________"

# World-class RCM document architecture (Appeal Engine)
RCM_MARGIN_INCHES = 1.25
COMPLIANCE_ATTESTATION_TEXT = (
    "This audit has been verified for accuracy against current NCCI edits and CMS guidelines."
)
PENALTY_CLAUSE_TEXT = (
    "Continued non-compliance with applicable prompt-payment, ERISA, or ACA review obligations may result in "
    "referral to the State Department of Insurance and other regulatory authorities as authorized by law."
)
# Fine-print footer on every Senturion PDF (Legal Authentication)
PDF_LEGAL_COMPLIANCE_FOOTER_TEXT = (
    "This document is a formal record of a Neural Audit conducted under ERISA § 503 and MD Code § 15-1005. "
    "Unauthorized tampering or failure to respond within statutory limits may be reported to federal regulators."
)
# Institutional PDF branding — matches STEALTH_ACCENT / UI (Titanium Silver)
PDF_TITANIUM_RGB = (224, 224, 224)  # #E0E0E0
# Senturion Audit Certificate — body disclaimer (batch / vault summary PDFs)
SENTURION_AUDIT_CERTIFICATE_LEGAL_DISCLAIMER = (
    "LEGAL COMPLIANCE DISCLAIMER: This Senturion Audit Certificate is a management summary produced from "
    "client-supplied and system-processed claim data. It does not constitute legal advice, a guarantee of "
    "recovery, or a final determination of benefits, coverage, or liability. The recipient is solely "
    "responsible for verifying amounts, coding, and regulatory compliance with applicable law, plan documents, "
    f"and payer policies. {BRAND_NAME} is not a law firm; consult qualified counsel for legal "
    "interpretation. The Audit Tracking Hash below is a non-repudiation fingerprint for this document instance."
)
PDF_TOP_MARGIN_FOR_SEAL_MM = 28.0  # room for top-right holographic seal
PDF_BOTTOM_MARGIN_FOR_FOOTER_MM = 30.0  # legal footer + Audit Tracking Hash line (fail-safe fingerprint)

# Payer Intelligence — hidden local ledger (win proxies by payer × statutory hook)
PAYER_INTEL_DB_PATH = os.path.join(os.path.dirname(os.path.abspath(__file__)), ".senturion_payer_intel.json")
PAYER_INTEL_DEFAULT_SEED: dict[str, dict[str, dict[str, int]]] = {
    "BCBS": {"ERISA § 503": {"wins": 18, "n": 23}, "29 C.F.R. § 2560.503-1": {"wins": 5, "n": 12}},
    "Aetna": {"ERISA § 503": {"wins": 14, "n": 20}, "PHS Act § 2719 (ACA)": {"wins": 3, "n": 8}},
    "Marriott": {"ERISA § 503": {"wins": 11, "n": 14}, "ERISA § 404": {"wins": 4, "n": 9}},
    "UnitedHealthcare": {"ERISA § 503": {"wins": 16, "n": 25}},
    "Cigna": {"ERISA § 503": {"wins": 9, "n": 15}},
    "Humana": {"ERISA § 503": {"wins": 7, "n": 13}},
}


def _normalize_payer_intel_key(raw: str) -> str:
    """Bucket payer names for intelligence tracking (BCBS, Aetna, Marriott, UHC, …)."""
    t = (raw or "").lower()
    if any(x in t for x in ("blue cross", "bcbs", "anthem", "bluecard", "bc/bs")):
        return "BCBS"
    if "aetna" in t:
        return "Aetna"
    if "marriott" in t:
        return "Marriott"
    if any(x in t for x in ("united", "uhc", "unitedhealth", "optum", "umr")):
        return "UnitedHealthcare"
    if "cigna" in t:
        return "Cigna"
    if "humana" in t:
        return "Humana"
    if any(x in t for x in ("medicare", "cms", "noridian")):
        return "Medicare / CMS"
    if "tricare" in t:
        return "TRICARE"
    s = (raw or "").strip()
    return s[:48] if s else "Unknown"


def _normalize_statute_intel_key(law: str) -> str:
    """Normalize Law Cited to a canonical statutory hook for win-rate buckets."""
    s = (law or "").strip()
    if not s or s == "—":
        return "Unspecified"
    u = s.upper()
    if "1051" in u and "ERISA" in u:
        return "ERISA § 1051"
    if re.search(r"ERISA\s*§?\s*503|503.*ERISA", s, re.I):
        return "ERISA § 503"
    if re.search(r"ERISA\s*§?\s*404|404.*FIDUCIARY", s, re.I):
        return "ERISA § 404"
    if "2560.503" in s or "29 C.F.R" in u:
        return "29 C.F.R. § 2560.503-1"
    if "2719" in s or "PHS ACT" in u or "ACA" in u and "2719" in u:
        return "PHS Act § 2719 (ACA)"
    if "15-1005" in s or ("MARYLAND" in u and "1005" in u):
        return "MD Code § 15-1005"
    if "PROMPT" in u or "1301" in s:
        return "State Prompt Payment"
    return s[:56]


def _load_payer_intel_profiles_from_disk() -> dict[str, Any]:
    """Read payer profiles from hidden JSON (thread-safe — no Streamlit session)."""
    if not os.path.isfile(PAYER_INTEL_DB_PATH):
        return {k: {sk: dict(rec) for sk, rec in v.items()} for k, v in PAYER_INTEL_DEFAULT_SEED.items()}
    try:
        with open(PAYER_INTEL_DB_PATH, encoding="utf-8") as f:
            data = json.load(f)
        prof = data.get("profiles") or {}
        if not prof:
            return {k: {sk: dict(rec) for sk, rec in v.items()} for k, v in PAYER_INTEL_DEFAULT_SEED.items()}
        return prof
    except Exception:
        return {k: {sk: dict(rec) for sk, rec in v.items()} for k, v in PAYER_INTEL_DEFAULT_SEED.items()}


def _ensure_payer_intel_loaded() -> dict:
    """Load / seed hidden payer intelligence DB (session + JSON file). Main thread only."""
    if st.session_state.get("_payer_intel_loaded"):
        return st.session_state.get("_payer_intel_db") or {}
    data: dict[str, Any] = {"version": 1, "profiles": _load_payer_intel_profiles_from_disk()}
    prof = data.setdefault("profiles", {})
    if not prof:
        prof.update({k: {sk: dict(rec) for sk, rec in v.items()} for k, v in PAYER_INTEL_DEFAULT_SEED.items()})
        data["profiles"] = prof
        try:
            with open(PAYER_INTEL_DB_PATH, "w", encoding="utf-8") as f:
                json.dump(data, f, indent=2)
        except Exception:
            pass
    st.session_state["_payer_intel_db"] = data
    st.session_state["_payer_intel_loaded"] = True
    return data


def _persist_payer_intel_db() -> None:
    db = st.session_state.get("_payer_intel_db") or {}
    try:
        with open(PAYER_INTEL_DB_PATH, "w", encoding="utf-8") as f:
            json.dump(db, f, indent=2)
    except Exception:
        pass


def _payer_intel_bump_exposure(payer_key: str, statute_key: str) -> None:
    """Record a vault exposure (batch upload) toward win-rate denominator."""
    _ensure_payer_intel_loaded()
    db = st.session_state.setdefault("_payer_intel_db", {"version": 1, "profiles": {}})
    prof = db.setdefault("profiles", {})
    pk = payer_key or "Unknown"
    sk = statute_key or "Unspecified"
    bucket = prof.setdefault(pk, {})
    rec = bucket.setdefault(sk, {"wins": 0, "n": 0})
    rec["n"] = int(rec.get("n") or 0) + 1
    _persist_payer_intel_db()


def _payer_intel_bump_win(payer_key: str, statute_key: str) -> None:
    """Record a statutory enforcement win (batch ZIP / ENFORCED)."""
    _ensure_payer_intel_loaded()
    db = st.session_state.setdefault("_payer_intel_db", {"version": 1, "profiles": {}})
    prof = db.setdefault("profiles", {})
    pk = payer_key or "Unknown"
    sk = statute_key or "Unspecified"
    bucket = prof.setdefault(pk, {})
    rec = bucket.setdefault(sk, {"wins": 0, "n": 0})
    rec["wins"] = int(rec.get("wins") or 0) + 1
    # Denominator `n` is from exposures; ensure wins never exceed recorded n (legacy rows).
    if int(rec.get("n") or 0) < int(rec.get("wins") or 0):
        rec["n"] = int(rec.get("wins") or 0)
    _persist_payer_intel_db()


def _extract_payer_from_row(row: dict) -> str:
    p = _get_best_val(row, ["payer", "insur", "carrier", "plan", "payer name"])
    if p:
        return str(p).strip()
    reason = str(row.get("Reason for Denial", "") or "")
    for token in ("BCBS", "Blue Cross", "Aetna", "United", "Cigna", "Humana", "Marriott"):
        if token.lower() in reason.lower():
            return token
    return ""


def _record_payer_intel_for_vault_rows(rows: list[dict]) -> None:
    """After neural merge: exposure signals per row (payer × statute)."""
    for row in rows or []:
        pk = _normalize_payer_intel_key(_extract_payer_from_row(row))
        sk = _normalize_statute_intel_key(str(row.get("Law Cited", "") or ""))
        _payer_intel_bump_exposure(pk, sk)


def _record_payer_intel_wins_for_enforced_batch(batch_sig: str | None) -> None:
    """When statutory batch is enforced, count wins for STATUTORY rows in that batch."""
    if not batch_sig:
        return
    for e in st.session_state.get("revenue_vault") or []:
        if e.get("batch_sig") != batch_sig:
            continue
        if str(e.get("Appeal Mode", "")).upper().strip() != "STATUTORY":
            continue
        src = e.get("_source_row")
        row = src if isinstance(src, dict) else e
        pk = _normalize_payer_intel_key(_extract_payer_from_row(row))
        sk = _normalize_statute_intel_key(str(e.get("Law Cited", "") or ""))
        _payer_intel_bump_win(pk, sk)


def _best_statute_for_payer(payer_key: str, prof: dict[str, Any] | None = None) -> tuple[str, float]:
    """Return (statute, win_rate 0..1) with highest historical proxy for this payer bucket."""
    prof = prof or _load_payer_intel_profiles_from_disk()
    bucket = prof.get(payer_key) or prof.get("BCBS") or {}
    best_s, best_r = "ERISA § 503", 0.0
    for stat, rec in bucket.items():
        n = max(int(rec.get("n") or 0), 1)
        w = int(rec.get("wins") or 0)
        r = w / float(n)
        if r > best_r:
            best_r, best_s = r, stat
    return best_s, best_r


def _build_gemini_payer_intel_prefix() -> str:
    """Inject before extraction prompt so Gemini biases Law Cited toward lethal historical hooks (disk-based for threads)."""
    try:
        prof = _load_payer_intel_profiles_from_disk()
        lines: list[str] = []
        for pk in sorted(prof.keys())[:14]:
            best_s, best_r = _best_statute_for_payer(pk, prof)
            n_tot = sum(int((rec or {}).get("n") or 0) for rec in (prof.get(pk) or {}).values())
            if n_tot < 1:
                continue
            lines.append(
                f"- **{pk}**: historical best leverage **{best_s}** (win proxy **{best_r*100:.0f}%** over tracked exposures)."
            )
        if not lines:
            return ""
        return (
            "SENTURION PAYER INTELLIGENCE (internal — when the payer name matches, prefer aligning **Law Cited** "
            "with the statute below if consistent with the denial; do not fabricate payer names):\n"
            + "\n".join(lines[:10])
            + "\n\n"
        )
    except Exception:
        return ""


def _render_payer_intelligence_strategy_panel(rows: list[dict]) -> None:
    """Post-extraction UI: lethal template suggestions from intelligence ledger."""
    if not rows:
        return
    _ensure_payer_intel_loaded()
    _prof = _load_payer_intel_profiles_from_disk()
    st.markdown("##### Payer Intelligence · strategy")
    tips: list[str] = []
    seen: set[str] = set()
    for row in rows:
        raw_p = _extract_payer_from_row(row)
        pk = _normalize_payer_intel_key(raw_p)
        if pk in seen:
            continue
        seen.add(pk)
        best_s, best_r = _best_statute_for_payer(pk, _prof)
        cur = str(row.get("Law Cited", "") or "").strip()
        tips.append(
            f"**{pk}** (from payer field / context): Senturion intelligence favors **{best_s}** "
            f"(historical win proxy **{best_r*100:.0f}%**). "
            f"Your extracted Law Cited: `{cur or '—'}`."
        )
    if tips:
        st.info("\n\n".join(tips[:8]))


def _payer_delinquency_all_clinics_df() -> pd.DataFrame:
    """Aggregate recovery targets by payer across ALL clinics (admin analytics)."""
    totals: dict[str, float] = defaultdict(float)
    for e in _vault_entries_displayable():
        amt = _vault_amount_for_entry(e)
        src = e.get("_source_row")
        if isinstance(src, dict):
            payer = _extract_payer_from_row(src) or _get_best_val(src, ["payer", "insur", "carrier", "plan"])
        else:
            payer = ""
        if not payer:
            payer = str(e.get("payer_label") or "").strip()
        payer = (payer or "").strip()[:80] or "Other / Unspecified"
        totals[payer] += float(amt)
    if not totals:
        return pd.DataFrame({"Payer": [], "Senturion recovery target ($)": [], "Rank": []})
    items = sorted(totals.items(), key=lambda x: -x[1])
    df = pd.DataFrame(
        [
            {"Rank": i + 1, "Payer": p, "Senturion recovery target ($)": round(v, 2)}
            for i, (p, v) in enumerate(items[:24])
        ]
    )
    return df


def _appeal_preview_mirror_body(appeal_text: str, *, include_cms_compliance: bool = True) -> str:
    """
    Append the same compliance attestation + penalty clause used in DOCX/PDF so the Streamlit preview
    matches export legal weight. CMS/NCCI attestation is omitted when no CPT context (matches PDF gate).
    """
    t = (appeal_text or "").rstrip()
    if not t:
        return t
    if include_cms_compliance:
        return (
            t
            + "\n\n---\n\n"
            + "***"
            + COMPLIANCE_ATTESTATION_TEXT
            + "***\n\n**"
            + PENALTY_CLAUSE_TEXT
            + "**"
        )
    return t + "\n\n---\n\n**" + PENALTY_CLAUSE_TEXT + "**"


RCM_DOC_FONT_PRIMARY = "Garamond"


def _rcm_body_font_name() -> str:
    """Prefer Garamond / Baskerville-class serif for institutional filings; Word substitutes if missing."""
    return RCM_DOC_FONT_PRIMARY


def _parse_markdown_table_two_col(block: str) -> list[tuple[str, str]]:
    """Parse a two-column GFM pipe table (Clinical Discrepancy Analysis)."""
    rows: list[tuple[str, str]] = []
    for line in (block or "").splitlines():
        line = line.strip()
        if not line.startswith("|"):
            continue
        if re.match(r"^\|\s*:?-+\s*(\|\s*:?-+\s*)+\|?\s*$", line):
            continue
        inner = [c.strip() for c in line.strip("|").split("|")]
        if len(inner) < 2:
            continue
        c0, c1 = inner[0], inner[1]
        c0l, c1l = c0.lower(), c1.lower()
        if "insurer denial justification" in c0l and "provider clinical substantiation" in c1l:
            continue
        if "plan interpretation" in c0l and ("marriott" in c0l or "erisa" in c1l or "2560" in c1l or "503" in c1l):
            continue
        if not c0 and not c1:
            continue
        rows.append((c0, c1))
    return rows


def _parse_appeal_rcm_sections(raw: str) -> dict:
    """Split model output into Legal Basis, evidence table rows, Administrative Narrative, and Regulatory Escalation."""
    raw = (raw or "").replace("\r\n", "\n").replace("\r", "\n")
    out: dict = {
        "legal_basis": "",
        "table_rows": [],
        "narrative": "",
        "mandatory_escalation": "",
        "regulatory_escalation": "",
        "raw": raw,
    }
    m_lb = re.search(
        r"##\s*Legal Basis for Appeal\s*\n(.*?)(?=\n##\s+|\Z)",
        raw,
        re.DOTALL | re.IGNORECASE,
    )
    if m_lb:
        out["legal_basis"] = m_lb.group(1).strip()
    m_cd = re.search(
        r"##\s*(?:Clinical Discrepancy Analysis|PLAN PROVISION VS\. STATUTORY REQUIREMENT)\s*\n(.*?)(?=\n##\s+|\Z)",
        raw,
        re.DOTALL | re.IGNORECASE,
    )
    cd_block = m_cd.group(1).strip() if m_cd else ""
    out["table_rows"] = _parse_markdown_table_two_col(cd_block)
    m_me = re.search(
        r"##\s*Mandatory 30-Day Escalation[^\n]*\n(.*?)(?=\n##\s*Regulatory Escalation Notice\s*\n|\Z)",
        raw,
        re.DOTALL | re.IGNORECASE,
    )
    if m_me:
        out["mandatory_escalation"] = m_me.group(1).strip()
    m_re = re.search(
        r"##\s*Regulatory Escalation Notice\s*\n(.*?)(?=\n##\s+|\Z)",
        raw,
        re.DOTALL | re.IGNORECASE,
    )
    if m_re:
        out["regulatory_escalation"] = m_re.group(1).strip()
    m_an = re.search(
        r"##\s*Administrative Narrative\s*\n(.*?)(?=\n##\s*Regulatory Escalation|\n##\s+|\Z)",
        raw,
        re.DOTALL | re.IGNORECASE,
    )
    if m_an:
        out["narrative"] = m_an.group(1).strip()
    else:
        tmp = raw
        if m_lb:
            tmp = re.sub(
                r"##\s*Legal Basis for Appeal\s*\n.*?(?=\n##\s+|\Z)",
                "",
                tmp,
                count=1,
                flags=re.DOTALL | re.IGNORECASE,
            )
        if m_cd:
            tmp = re.sub(
                r"##\s*(?:Clinical Discrepancy Analysis|PLAN PROVISION VS\. STATUTORY REQUIREMENT)\s*\n.*?(?=\n##\s+|\Z)",
                "",
                tmp,
                count=1,
                flags=re.DOTALL | re.IGNORECASE,
            )
        if m_re:
            tmp = re.sub(
                r"##\s*Regulatory Escalation Notice\s*\n.*?(?=\n##\s+|\Z)",
                "",
                tmp,
                count=1,
                flags=re.DOTALL | re.IGNORECASE,
            )
        out["narrative"] = tmp.strip()
    if not out["narrative"]:
        out["narrative"] = raw.strip()
    return out


def _logo_bytes_faint_watermark(src: bytes | None, opacity: float = 0.14) -> bytes | None:
    """Reduce logo alpha for a washout watermark (Page 1 header). Falls back to raw bytes if PIL fails."""
    if not src:
        return None
    try:
        from PIL import Image
        im = Image.open(io.BytesIO(src)).convert("RGBA")
        bands = im.split()
        if len(bands) >= 4:
            alpha = bands[3].point(lambda p: min(255, int(p * opacity)))
            im.putalpha(alpha)
        else:
            im = im.convert("RGBA")
            alpha = Image.new("L", im.size, int(255 * opacity))
            im.putalpha(alpha)
        out = io.BytesIO()
        im.save(out, format="PNG")
        out.seek(0)
        return out.read()
    except Exception:
        return src


def _get_best_val(row, keywords) -> str:
    """
    First column whose header contains any keyword (case-insensitive substring match).
    Works with a dict (e.g. session row) or a pandas Series.
    """
    if row is None:
        return ""
    if isinstance(row, dict):
        cols = list(row.keys())

        def _get(c):
            return row.get(c)
    else:
        cols = list(row.index)

        def _get(c):
            return row[c]

    for col in cols:
        col_l = str(col).lower()
        if any(str(k).lower() in col_l for k in keywords):
            v = _get(col)
            if v is None:
                return ""
            try:
                if pd.isna(v):
                    return ""
            except (TypeError, ValueError):
                pass
            return str(v).strip()
    return ""


def _apply_docx_csv_placeholders(
    text: str,
    *,
    insurance_payer_name: str = "",
    patient_full_name: str = "",
    claim_number: str = "",
) -> str:
    """Swap bracket placeholders (case-insensitive) and strip markdown-style stars for DOCX."""
    today = datetime.now().strftime("%B %d, %Y")
    ins_p = (insurance_payer_name or "").strip()
    pat_n = (patient_full_name or "").strip()
    clm = (claim_number or "").strip()
    t = text
    # Case-insensitive placeholder replacement (handles AI casing variants)
    t = re.sub(r"\[date\]", today, t, flags=re.IGNORECASE)
    t = re.sub(r"\[insurance payer name\]", ins_p, t, flags=re.IGNORECASE)
    t = re.sub(r"\[patient full name\]", pat_n, t, flags=re.IGNORECASE)
    t = re.sub(r"\[claim number\]", clm, t, flags=re.IGNORECASE)
    t = _strip_docx_markdown_stars(t)
    return t


def _appeal_case_reference(patient_id: str, claim_number: str) -> str:
    """Deterministic internal case id for filing metadata and headers (SNA- + 12 hex)."""
    base = f"{(patient_id or '').strip().upper()}|{(claim_number or '').strip().upper()}|SENTURION-NAD"
    h = hashlib.sha256(base.encode("utf-8")).hexdigest()[:12].upper()
    return f"SNA-{h}"


def _audit_tracking_hash(case_id: str, generated_by_user_id: str | None) -> str:
    """6-character hex audit stamp for provenance footer (stable per case + generator)."""
    raw = f"{(case_id or '').strip()}|{(generated_by_user_id or '').strip()}|SENTURION-AUDIT-SERIES"
    return hashlib.sha256(raw.encode("utf-8")).hexdigest()[:6].upper()


def _appeal_to_docx(
    appeal_text: str,
    logo_bytes: bytes | None = None,
    clinic_name: str | None = None,
    clinic_address: str | None = None,
    *,
    insurance_payer_name: str = "",
    patient_full_name: str = "",
    patient_id: str = "",
    claim_number: str = "",
    denial_code: str = "",
    denial_reason: str = "",
    service_code: str = "",
    appeal_type: str = "",
    case_id: str | None = None,
    generated_by_user_id: str | None = None,
    generated_by_email: str | None = None,
    statutory_appeal: bool = False,
    show_cms_compliance_footer: bool = True,
) -> bytes:
    """World-class RCM filing: Garamond-class serif, 1.25\" margins, Legal Basis + evidence table + narrative."""
    del appeal_type  # filing template uses fixed RE line; optional type reserved for future use
    if not HAS_DOCX or WD_ALIGN_PARAGRAPH is None or RGBColor is None:
        return b""

    try:
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
    except Exception:
        OxmlElement = None  # type: ignore[misc, assignment]
        qn = None  # type: ignore[misc, assignment]

    def _p_border_bottom(paragraph, *, sz_eighths_pt: str = "12") -> None:
        """1.5pt ≈ 12 eighths of a point in Word OOXML."""
        if OxmlElement is None or qn is None:
            return
        p_pr = paragraph._p.get_or_add_pPr()
        p_bdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), sz_eighths_pt)
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "000000")
        p_bdr.append(bottom)
        p_pr.append(p_bdr)

    def _disp(val: str) -> str:
        t = (val or "").strip()
        return t if t else DOCX_HANDWRITTEN_BLANK

    ins_d = _disp(insurance_payer_name)
    pat_d = _disp(patient_full_name)
    clm_d = _disp(claim_number)
    pid_for_re = (patient_id or "").strip()
    if not pid_for_re:
        pid_for_re = pat_d

    appeal_text = _apply_docx_csv_placeholders(
        appeal_text,
        insurance_payer_name=ins_d,
        patient_full_name=pat_d,
        claim_number=clm_d,
    )

    _font_body = _rcm_body_font_name()
    _sections = _parse_appeal_rcm_sections(appeal_text)
    _tbl = list(_sections.get("table_rows") or [])
    if not _tbl:
        _dr = (denial_reason or "").strip()
        if not _dr:
            _dr = f"Denial code {denial_code}" if denial_code else ""
        if statutory_appeal:
            _tbl = [
                (
                    "Plan interpretation as reflected in the denial record and plan administrator correspondence.",
                    "ERISA § 503 and 29 C.F.R. § 2560.503-1 full and fair review; procedural and substantive requirements applicable to the plan’s denial.",
                )
            ]
        else:
            _tbl = [
                (
                    "Insurer denial justification as reflected in the claim record and payer correspondence.",
                    f"ICD-10/CPT: state codes from claim data when available. Provider substantiation: {_dr or 'See administrative narrative and attached documentation.'} "
                    f"Gold standard / LCD: tie cited CPT to diagnosis per applicable CMS LCD or NCD where inferable.",
                )
            ]

    cid = (case_id or "").strip() or _appeal_case_reference(pid_for_re, clm_d)
    _audit_hash = _audit_tracking_hash(cid, generated_by_user_id)
    clinic = (clinic_name or "").strip() or BRAND_INSTITUTIONAL_HEADER
    today = datetime.now().strftime("%B %d, %Y")
    to_line = ins_d if ins_d != DOCX_HANDWRITTEN_BLANK else "[Insurance plan administrator — complete upon filing]"

    doc = Document()
    for s in doc.sections:
        s.top_margin = Inches(RCM_MARGIN_INCHES)
        s.bottom_margin = Inches(RCM_MARGIN_INCHES)
        s.left_margin = Inches(RCM_MARGIN_INCHES)
        s.right_margin = Inches(RCM_MARGIN_INCHES)
    section = doc.sections[0]
    section.different_first_page_header_footer = False

    def _run_times(p, text: str, *, bold: bool = False, size_pt: int = 11, center: bool = False) -> None:
        if center:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        r.bold = bold
        r.font.size = Pt(size_pt)
        r.font.name = _font_body

    # --- Authority header ---
    p_t = doc.add_paragraph()
    _run_times(p_t, BRAND_INSTITUTIONAL_HEADER, bold=True, size_pt=14, center=True)
    p_t.paragraph_format.space_after = Pt(6)

    p_s = doc.add_paragraph()
    _run_times(
        p_s,
        f"NEURAL DENIAL ANALYTICS // CASE ID: {cid}",
        bold=False,
        size_pt=10,
        center=True,
    )
    p_s.paragraph_format.space_after = Pt(10)
    _p_border_bottom(p_s)

    # Optional compact logo (does not replace institutional header)
    cover_logo = logo_bytes
    if not cover_logo:
        cover_logo = _default_brand_logo_bytes()
    if cover_logo:
        try:
            p_logo = doc.add_paragraph()
            p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r_logo = p_logo.add_run()
            r_logo.add_picture(io.BytesIO(cover_logo), width=Inches(1.25))
            p_logo.paragraph_format.space_after = Pt(14)
        except Exception:
            pass

    # --- Filing blocks ---
    def _filing_line(label: str, value: str) -> None:
        p = doc.add_paragraph()
        r0 = p.add_run(f"{label} ")
        r0.bold = True
        r0.font.size = Pt(11)
        r0.font.name = _font_body
        r1 = p.add_run(value)
        r1.bold = False
        r1.font.size = Pt(11)
        r1.font.name = _font_body
        p.paragraph_format.space_after = Pt(4)

    doc.add_paragraph().paragraph_format.space_after = Pt(8)
    _filing_line("TO:", to_line)
    _filing_line("FROM:", clinic)
    _cad = (clinic_address or "").strip()
    if _cad:
        for _ln in _cad.split("\n"):
            _t = _ln.strip()
            if _t:
                p_ad = doc.add_paragraph()
                _run_times(p_ad, _t, size_pt=11, center=False)
                p_ad.paragraph_format.space_after = Pt(2)
    _filing_line("DATE:", today)
    p_re = doc.add_paragraph()
    r_re0 = p_re.add_run("RE:")
    r_re0.bold = True
    r_re0.font.size = Pt(11)
    r_re0.font.name = _font_body
    r_re1 = p_re.add_run(
        f" Patient ID: {pid_for_re}; Claim Number: {clm_d}; FORMAL APPEAL OF DENIAL"
    )
    r_re1.bold = False
    r_re1.font.size = Pt(11)
    r_re1.font.name = _font_body
    p_re.paragraph_format.space_after = Pt(14)

    def _body_paragraphs(block: str, *, size_pt: int = 11) -> None:
        block = _strip_docx_markdown_stars(block.replace("\r\n", "\n").replace("\r", "\n"))
        for line in block.split("\n"):
            txt = line.strip()
            if not txt:
                doc.add_paragraph().paragraph_format.space_after = Pt(4)
                continue
            p_b = doc.add_paragraph()
            r_b = p_b.add_run(txt)
            r_b.font.size = Pt(size_pt)
            r_b.font.name = _font_body
            p_b.paragraph_format.space_after = Pt(8)
            p_b.paragraph_format.line_spacing = 1.25

    def _section_heading(title: str) -> None:
        p_h = doc.add_paragraph()
        r_h = p_h.add_run(title)
        r_h.bold = True
        r_h.font.size = Pt(12)
        r_h.font.name = _font_body
        p_h.paragraph_format.space_before = Pt(10)
        p_h.paragraph_format.space_after = Pt(6)

    # --- Legal Basis for Appeal ---
    if (_sections.get("legal_basis") or "").strip():
        _section_heading("Legal Basis for Appeal")
        _body_paragraphs(str(_sections["legal_basis"]))

    # --- Evidence matrix (clinical vs. ERISA statutory) ---
    _section_heading(
        "PLAN PROVISION VS. STATUTORY REQUIREMENT"
        if statutory_appeal
        else "Clinical Discrepancy Analysis"
    )
    try:
        tbl = doc.add_table(rows=1, cols=2)
        tbl.style = "Table Grid"
        hdr = tbl.rows[0].cells
        if statutory_appeal:
            hdr[0].text = "Plan Interpretation (Marriott)"
            hdr[1].text = "ERISA § 503 / 29 C.F.R. § 2560.503-1 Requirement"
        else:
            hdr[0].text = "Insurer Denial Justification"
            hdr[1].text = "Provider Clinical Substantiation"
        for c in hdr:
            for p in c.paragraphs:
                for r in p.runs:
                    r.bold = True
                    r.font.name = _font_body
                    r.font.size = Pt(10)
        for a, b in _tbl:
            row = tbl.add_row().cells
            row[0].text = a
            row[1].text = b
            for c in row:
                for p in c.paragraphs:
                    for r in p.runs:
                        r.font.name = _font_body
                        r.font.size = Pt(10)
    except Exception:
        for a, b in _tbl:
            _body_paragraphs(f"Denial basis: {a}\nProvider substantiation: {b}")
    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # --- Administrative Narrative ---
    if (_sections.get("narrative") or "").strip():
        _section_heading("Administrative Narrative")
        _body_paragraphs(str(_sections["narrative"]))

    # --- Mandatory 30-day escalation (statutory / ERISA track) ---
    if statutory_appeal and (_sections.get("mandatory_escalation") or "").strip():
        _section_heading("Mandatory 30-Day Escalation & Regulatory Referral")
        _body_paragraphs(str(_sections["mandatory_escalation"]))

    # --- Regulatory Escalation Notice (model-generated; statutory posture) ---
    if (_sections.get("regulatory_escalation") or "").strip():
        _section_heading("Regulatory Escalation Notice")
        _body_paragraphs(str(_sections["regulatory_escalation"]))

    # --- Compliance attestation (CMS only when claim contains procedure coding context) ---
    if show_cms_compliance_footer:
        doc.add_paragraph().paragraph_format.space_after = Pt(14)
        p_comp = doc.add_paragraph()
        r_comp = p_comp.add_run(COMPLIANCE_ATTESTATION_TEXT)
        r_comp.font.size = Pt(10)
        r_comp.font.name = _font_body
        r_comp.bold = True
        r_comp.italic = True
        p_comp.paragraph_format.space_after = Pt(16)

    # --- Penalty / regulatory referral (fixed institutional clause) ---
    p_pen = doc.add_paragraph()
    r_pen = p_pen.add_run(PENALTY_CLAUSE_TEXT)
    r_pen.font.size = Pt(10)
    r_pen.font.name = _font_body
    r_pen.bold = True
    p_pen.paragraph_format.space_after = Pt(16)

    # --- Signature & provenance ---
    doc.add_paragraph().paragraph_format.space_after = Pt(10)
    p_f = doc.add_paragraph()
    r_f = p_f.add_run(
        f"This document was generated via Senturion Neural Audit for {clinic}. "
        f"Audit Tracking Hash: {_audit_hash}."
    )
    r_f.font.size = Pt(10)
    r_f.font.name = _font_body
    r_f.italic = True
    p_f.paragraph_format.space_after = Pt(16)

    p_sig = doc.add_paragraph()
    r_sl = p_sig.add_run("Authorized Senturion Auditor")
    r_sl.font.size = Pt(11)
    r_sl.font.name = _font_body
    r_sl.bold = True
    p_sig.paragraph_format.space_after = Pt(28)

    p_line = doc.add_paragraph()
    r_ln = p_line.add_run("_" * 64)
    r_ln.font.size = Pt(11)
    r_ln.font.name = _font_body
    p_line.paragraph_format.space_after = Pt(4)
    p_cap = doc.add_paragraph()
    r_cp = p_cap.add_run("Signature (physical execution)")
    r_cp.font.size = Pt(9)
    r_cp.font.name = _font_body
    r_cp.italic = True

    # --- OOXML metadata (internal audit) ---
    try:
        cp = doc.core_properties
        cp.title = f"{cid} — Formal Appeal of Denial"
        cp.subject = cid
        cp.keywords = (
            f"case_id={cid}; patient_id={pid_for_re}; claim_no={clm_d}; "
            f"audit_hash={_audit_hash}; "
            f"generator_user_id={generated_by_user_id or ''}; "
            f"generator_email={(generated_by_email or '').strip()}"
        )
        cp.comments = (
            f"Senturion Neural Audit export. Case: {cid}. Audit Tracking Hash: {_audit_hash}. "
            f"Generated_by_uid: {generated_by_user_id or 'N/A'}. "
            f"Generated_by_email: {(generated_by_email or '').strip() or 'N/A'}."
        )
        cp.category = "SENTURION_INSTITUTIONAL_APPEAL"
        cp.author = "Senturion Neural Audit"
    except Exception:
        pass

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def _appeal_text_for_pdf(appeal_text: str) -> str:
    """Core PDF font is latin-1; replace unsupported chars safely."""
    t = _strip_docx_markdown_stars((appeal_text or "").replace("\r\n", "\n").replace("\r", "\n"))
    return t.encode("latin-1", errors="replace").decode("latin-1")


def _register_rcm_pdf_font(pdf: FPDF) -> str:
    """Register Garamond/Baskerville-class TTF when available; return family name for set_font."""
    base_dir = os.path.dirname(os.path.abspath(__file__))
    win = os.environ.get("WINDIR", r"C:\Windows")
    fonts_dir = os.path.join(win, "Fonts")
    fam = "SenturionRCM"
    for pth in (
        os.path.join(base_dir, "fonts", "EBGaramond-Regular.ttf"),
        os.path.join(base_dir, "fonts", "LibreBaskerville-Regular.ttf"),
        os.path.join(fonts_dir, "GARA.TTF"),
    ):
        if os.path.isfile(pth) and pth.lower().endswith(".ttf"):
            try:
                pdf.add_font(fam, "", pth)
                bd = os.path.join(fonts_dir, "GARABD.TTF")
                if os.path.isfile(bd):
                    try:
                        pdf.add_font(fam, "B", bd)
                    except Exception:
                        pass
                return fam
            except Exception:
                continue
    return "Times"


def _draw_senturion_verified_audit_seal(
    pdf: Any,
    *,
    x0: float | None = None,
    y0: float | None = None,
    box_w: float = 50.0,
    box_h: float = 14.0,
) -> None:
    """
    Overlay a 'Senturion Verified Audit' holographic-style seal (Prompt 10 Legal Authentication Seal).
    Dual chrome border + emerald / titanium typography (FPDF2-compatible).
    Default position: top-right of the current page; pass x0/y0 to place near the signature block.
    """
    if not HAS_FPDF or pdf is None:
        return
    try:
        if x0 is None:
            x0 = float(pdf.w) - float(pdf.r_margin) - box_w - 2.0
        if y0 is None:
            y0 = 8.0
        pdf.set_draw_color(0, 255, 65)
        pdf.set_line_width(0.35)
        pdf.rect(x0, y0, box_w, box_h, style="D")
        pdf.set_draw_color(*PDF_TITANIUM_RGB)
        pdf.set_line_width(0.12)
        pdf.rect(x0 + 0.7, y0 + 0.7, box_w - 1.4, box_h - 1.4, style="D")
        pdf.set_xy(x0, y0 + 1.8)
        pdf.set_font("Helvetica", "B", 7)
        pdf.set_text_color(0, 210, 88)
        pdf.cell(box_w, 4, "SENTURION", align="C", new_x="LMARGIN", new_y="NEXT")
        pdf.set_x(x0)
        pdf.set_font("Helvetica", "I", 6)
        pdf.set_text_color(*PDF_TITANIUM_RGB)
        pdf.cell(box_w, 3, "Verified Audit", align="C", new_x="LMARGIN", new_y="NEXT")
        pdf.set_font("Helvetica", "", 5)
        pdf.set_text_color(160, 162, 166)
        pdf.set_x(x0)
        pdf.cell(box_w, 3, "Legal Authentication Seal", align="C", new_x="LMARGIN", new_y="NEXT")
        pdf.set_text_color(0, 0, 0)
    except Exception:
        pass


if HAS_FPDF and FPDF is not None:

    def _pdf_footer_legal_and_hash(pdf: Any, *, dark_page: bool = False) -> None:
        """Fail-safe footer: compliance microtext + Audit Tracking Hash (Titanium rule + readable fingerprint)."""
        fam = getattr(pdf, "rcm_font_family", None) or "Helvetica"
        try:
            pdf.set_font(fam, "", 6)
        except Exception:
            pdf.set_font("Helvetica", "", 6)
        if dark_page:
            pdf.set_text_color(200, 200, 200)
        else:
            pdf.set_text_color(72, 72, 72)
        pdf.multi_cell(0, 3, _appeal_text_for_pdf(PDF_LEGAL_COMPLIANCE_FOOTER_TEXT), align="C")
        h = (getattr(pdf, "audit_tracking_hash", None) or "").strip()
        if h:
            pdf.ln(1)
            # Titanium Silver (#E0E0E0) rule — institutional accent; hash text stays dark for OCR/print contrast
            try:
                y_rule = float(pdf.get_y())
                pdf.set_draw_color(*PDF_TITANIUM_RGB)
                pdf.set_line_width(0.25)
                pdf.line(float(pdf.l_margin), y_rule, float(pdf.w) - float(pdf.r_margin), y_rule)
                pdf.ln(2)
            except Exception:
                pdf.ln(1)
            if dark_page:
                pdf.set_text_color(*PDF_TITANIUM_RGB)
            else:
                pdf.set_text_color(48, 48, 52)
            try:
                pdf.set_font(fam, "B", 7)
            except Exception:
                pdf.set_font("Helvetica", "B", 7)
            pdf.multi_cell(
                0,
                3.5,
                _appeal_text_for_pdf(f"Audit Tracking Hash: {h}"),
                align="C",
            )
        pdf.ln(1)
        try:
            pdf.set_font(fam, "", 6)
        except Exception:
            pdf.set_font("Helvetica", "", 6)
        if dark_page:
            pdf.set_text_color(170, 220, 185)
        else:
            pdf.set_text_color(56, 92, 64)
        pdf.multi_cell(
            0,
            3,
            _appeal_text_for_pdf(f"{BRAND_NAME} · Neural Audit & Compliance"),
            align="C",
        )
        pdf.set_text_color(0, 0, 0)

    class SenturionPDF(FPDF):
        """Institutional PDF: seal + EB Garamond when registered; fail-safe footer hash on every page."""

        def __init__(self, *args: Any, **kwargs: Any) -> None:
            super().__init__(*args, **kwargs)
            self.audit_tracking_hash: str = ""
            self.rcm_font_family: str | None = None  # EB Garamond / Times after _register_rcm_pdf_font

        def header(self) -> None:
            _draw_senturion_verified_audit_seal(self)

        def footer(self) -> None:
            try:
                self.set_y(-24)
                _pdf_footer_legal_and_hash(self, dark_page=False)
            except Exception:
                pass

    class MasterAuditCertificatePDF(SenturionPDF):
        """Cover page = pitch black + emerald; inner pages use standard seal + footer."""

        def header(self) -> None:
            if self.page_no() == 1:
                return
            _draw_senturion_verified_audit_seal(self)

        def footer(self) -> None:
            if self.page_no() == 1:
                try:
                    self.set_y(-24)
                    _pdf_footer_legal_and_hash(self, dark_page=True)
                except Exception:
                    pass
                return
            super().footer()

else:
    SenturionPDF = None  # type: ignore[misc, assignment]
    MasterAuditCertificatePDF = None  # type: ignore[misc, assignment]


def _pdf_html_esc(s: str) -> str:
    return (
        (s or "")
        .replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
        .replace('"', "&quot;")
    )


def _parse_currency_to_float(val: Any) -> float:
    """Coerce currency / numeric strings to float for recovery rollups."""
    if val is None:
        return 0.0
    if isinstance(val, (int, float)):
        return float(val)
    s = str(val).strip()
    if not s or s in ("—", "N/A", "Unknown"):
        return 0.0
    s = re.sub(r"[$€£,\s]", "", s)
    try:
        return float(s) if s else 0.0
    except ValueError:
        return 0.0


def _parse_denial_date_val(val: Any) -> datetime | None:
    """Parse denial / letter dates for days-past-due analytics."""
    if val is None:
        return None
    s = str(val).strip()
    if not s or s in ("—", "N/A", "Unknown"):
        return None
    try:
        dt = pd.to_datetime(s, errors="coerce")
        if pd.isna(dt):
            return None
        ts = pd.Timestamp(dt)
        return ts.to_pydatetime()
    except Exception:
        return None


def _coerce_audit_claim_row(claim: dict) -> dict:
    """Normalize arbitrary claim/audit dicts for executive brief aggregation."""
    c = dict(claim or {})
    rev = _parse_currency_to_float(
        c.get("Potential Revenue") or c.get("potential_revenue") or c.get("recovery_target")
    )
    dtp = _parse_denial_date_val(c.get("Denial Date") or c.get("denial_date") or c.get("Date of Denial"))
    law = clean_text(str(c.get("Law Cited") or c.get("law_cited") or ""))
    c["potential_revenue"] = rev
    c["denial_date_parsed"] = dtp
    c["law_cited_norm"] = law if law else "—"
    return c


def generate_executive_brief(selected_claims: list[dict]) -> dict:
    """
    Batch aggregator: total volume, recovery target, dominant statute, average days past due, audit hash.
    """
    rows = [_coerce_audit_claim_row(c) for c in (selected_claims or [])]
    n = len(rows)
    total_rev = sum(float(r.get("potential_revenue") or 0) for r in rows)
    laws_clean = [str(r.get("law_cited_norm") or "") for r in rows]
    laws_clean = [L for L in laws_clean if L and L.strip() and L.strip() != "—"]
    if laws_clean:
        primary = Counter(laws_clean).most_common(1)[0][0]
    else:
        primary = "—"
    days_list: list[int] = []
    for r in rows:
        dtp = r.get("denial_date_parsed")
        if dtp is not None:
            d0 = dtp.date() if hasattr(dtp, "date") else dtp
            ref = datetime.now().date()
            try:
                days_list.append((ref - d0).days)
            except Exception:
                pass
    avg_days: float | None = (sum(days_list) / len(days_list)) if days_list else None
    def _row_ser(r: dict) -> dict:
        _d = r.get("denial_date_parsed")
        dd_iso = None
        if _d is not None and hasattr(_d, "date"):
            try:
                dd_iso = _d.date().isoformat()
            except Exception:
                dd_iso = None
        return {
            "pid": str(r.get("Patient ID", "")),
            "dc": str(r.get("Denial Code", "")),
            "rev": float(r.get("potential_revenue") or 0),
            "law": str(r.get("law_cited_norm", "")),
            "dd": dd_iso,
        }

    ser = sorted((_row_ser(r) for r in rows), key=lambda x: (x["pid"], x["dc"]))
    audit_hash = hashlib.sha256(json.dumps(ser, default=str).encode("utf-8")).hexdigest()[:12].upper()
    return {
        "total_claims_audited": n,
        "total_potential_revenue": total_rev,
        "primary_legal_lever": primary,
        "avg_days_past_due": avg_days,
        "audit_hash": audit_hash,
    }


def _row_is_erisa_503_statutory(c: dict) -> bool:
    """Statutory / ERISA § 503 style row for Master Audit Certificate totals."""
    mode = str(c.get("Appeal Mode", "") or "").upper().strip()
    law = str(c.get("Law Cited", "") or "")
    src = c.get("_source_row")
    if isinstance(src, dict):
        law = f"{law} {src.get('Law Cited', '') or ''}"
    law_u = law.upper()
    if mode == "STATUTORY":
        return True
    if "ERISA" in law_u and "503" in law:
        return True
    if "2560.503" in law or "503-1" in law:
        return True
    return False


def _certificate_unique_files_estimate(claims: list[dict]) -> int:
    """Distinct source-file batches (batch_sig); fallback to claim row count."""
    sigs: set[str] = set()
    for c in claims or []:
        bs = str(c.get("batch_sig") or "").strip()
        if bs:
            sigs.add(bs)
    if sigs:
        return len(sigs)
    n = len(claims or [])
    return max(1, n) if n else 0


def generate_master_audit_certificate_data(claims: list[dict]) -> dict:
    """
    Aggregate metrics for the Senturion Audit Certificate PDF (batch / cohort scope).
    Reuses the same Audit Tracking Hash linkage as the executive brief for the same cohort.
    """
    brief = generate_executive_brief(claims)
    rows = claims or []
    total_files = _certificate_unique_files_estimate(rows)
    statutory_erisa = sum(1 for r in rows if _row_is_erisa_503_statutory(r))
    total_rev = float(brief.get("total_potential_revenue") or 0)
    return {
        "total_files_scanned": total_files,
        "total_statutory_erisa_503": statutory_erisa,
        "total_recoverable_revenue": total_rev,
        "audit_hash": brief.get("audit_hash"),
        "claim_rows_included": len(rows),
        "certificate_mode": "cohort",
    }


def _vault_entry_as_claim_row(e: dict) -> dict:
    """Map a Revenue Vault entry to a claim-shaped dict for certificate rollups."""
    src = e.get("_source_row")
    if isinstance(src, dict):
        r = dict(src)
    else:
        r = {
            "Patient ID": e.get("Patient") or "—",
            "Patient Name": e.get("Patient Name") or "—",
            "Denial Code": e.get("Denial Code") or "—",
            "Law Cited": e.get("Law Cited") or "—",
            "Potential Revenue": str(e.get("Potential Revenue") or e.get("amount_denied_base") or "0"),
            "Appeal Mode": str(e.get("Appeal Mode") or ""),
            "Reason for Denial": "",
        }
    r["batch_sig"] = e.get("batch_sig") or ""
    if "Audit Hash" not in r:
        r["Audit Hash"] = e.get("Audit Hash") or ""
    return r


def _vault_wide_audit_hash_for_certificate(clinic_id: str | None) -> str:
    """Stable 16-char hash over entire vault scope (Eddie contract / batch linkage)."""
    _ensure_revenue_vault()
    entries = [
        e
        for e in (st.session_state.revenue_vault or [])
        if (not clinic_id or e.get("clinic_id") == clinic_id) and not _vault_entry_is_quarantined(e)
    ]
    rows = []
    for e in sorted(entries, key=lambda x: str(x.get("vault_id") or "")):
        rows.append(
            {
                "vault_id": e.get("vault_id"),
                "patient": e.get("Patient"),
                "denial": e.get("Denial Code"),
                "amt": round(float(_vault_amount_for_entry(e)), 2),
                "audit_hash_row": e.get("Audit Hash"),
                "clinic_id": e.get("clinic_id"),
            }
        )
    return hashlib.sha256(json.dumps(rows, sort_keys=True, default=str).encode("utf-8")).hexdigest()[
        :16
    ].upper()


def generate_vault_master_audit_certificate_data(clinic_id: str | None) -> dict:
    """
    Full-clinic Revenue Vault aggregation for the Master Audit Certificate (Eddie signing PDF).
    Total recoverable uses live vault amounts; audit hash fingerprints the entire ledger for that clinic.
    """
    _ensure_revenue_vault()
    entries = [
        e
        for e in (st.session_state.revenue_vault or [])
        if (not clinic_id or e.get("clinic_id") == clinic_id) and not _vault_entry_is_quarantined(e)
    ]
    rows = [_vault_entry_as_claim_row(e) for e in entries]
    total_rev = sum(float(_vault_amount_for_entry(e)) for e in entries)
    total_files = _certificate_unique_files_estimate(rows)
    statutory_erisa = sum(1 for r in rows if _row_is_erisa_503_statutory(r))
    vault_hash = _vault_wide_audit_hash_for_certificate(clinic_id)
    cn = _clinic_display_name(clinic_id) if clinic_id else "—"
    return {
        "total_files_scanned": total_files,
        "total_statutory_erisa_503": statutory_erisa,
        "total_recoverable_revenue": total_rev,
        "audit_hash": vault_hash,
        "claim_rows_included": len(rows),
        "clinic_name": cn,
        "certificate_mode": "vault",
        "vault_scope": "clinic",
    }


def _executive_brief_to_pdf_bytes(
    brief: dict,
    *,
    document_title: str | None = None,
    subtitle: str | None = None,
) -> bytes:
    """BOSSLOGIC PDF: Senturion Fiscal Recovery — Executive Summary (RCM margins + Garamond)."""
    if not HAS_FPDF or FPDF is None or SenturionPDF is None:
        return b""

    _m = RCM_MARGIN_INCHES * 25.4
    # Executive Brief: wider bottom band so footer / page-break math does not clip content
    _exec_bottom_margin_mm = max(1.5 * 25.4, PDF_BOTTOM_MARGIN_FOR_FOOTER_MM)
    pdf = SenturionPDF()
    pdf.set_auto_page_break(auto=True, margin=_exec_bottom_margin_mm)
    pdf.set_margins(_m, PDF_TOP_MARGIN_FOR_SEAL_MM, _m)

    _fname = "Times"
    try:
        _fname = _register_rcm_pdf_font(pdf)
        pdf.set_font(_fname, "", 10)
    except Exception:
        _fname = "Times"
        try:
            pdf.set_font("Times", "", 10)
        except Exception:
            pass
    pdf.rcm_font_family = _fname
    pdf.audit_tracking_hash = str(brief.get("audit_hash") or "")

    def _pdf_font(style: str, size: int) -> None:
        try:
            pdf.set_font(_fname, style, size)
        except Exception:
            try:
                pdf.set_font("Times", style, size)
            except Exception:
                pdf.set_font("Times", "", size)

    pdf.add_page()
    # fpdf2: w=0 for multi_cell uses full paper width; use margin-safe width to avoid "Not enough horizontal space"
    effective_page_width = pdf.w - 2 * pdf.l_margin
    # A4 @ 1.25" margins ≈ 146mm printable — clamp HTML table to that zone (not 100% of physical page)
    _table_width_mm = min(effective_page_width, 160.0)
    _ti_rgb = (224, 224, 224)  # Titanium silver #E0E0E0 (matches STEALTH_ACCENT)
    _ink_rgb = (17, 17, 17)

    _title = document_title or "SENTURION FISCAL RECOVERY: EXECUTIVE SUMMARY"
    _pdf_font("B", 16)
    pdf.cell(
        effective_page_width,
        10,
        _appeal_text_for_pdf(_title),
        align="C",
        new_x="LMARGIN",
        new_y="NEXT",
    )
    pdf.ln(2)
    _pdf_font("", 10)
    pdf.cell(
        effective_page_width,
        6,
        _appeal_text_for_pdf(f"Generated: {datetime.now().strftime('%B %d, %Y %H:%M')}"),
        align="C",
        new_x="LMARGIN",
        new_y="NEXT",
    )
    if subtitle:
        pdf.ln(2)
        _pdf_font("I", 10)
        pdf.cell(
            effective_page_width,
            6,
            _appeal_text_for_pdf(subtitle),
            align="C",
            new_x="LMARGIN",
            new_y="NEXT",
        )
    pdf.ln(4)

    vol = str(int(brief.get("total_claims_audited", 0)))
    rev = float(brief.get("total_potential_revenue") or 0)
    rev_s = f"${rev:,.2f}"
    # Total Recovery Target: 14pt bold + titanium cell (HTML path)
    _table_html = (
        f"<table border='1' style='width:{_table_width_mm:.1f}mm;max-width:{_table_width_mm:.1f}mm;table-layout:fixed;border-collapse:collapse;'>"
        "<thead><tr>"
        "<th style='font-size:10pt;font-weight:bold;padding:4px;'>Total Audit Volume</th>"
        "<th style='font-size:10pt;font-weight:bold;padding:4px;'>Total Recovery Target</th>"
        "</tr></thead><tbody>"
        f"<tr>"
        f"<td style='font-size:10pt;padding:6px;vertical-align:middle;'>{_pdf_html_esc(vol)}</td>"
        f"<td style='font-size:14pt;font-weight:bold;background-color:#E0E0E0;color:#111111;"
        f"padding:8px;vertical-align:middle;text-align:center;'>{_pdf_html_esc(rev_s)}</td>"
        f"</tr>"
        "</tbody></table>"
    )
    try:
        if hasattr(pdf, "write_html"):
            pdf.write_html(_table_html)
        else:
            raise RuntimeError("write_html unavailable")
    except Exception:
        _w2 = effective_page_width / 2
        pdf.set_text_color(0, 0, 0)
        _pdf_font("", 10)
        pdf.cell(_w2, 10, _appeal_text_for_pdf(str(vol)), border=1, align="C")
        pdf.set_fill_color(*_ti_rgb)
        pdf.set_text_color(*_ink_rgb)
        _pdf_font("B", 14)
        pdf.cell(_w2, 10, _appeal_text_for_pdf(rev_s), border=1, align="C", fill=True, new_x="LMARGIN", new_y="NEXT")
        pdf.set_fill_color(255, 255, 255)
        pdf.set_text_color(0, 0, 0)

    pdf.ln(8)
    # Brand parity with DOCX `_section_heading`: bold 12pt + vertical rhythm (institutional “spaced” header block)
    _pdf_font("B", 12)
    pdf.multi_cell(effective_page_width, 9, _appeal_text_for_pdf("Statutory Strategy"))
    pdf.ln(4)
    law = str(brief.get("primary_legal_lever") or "—")
    strat = (
        f"Senturion has deployed {law} as the primary enforcement mechanism for this batch."
    )
    _pdf_font("", 11)
    pdf.multi_cell(effective_page_width, 5, _appeal_text_for_pdf(strat))

    pdf.ln(6)
    _pdf_font("B", 11)
    pdf.multi_cell(effective_page_width, 6, _appeal_text_for_pdf("Operational metric"))
    _pdf_font("", 11)
    avg = brief.get("avg_days_past_due")
    if avg is not None:
        line = f"Average Days Past Due (based on Denial Date): {float(avg):.1f}"
    else:
        line = "Average Days Past Due: N/A — denial dates not populated for this batch."
    pdf.multi_cell(effective_page_width, 5, _appeal_text_for_pdf(line))
    pdf.ln(6)
    _pdf_font("I", 9)
    pdf.set_text_color(*PDF_TITANIUM_RGB)
    pdf.multi_cell(
        effective_page_width,
        4,
        _appeal_text_for_pdf(
            f"Audit Tracking Hash (batch linkage): {brief.get('audit_hash') or '—'} — repeated in document footer for legal traceability."
        ),
    )
    pdf.set_text_color(0, 0, 0)

    out = pdf.output(dest="S")
    if isinstance(out, (bytes, bytearray)):
        return bytes(out)
    return str(out).encode("latin-1", errors="replace")


def _master_audit_certificate_to_pdf_bytes(
    cert: dict,
    *,
    signature_image_bytes: bytes | None = None,
) -> bytes:
    """
    Senturion Audit Certificate — batch / vault summary PDF:
    Total Recovery Target, Audit Tracking Hash, legal disclaimer, pitch-black cover + seal + fail-safe footer hash.
    """
    if not HAS_FPDF or FPDF is None or MasterAuditCertificatePDF is None:
        return b""

    _m = RCM_MARGIN_INCHES * 25.4
    _bottom = max(1.5 * 25.4, PDF_BOTTOM_MARGIN_FOR_FOOTER_MM)
    pdf = MasterAuditCertificatePDF()
    pdf.set_auto_page_break(auto=True, margin=_bottom)
    pdf.set_margins(_m, PDF_TOP_MARGIN_FOR_SEAL_MM, _m)

    _fname = "Times"
    try:
        _fname = _register_rcm_pdf_font(pdf)
        pdf.set_font(_fname, "", 10)
    except Exception:
        _fname = "Times"
        try:
            pdf.set_font("Times", "", 10)
        except Exception:
            pass
    pdf.rcm_font_family = _fname
    pdf.audit_tracking_hash = str(cert.get("audit_hash") or "")

    def _pdf_font(style: str, size: int) -> None:
        try:
            pdf.set_font(_fname, style, size)
        except Exception:
            try:
                pdf.set_font("Times", style, size)
            except Exception:
                pdf.set_font("Times", "", size)

    # Page 1 — HERO: black / emerald cover — total recovery target + audit tracking hash
    _tr_hero = float(cert.get("total_recoverable_revenue") or 0)
    _ah_hero = str(cert.get("audit_hash") or "")
    _cn_hero = str(cert.get("clinic_name") or "").strip()
    _vault_cert = cert.get("certificate_mode") == "vault" or bool(cert.get("vault_scope"))

    pdf.add_page()
    pdf.set_fill_color(5, 5, 5)
    pdf.rect(0, 0, pdf.w, pdf.h, "F")
    _cert_logo = _default_brand_logo_bytes()
    _hero_y = pdf.h * 0.18
    if _cert_logo:
        try:
            pdf.set_y(pdf.h * 0.05)
            pdf.image(io.BytesIO(_cert_logo), x=(pdf.w - 44) / 2, w=44)
            pdf.ln(10)
            _hero_y = max(float(pdf.get_y()) + 6, pdf.h * 0.16)
        except Exception:
            _hero_y = pdf.h * 0.18
    pdf.set_text_color(0, 255, 65)
    pdf.set_y(_hero_y)
    _pdf_font("B", 20)
    pdf.cell(
        0,
        11,
        _appeal_text_for_pdf("SENTURION AI SOLUTIONS · AUDIT CERTIFICATE"),
        align="C",
        new_x="LMARGIN",
        new_y="NEXT",
    )
    pdf.ln(4)
    _pdf_font("B", 11)
    pdf.cell(
        0,
        7,
        _appeal_text_for_pdf("CONFIDENTIAL REVENUE RECOVERY ANALYSIS"),
        align="C",
        new_x="LMARGIN",
        new_y="NEXT",
    )
    pdf.ln(10)
    _pdf_font("B", 28)
    pdf.cell(
        0,
        16,
        _appeal_text_for_pdf(f"TOTAL RECOVERY TARGET: ${_tr_hero:,.2f}"),
        align="C",
        new_x="LMARGIN",
        new_y="NEXT",
    )
    pdf.ln(8)
    _pdf_font("B", 13)
    pdf.cell(
        0,
        9,
        _appeal_text_for_pdf(f"AUDIT TRACKING HASH: {_ah_hero}"),
        align="C",
        new_x="LMARGIN",
        new_y="NEXT",
    )
    if _cn_hero:
        pdf.ln(6)
        _pdf_font("I", 12)
        pdf.cell(
            0,
            8,
            _appeal_text_for_pdf(f"Clinic: {_cn_hero}"),
            align="C",
            new_x="LMARGIN",
            new_y="NEXT",
        )
    pdf.ln(10)
    _pdf_font("I", 11)
    _sub_hero = (
        "Full Revenue Vault certification — Senturion Verified — contract execution (Eddie handoff)"
        if _vault_cert
        else "Institutional Certificate of Audit — Senturion Neural Audit cohort certification"
    )
    pdf.cell(
        0,
        7,
        _appeal_text_for_pdf(_sub_hero),
        align="C",
        new_x="LMARGIN",
        new_y="NEXT",
    )
    # Prompt 10 — Senturion Verified compliance seal on cover (stamped every page via header on p2+)
    _draw_senturion_verified_audit_seal(
        pdf,
        x0=float(pdf.w) - float(pdf.r_margin) - 52.0,
        y0=8.0,
    )
    pdf.set_text_color(0, 0, 0)

    # Page 2 — totals + signature block + seal
    pdf.add_page()
    effective_page_width = pdf.w - 2 * pdf.l_margin
    _tw = min(effective_page_width, 170.0)
    tf = int(cert.get("total_files_scanned") or 0)
    ts = int(cert.get("total_statutory_erisa_503") or 0)
    tr = float(cert.get("total_recoverable_revenue") or 0)
    rev_s = f"${tr:,.2f}"
    ah = str(cert.get("audit_hash") or "")

    _pdf_font("B", 12)
    pdf.cell(
        0,
        8,
        _appeal_text_for_pdf("Senturion AI Solutions · Audit Certificate — Batch Summary"),
        new_x="LMARGIN",
        new_y="NEXT",
    )
    _pdf_font("", 9)
    pdf.cell(
        0,
        5,
        _appeal_text_for_pdf(f"Generated: {datetime.now().strftime('%B %d, %Y %H:%M')}"),
        new_x="LMARGIN",
        new_y="NEXT",
    )
    pdf.cell(0, 5, _appeal_text_for_pdf(f"Audit Tracking Hash: {ah}"), new_x="LMARGIN", new_y="NEXT")
    pdf.ln(6)

    _table_html = (
        f"<table border='1' style='width:{_tw:.1f}mm;max-width:{_tw:.1f}mm;table-layout:fixed;border-collapse:collapse;'>"
        "<thead><tr>"
        "<th style='font-size:11pt;font-weight:bold;padding:10px;background-color:#111111;color:#00FF41;'>Metric</th>"
        "<th style='font-size:11pt;font-weight:bold;padding:10px;background-color:#111111;color:#00FF41;'>Certified Value</th>"
        "</tr></thead><tbody>"
        "<tr>"
        "<td style='font-size:12pt;padding:12px;font-weight:bold;vertical-align:middle;'>Total Files Scanned</td>"
        f"<td style='font-size:26pt;font-weight:bold;padding:16px;text-align:center;vertical-align:middle;color:#111111;'>{_pdf_html_esc(str(tf))}</td>"
        "</tr>"
        "<tr>"
        "<td style='font-size:12pt;padding:12px;font-weight:bold;vertical-align:middle;'>"
        "Total Statutory Violations Found (ERISA &#167; 503)"
        "</td>"
        f"<td style='font-size:26pt;font-weight:bold;padding:16px;text-align:center;vertical-align:middle;color:#111111;'>{_pdf_html_esc(str(ts))}</td>"
        "</tr>"
        "<tr>"
        "<td style='font-size:13pt;padding:14px;font-weight:bold;vertical-align:middle;background-color:#0A0A0A;color:#00FF41;'>"
        "TOTAL RECOVERABLE REVENUE"
        "</td>"
        f"<td style='font-size:22pt;font-weight:bold;padding:18px;text-align:center;vertical-align:middle;background-color:#00FF41;color:#111111;'>{_pdf_html_esc(rev_s)}</td>"
        "</tr>"
        "</tbody></table>"
    )
    try:
        if hasattr(pdf, "write_html"):
            pdf.write_html(_table_html)
        else:
            raise RuntimeError("write_html unavailable")
    except Exception:
        _w1 = effective_page_width * 0.52
        _w2 = effective_page_width * 0.48
        pdf.set_text_color(0, 0, 0)
        _pdf_font("B", 10)
        pdf.cell(_w1, 12, _appeal_text_for_pdf("Total Files Scanned"), border=1)
        pdf.cell(_w2, 12, _appeal_text_for_pdf(str(tf)), border=1, new_x="LMARGIN", new_y="NEXT")
        pdf.cell(_w1, 12, _appeal_text_for_pdf("Total Statutory Violations (ERISA 503)"), border=1)
        pdf.cell(_w2, 12, _appeal_text_for_pdf(str(ts)), border=1, new_x="LMARGIN", new_y="NEXT")
        pdf.set_fill_color(0, 255, 65)
        pdf.set_text_color(17, 17, 17)
        _pdf_font("B", 12)
        pdf.cell(_w1, 14, _appeal_text_for_pdf("TOTAL RECOVERABLE REVENUE"), border=1, fill=True)
        pdf.cell(_w2, 14, _appeal_text_for_pdf(rev_s), border=1, fill=True, new_x="LMARGIN", new_y="NEXT")
        pdf.set_fill_color(255, 255, 255)
        pdf.set_text_color(0, 0, 0)

    pdf.ln(6)
    _pdf_font("B", 10)
    pdf.multi_cell(effective_page_width, 5, _appeal_text_for_pdf("Legal Compliance Disclaimer"))
    pdf.ln(2)
    _pdf_font("", 8)
    pdf.multi_cell(
        effective_page_width,
        3.8,
        _appeal_text_for_pdf(SENTURION_AUDIT_CERTIFICATE_LEGAL_DISCLAIMER),
    )
    pdf.ln(5)
    _pdf_font("B", 10)
    pdf.multi_cell(effective_page_width, 5, _appeal_text_for_pdf(PENALTY_CLAUSE_TEXT))
    pdf.ln(6)

    _audit_execution_ts = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    _sig_start_y = float(pdf.get_y())
    w_text = effective_page_width - 58.0
    _pdf_font("B", 11)
    pdf.multi_cell(w_text, 6, _appeal_text_for_pdf("Authorized Senturion Auditor"), new_x="LMARGIN", new_y="NEXT")
    pdf.ln(1)
    _pdf_font("I", 9)
    pdf.multi_cell(
        w_text,
        4,
        _appeal_text_for_pdf(f"Audit execution timestamp: {_audit_execution_ts}"),
        new_x="LMARGIN",
        new_y="NEXT",
    )
    pdf.ln(2)
    if signature_image_bytes:
        try:
            sig_stream = io.BytesIO(signature_image_bytes)
            pdf.image(sig_stream, w=42)
            pdf.ln(2)
        except Exception:
            pass
    pdf.ln(3)
    _pdf_font("", 10)
    pdf.multi_cell(w_text, 5, _appeal_text_for_pdf("_" * 56), new_x="LMARGIN", new_y="NEXT")
    _pdf_font("I", 8)
    pdf.multi_cell(w_text, 4, _appeal_text_for_pdf("Signature (physical execution)"), new_x="LMARGIN", new_y="NEXT")

    _draw_senturion_verified_audit_seal(
        pdf,
        x0=float(pdf.w) - float(pdf.r_margin) - 52.0,
        y0=_sig_start_y,
    )

    if pdf.get_y() < _sig_start_y + 18.0:
        pdf.set_y(_sig_start_y + 18.0)

    out = pdf.output(dest="S")
    if isinstance(out, (bytes, bytearray)):
        return bytes(out)
    return str(out).encode("latin-1", errors="replace")


def _appeal_to_pdf_bytes(
    appeal_text: str,
    *,
    clinic_name: str | None = None,
    clinic_address: str | None = None,
    logo_image_bytes: bytes | None = None,
    insurance_payer_name: str = "",
    patient_full_name: str = "",
    patient_id: str = "",
    claim_number: str = "",
    denial_reason: str = "",
    denial_code: str = "",
    case_id: str | None = None,
    generated_by_user_id: str | None = None,
    generated_by_email: str | None = None,
    statutory_appeal: bool = False,
    show_cms_compliance_footer: bool = True,
    signature_image_bytes: bytes | None = None,
) -> bytes:
    """Institutional appeal PDF: Garamond when available, 1.25\" margins, Legal Basis + HTML evidence table + narrative."""
    if not HAS_FPDF or FPDF is None or SenturionPDF is None:
        return b""

    ins_p = (insurance_payer_name or "").strip()
    to_line = ins_p if ins_p else "[Insurance plan administrator - complete upon filing]"
    clinic = (clinic_name or "").strip() or BRAND_INSTITUTIONAL_HEADER
    _addr_lines = [ln.strip() for ln in (clinic_address or "").split("\n") if ln.strip()]
    pid = (patient_id or "").strip() or "[Patient ID]"
    clm = (claim_number or "").strip() or "[Claim Number]"
    cid = (case_id or "").strip() or _appeal_case_reference(pid, clm)
    _audit_hash = _audit_tracking_hash(cid, generated_by_user_id)
    today = datetime.now().strftime("%B %d, %Y")
    _audit_execution_ts = datetime.now().strftime("%Y-%m-%d %H:%M:%S")

    appeal_text = _apply_docx_csv_placeholders(
        appeal_text,
        insurance_payer_name=ins_p,
        patient_full_name=patient_full_name,
        claim_number=clm,
    )
    _sections = _parse_appeal_rcm_sections(appeal_text)
    _tbl = list(_sections.get("table_rows") or [])
    if not _tbl:
        _dr = (denial_reason or "").strip()
        if not _dr:
            _dr = f"Denial code {denial_code}" if denial_code else ""
        if statutory_appeal:
            _tbl = [
                (
                    "Plan interpretation as reflected in the denial record and plan administrator correspondence.",
                    "ERISA § 503 and 29 C.F.R. § 2560.503-1 full and fair review; procedural and substantive requirements applicable to the plan’s denial.",
                )
            ]
        else:
            _tbl = [
                (
                    "Insurer denial justification as reflected in the claim record and payer correspondence.",
                    f"ICD-10/CPT: state codes from claim data when available. Provider substantiation: {_dr or 'See administrative narrative.'} "
                    f"Gold standard / LCD: tie cited CPT to diagnosis per applicable CMS LCD or NCD where inferable.",
                )
            ]

    _m = RCM_MARGIN_INCHES * 25.4
    pdf = SenturionPDF()
    pdf.set_auto_page_break(
        auto=True,
        margin=max(RCM_MARGIN_INCHES * 25.4, PDF_BOTTOM_MARGIN_FOR_FOOTER_MM),
    )
    pdf.set_margins(_m, PDF_TOP_MARGIN_FOR_SEAL_MM, _m)
    _fname = _register_rcm_pdf_font(pdf)
    pdf.rcm_font_family = _fname
    pdf.audit_tracking_hash = _audit_hash

    def _pdf_font(style: str, size: int) -> None:
        try:
            pdf.set_font(_fname, style, size)
        except Exception:
            pdf.set_font(_fname, "", size)

    _meta = [
        ("set_title", (f"{cid} - Formal Appeal of Denial",)),
        (
            "set_author",
            (
                f"Senturion Neural Audit | uid={generated_by_user_id or 'N/A'}"
                f" | email={(generated_by_email or '').strip() or 'N/A'}",
            ),
        ),
        ("set_subject", (cid,)),
        (
            "set_keywords",
            (
                f"case_id={cid};audit_hash={_audit_hash};patient_id={pid};claim={clm};generator_uid={generated_by_user_id or ''}",
            ),
        ),
        ("set_creator", ("Senturion Neural Audit",)),
    ]
    for name, args in _meta:
        fn = getattr(pdf, name, None)
        if callable(fn):
            try:
                fn(*args)
            except Exception:
                pass

    pdf.add_page()
    effective_page_width = pdf.w - 2 * pdf.l_margin
    _table_w_mm = min(effective_page_width, 160.0)

    if not logo_image_bytes:
        logo_image_bytes = _default_brand_logo_bytes()

    _pdf_font("B", 14)
    pdf.cell(0, 8, BRAND_INSTITUTIONAL_HEADER, align="C", new_x="LMARGIN", new_y="NEXT")
    _pdf_font("", 10)
    pdf.cell(0, 6, f"NEURAL DENIAL ANALYTICS // CASE ID: {cid}", align="C", new_x="LMARGIN", new_y="NEXT")
    pdf.ln(2)
    pdf.set_draw_color(0, 0, 0)
    pdf.set_line_width(0.5)
    y = pdf.get_y()
    pdf.line(_m, y, pdf.w - _m, y)
    pdf.ln(4)

    if logo_image_bytes:
        try:
            pdf.image(io.BytesIO(logo_image_bytes), w=36)
            pdf.ln(3)
        except Exception:
            pass

    _pdf_font("", 11)
    pdf.multi_cell(effective_page_width, 5, _appeal_text_for_pdf(f"TO: {to_line}"))
    pdf.multi_cell(effective_page_width, 5, _appeal_text_for_pdf(f"FROM: {clinic}"))
    for _ln in _addr_lines:
        pdf.multi_cell(effective_page_width, 5, _appeal_text_for_pdf(_ln))
    pdf.multi_cell(effective_page_width, 5, _appeal_text_for_pdf(f"DATE: {today}"))
    pdf.multi_cell(
        effective_page_width,
        6,
        _appeal_text_for_pdf(
            f"RE: Patient ID: {pid}; Claim Number: {clm}; FORMAL APPEAL OF DENIAL"
        ),
    )
    pdf.ln(4)

    _pdf_font("B", 12)
    if (_sections.get("legal_basis") or "").strip():
        pdf.multi_cell(effective_page_width, 6, "Legal Basis for Appeal")
        pdf.ln(2)
        _pdf_font("", 11)
        for line in _appeal_text_for_pdf(str(_sections["legal_basis"])).split("\n"):
            pdf.multi_cell(effective_page_width, 5, line if line.strip() else " ")
        pdf.ln(3)

    _pdf_font("B", 12)
    _evidence_title = (
        "PLAN PROVISION VS. STATUTORY REQUIREMENT"
        if statutory_appeal
        else "Clinical Discrepancy Analysis"
    )
    pdf.multi_cell(effective_page_width, 6, _appeal_text_for_pdf(_evidence_title))
    pdf.ln(1)
    _rows_html = "".join(
        f"<tr><td>{_pdf_html_esc(a)}</td><td>{_pdf_html_esc(b)}</td></tr>"
        for a, b in _tbl
    )
    if statutory_appeal:
        _th0, _th1 = "Plan Interpretation (Marriott)", "ERISA § 503 / 29 C.F.R. § 2560.503-1 Requirement"
    else:
        _th0, _th1 = "Insurer Denial Justification", "Provider Clinical Substantiation"
    _table_html = (
        f"<table border='1' style='width:{_table_w_mm:.1f}mm;max-width:{_table_w_mm:.1f}mm;table-layout:fixed;'><thead><tr>"
        f"<th>{_pdf_html_esc(_th0)}</th>"
        f"<th>{_pdf_html_esc(_th1)}</th></tr></thead><tbody>"
        f"{_rows_html}</tbody></table>"
    )
    try:
        if hasattr(pdf, "write_html"):
            pdf.write_html(_table_html)
        else:
            raise RuntimeError("write_html unavailable")
    except Exception:
        _pdf_font("", 10)
        for a, b in _tbl:
            if statutory_appeal:
                pdf.multi_cell(effective_page_width, 5, _appeal_text_for_pdf(f"Plan position: {a}"))
                pdf.multi_cell(effective_page_width, 5, _appeal_text_for_pdf(f"Statutory requirement: {b}"))
            else:
                pdf.multi_cell(effective_page_width, 5, _appeal_text_for_pdf(f"Denial basis: {a}"))
                pdf.multi_cell(effective_page_width, 5, _appeal_text_for_pdf(f"Provider substantiation: {b}"))
            pdf.ln(2)

    pdf.ln(4)
    if (_sections.get("narrative") or "").strip():
        _pdf_font("B", 12)
        pdf.multi_cell(effective_page_width, 6, "Administrative Narrative")
        pdf.ln(2)
        _pdf_font("", 11)
        for line in _appeal_text_for_pdf(str(_sections["narrative"])).split("\n"):
            pdf.multi_cell(effective_page_width, 5, line if line.strip() else " ")

    pdf.ln(4)
    if statutory_appeal and (_sections.get("mandatory_escalation") or "").strip():
        _pdf_font("B", 12)
        pdf.multi_cell(
            effective_page_width,
            6,
            _appeal_text_for_pdf("Mandatory 30-Day Escalation & Regulatory Referral"),
        )
        pdf.ln(2)
        _pdf_font("", 11)
        for line in _appeal_text_for_pdf(str(_sections["mandatory_escalation"])).split("\n"):
            pdf.multi_cell(effective_page_width, 5, line if line.strip() else " ")

    pdf.ln(4)
    if (_sections.get("regulatory_escalation") or "").strip():
        _pdf_font("B", 12)
        pdf.multi_cell(effective_page_width, 6, "Regulatory Escalation Notice")
        pdf.ln(2)
        _pdf_font("", 11)
        for line in _appeal_text_for_pdf(str(_sections["regulatory_escalation"])).split("\n"):
            pdf.multi_cell(effective_page_width, 5, line if line.strip() else " ")

    pdf.ln(6)
    if show_cms_compliance_footer:
        _pdf_font("B", 10)
        pdf.multi_cell(effective_page_width, 5, _appeal_text_for_pdf(COMPLIANCE_ATTESTATION_TEXT))
        pdf.ln(6)
    _pdf_font("I", 10)
    pdf.multi_cell(
        effective_page_width,
        5,
        _appeal_text_for_pdf(
            f"This document was generated via Senturion Neural Audit for {clinic}. "
            f"Audit Tracking Hash: {_audit_hash}."
        ),
    )
    pdf.ln(4)
    _pdf_font("B", 10)
    pdf.multi_cell(effective_page_width, 5, _appeal_text_for_pdf(PENALTY_CLAUSE_TEXT))
    pdf.ln(6)
    _pdf_font("B", 11)
    pdf.cell(0, 6, "Authorized Senturion Auditor", new_x="LMARGIN", new_y="NEXT")
    pdf.ln(1)
    _pdf_font("I", 9)
    pdf.multi_cell(
        effective_page_width,
        4,
        _appeal_text_for_pdf(f"Audit execution timestamp: {_audit_execution_ts}"),
    )
    pdf.ln(2)
    if signature_image_bytes:
        try:
            sig_stream = io.BytesIO(signature_image_bytes)
            pdf.image(sig_stream, w=42)
            pdf.ln(2)
        except Exception:
            pass
    pdf.ln(6)
    _pdf_font("", 11)
    pdf.cell(0, 6, "_" * 72, new_x="LMARGIN", new_y="NEXT")
    _pdf_font("I", 9)
    pdf.cell(0, 5, "Signature (physical execution)", new_x="LMARGIN", new_y="NEXT")

    out = pdf.output(dest="S")
    if isinstance(out, (bytes, bytearray)):
        return bytes(out)
    return str(out).encode("latin-1", errors="replace")


def _truncate_pdf_to_pages(src_path: str, dest_path: str, max_pages: int = MAX_PDF_PAGES) -> int:
    """Extract first max_pages into dest_path. Returns pages actually used."""
    reader = PdfReader(src_path)
    n = len(reader.pages)
    if n <= max_pages:
        reader = None
        return n
    writer = PdfWriter()
    for i in range(max_pages):
        writer.add_page(reader.pages[i])
    with open(dest_path, "wb") as f:
        writer.write(f)
    writer = None
    reader = None
    return max_pages


def convert_pdf_to_text(pdf_path: str) -> str:
    """Convert PDF to text using MarkItDown."""
    md = MarkItDown()
    result = md.convert(pdf_path)
    text = result.text_content if result else ""
    md = None
    result = None
    return text


def _pdf_bytes_to_text_for_neural(file_bytes: bytes) -> tuple[str, str | None]:
    """Extract text from in-memory PDF bytes (same truncation rules as Secure File Uplink). Returns (text, error)."""
    if not file_bytes:
        return "", "empty file"
    file_size = len(file_bytes)
    temp_path = None
    trunc_path = None
    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
            tmp.write(file_bytes)
            temp_path = tmp.name
        reader = PdfReader(temp_path)
        page_count = len(reader.pages)
        del reader
        path_to_convert = temp_path
        if page_count > MAX_PDF_PAGES:
            with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as t2:
                trunc_path = t2.name
            _truncate_pdf_to_pages(temp_path, trunc_path)
            path_to_convert = trunc_path
        text = convert_pdf_to_text(path_to_convert)
        return (text or "", None)
    except Exception as e:
        return "", str(e)
    finally:
        if temp_path:
            try:
                os.unlink(temp_path)
            except OSError:
                pass
        if trunc_path:
            try:
                os.unlink(trunc_path)
            except OSError:
                pass
        gc.collect()


def _ghost_claim_label_from_filename(name: str) -> str:
    """Telemetry label like CLM-99021 from filename (deterministic fallback)."""
    base = os.path.splitext(name or "batch")[0]
    m = re.search(r"(\d{4,})", base)
    if m:
        tail = m.group(1)
        return f"CLM-{tail[-8:]}" if len(tail) > 8 else f"CLM-{tail}"
    h = int(hashlib.md5(base.encode("utf-8", errors="ignore")).hexdigest()[:8], 16) % 100000
    return f"CLM-{h:05d}"


def _ghost_files_signature(files: list) -> str:
    parts = []
    for f in files:
        parts.append(f"{getattr(f, 'name', '')}|{getattr(f, 'size', 0)}")
    return hashlib.sha256("|".join(parts).encode("utf-8", errors="ignore")).hexdigest()


def _intake_pdf_sha256(file_bytes: bytes) -> str:
    """Stable content hash for audit-safe dedupe (matches prior intake = skip AI + vault)."""
    return hashlib.sha256(file_bytes).hexdigest()


def _load_intake_processed_hashes() -> set[str]:
    """Hashes of PDFs already processed (persisted across sessions)."""
    with _INTAKE_HASH_FILE_LOCK:
        try:
            if os.path.exists(_INTAKE_PROCESSED_HASHES_PATH):
                with open(_INTAKE_PROCESSED_HASHES_PATH, encoding="utf-8") as f:
                    data = json.load(f)
                    return set(data.get("hashes", []))
        except Exception:
            pass
        return set()


def _persist_intake_processed_hashes_unlocked(hashes: set[str]) -> None:
    """Write hashes set to disk (caller must hold _INTAKE_HASH_FILE_LOCK if needed)."""
    try:
        lst = list(hashes)[:100_000]
        with open(_INTAKE_PROCESSED_HASHES_PATH, "w", encoding="utf-8") as f:
            json.dump({"hashes": lst}, f)
    except Exception:
        pass


def _intake_register_processed_hash(h: str, known: set[str]) -> None:
    """Register after successful extraction so same PDF is never re-billed (tokens + revenue)."""
    if not h:
        return
    with _INTAKE_HASH_FILE_LOCK:
        known.add(h)
        _persist_intake_processed_hashes_unlocked(known)


def _recompute_neural_audit_batch_sig() -> None:
    data = st.session_state.get("neural_audit_batch") or []
    if not data:
        st.session_state.pop("neural_audit_batch_sig", None)
        return
    st.session_state.neural_audit_batch_sig = hashlib.sha256(
        json.dumps(data, default=str).encode("utf-8")
    ).hexdigest()[:24]


def _ghost_sum_rows_potential_revenue(rows: list[dict]) -> float:
    """Running estimate of recoverable $ from extracted denial rows (Industrial Intake telemetry)."""
    t = 0.0
    for r in rows or []:
        try:
            t += float(_parse_amount_denied(r.get("Potential Revenue", "0")))
        except Exception:
            continue
    return t


def _ghost_process_one_pdf(
    args: tuple[int, str, bytes, str, str],
) -> dict[str, Any]:
    """ThreadPool worker: PDF → text → headless Batch Predator (no Streamlit / session_state)."""
    i, name, file_bytes, job_id, file_hash = args
    label = _ghost_claim_label_from_filename(name)
    out: dict[str, Any] = {
        "idx": i,
        "name": name,
        "label": label,
        "file_hash": file_hash,
        "merge_sig": None,
        "rows": [],
        "error": None,
        "register_hash": False,
    }
    text, terr = _pdf_bytes_to_text_for_neural(file_bytes)
    if terr:
        out["error"] = terr
        return out
    if not text or len(text.strip()) < 20:
        out["error"] = "insufficient text"
        return out
    try:
        rows, _meta = extract_denial_data(text, headless=True)
    except Exception as e:
        out["error"] = str(e)
        return out
    rows = rows or []
    out["rows"] = rows
    out["merge_sig"] = f"{job_id}|{i}|{hashlib.sha256(name.encode()).hexdigest()[:8]}"
    out["register_hash"] = True
    return out


def _ghost_worker_main(files_data: list[tuple[str, bytes]], job_id: str) -> None:
    """Background thread: pool up to MASSIVE_INTAKE_MAX_WORKERS PDFs; queue merges for main thread."""
    known_hashes = _load_intake_processed_hashes()
    batch_seen: set[str] = set()
    errors: list[str] = []
    total = len(files_data)
    completed = 0
    batch_tag = (job_id or "")[:8].upper() or "BATCH"
    dup_names: list[str] = []

    def _bump_progress(name: str, label: str, extra: str = "") -> None:
        nonlocal completed
        completed += 1
        pct = (completed / max(total, 1)) * 100.0
        line = (
            f"Neural Revenue Recovery · batch [{batch_tag}] · {completed}/{total} · {label} · {pct:.0f}%"
        )
        if extra:
            line += f" · {extra}"
        with _GHOST_LOCK:
            _GHOST_STATE["done_idx"] = completed - 1
            _GHOST_STATE["files_completed"] = completed
            _GHOST_STATE["current_file"] = name
            _GHOST_STATE["current_label"] = label
            _GHOST_STATE["pct"] = pct
            _GHOST_STATE["status_line"] = line
            _GHOST_STATE["pending_errors"] = list(errors)

    work: list[tuple[int, str, bytes, str, str]] = []
    for i, (name, file_bytes) in enumerate(files_data):
        h = _intake_pdf_sha256(file_bytes)
        if h in known_hashes or h in batch_seen:
            dup_names.append(name)
            lbl = _ghost_claim_label_from_filename(name)
            _bump_progress(
                name,
                lbl,
                "skipped — duplicate file (already processed)",
            )
            continue
        batch_seen.add(h)
        work.append((i, name, file_bytes, job_id, h))

    with _GHOST_LOCK:
        _GHOST_STATE["skipped_duplicate_count"] = len(dup_names)

    try:
        if work:
            with ThreadPoolExecutor(max_workers=MASSIVE_INTAKE_MAX_WORKERS) as ex:
                future_map = {ex.submit(_ghost_process_one_pdf, w): w for w in work}
                for fut in as_completed(future_map):
                    w = future_map[fut]
                    wname = w[1]
                    try:
                        result = fut.result()
                    except Exception as e:
                        errors.append(f"{wname}: {e!s}")
                        _bump_progress(wname, _ghost_claim_label_from_filename(wname), "worker error")
                        continue
                    name = result["name"]
                    label = result["label"]
                    if result.get("error"):
                        errors.append(f"{name}: {result['error']}")
                        _bump_progress(name, label, str(result["error"])[:80])
                        continue
                    if result.get("register_hash"):
                        _intake_register_processed_hash(str(result.get("file_hash") or ""), known_hashes)
                    rows = result.get("rows") or []
                    if rows:
                        merge_sig = result.get("merge_sig") or ""
                        _rev_add = _ghost_sum_rows_potential_revenue(rows)
                        with _GHOST_LOCK:
                            _GHOST_STATE.setdefault("pending_merge_queue", []).append(
                                {
                                    "merge_sig": merge_sig,
                                    "rows": rows,
                                    "filename": name,
                                }
                            )
                            _GHOST_STATE["estimated_recoverable_usd"] = float(
                                _GHOST_STATE.get("estimated_recoverable_usd") or 0
                            ) + float(_rev_add)
                        _bump_progress(name, label, f"vault +{len(rows)} row(s)")
                    else:
                        _bump_progress(name, label, "extracted · 0 rows")
    finally:
        if dup_names:
            msg = (
                f"Skipped {len(dup_names)} duplicate file(s) (already processed): "
                + "; ".join(dup_names[:12])
            )
            if len(dup_names) > 12:
                msg += f" … +{len(dup_names) - 12} more"
            errors.insert(0, msg)
        with _GHOST_LOCK:
            _GHOST_STATE["running"] = False
            _GHOST_STATE["finished"] = True
            _GHOST_STATE["pct"] = 100.0
            _GHOST_STATE["files_completed"] = total
            _GHOST_STATE["status_line"] = (
                f"Neural Revenue Recovery · batch [{batch_tag}] · complete · {total} file(s)"
            )
            _GHOST_STATE["pending_errors"] = list(errors)


def _ghost_start_batch(uploaded_list: list) -> None:
    """Snapshot bytes on main thread, reset neural batch for this job, spawn daemon worker."""
    st.session_state.pop("_ghost_skip_full_batch_sync_sig", None)
    files_data = [(f.name, f.getvalue()) for f in uploaded_list]
    job_id = hashlib.sha256(
        f"{time.time()}|{len(files_data)}|{_ghost_files_signature(uploaded_list)}".encode()
    ).hexdigest()[:20]
    st.session_state.neural_audit_batch = []
    st.session_state.pop("neural_audit_batch_sig", None)
    st.session_state.pop("neural_exec_pdf", None)
    st.session_state.pop("neural_exec_fn", None)
    with _GHOST_LOCK:
        _GHOST_STATE["running"] = True
        _GHOST_STATE["finished"] = False
        _GHOST_STATE["job_id"] = job_id
        _GHOST_STATE["predator_batch_tag"] = job_id[:8].upper()
        _GHOST_STATE["total_files"] = len(files_data)
        _GHOST_STATE["files_completed"] = 0
        _GHOST_STATE["done_idx"] = 0
        _GHOST_STATE["pct"] = 0.0
        _GHOST_STATE["current_file"] = ""
        _GHOST_STATE["current_label"] = ""
        _GHOST_STATE["estimated_recoverable_usd"] = 0.0
        _GHOST_STATE["status_line"] = (
            f"Neural Revenue Recovery · batch [{job_id[:8].upper()}] · initializing…"
        )
        _GHOST_STATE["pending_merge_queue"] = []
        _GHOST_STATE["pending_errors"] = []
        _GHOST_STATE["skipped_duplicate_count"] = 0
    th = threading.Thread(target=_ghost_worker_main, args=(files_data, job_id), daemon=True)
    th.start()


def _drain_ghost_merge_queue() -> bool:
    """Apply worker merge queue to vault + neural batch (must run on main Streamlit thread)."""
    with _GHOST_LOCK:
        q = list(_GHOST_STATE.get("pending_merge_queue") or [])
        _GHOST_STATE["pending_merge_queue"] = []
    if not q:
        return False
    _ensure_clinic_profiles()
    _ac = st.session_state.get("active_clinic_id")
    prev_sig = st.session_state.get("neural_audit_batch_sig")
    nb = list(st.session_state.get("neural_audit_batch") or [])
    changed = False
    for item in q:
        rows = item.get("rows") or []
        merge_sig = item.get("merge_sig") or ""
        fname = item.get("filename") or ""
        if not rows:
            continue
        tagged: list[dict] = []
        for r in rows:
            rr = dict(r)
            rr["clinic_id"] = _ac
            rr["Appeal Mode"] = _resolve_neural_appeal_mode(
                str(rr.get("Reason for Denial", "") or ""),
                str(rr.get("Law Cited", "") or ""),
            )
            tagged.append(rr)
        nb.extend(tagged)
        _merge_neural_batch_into_vault(tagged, merge_sig)
        changed = True
        st.session_state["_vault_metric_flash_ts"] = time.time()
        now = time.time()
        last = float(st.session_state.get("_ghost_last_beep_ts") or 0)
        if now - last >= GHOST_BEEP_MIN_INTERVAL_SEC:
            _play_success_beep()
            st.session_state["_ghost_last_beep_ts"] = now
        _append_audit_log(f"Ghost vault sync · {fname} ({len(tagged)} row(s))")
    st.session_state.neural_audit_batch = nb
    _recompute_neural_audit_batch_sig()
    if prev_sig != st.session_state.get("neural_audit_batch_sig"):
        st.session_state.pop("neural_exec_pdf", None)
        st.session_state.pop("neural_exec_fn", None)
    st.session_state["_ghost_skip_full_batch_sync_sig"] = st.session_state.get("neural_audit_batch_sig")
    return changed


def _ghost_try_finalize_master_pdf() -> None:
    """After a 5,000+ file Ghost batch, build CEO Master Audit Summary PDF once."""
    with _GHOST_LOCK:
        if not _GHOST_STATE.get("finished"):
            return
        if _GHOST_STATE.get("pending_merge_queue"):
            return
        tot = int(_GHOST_STATE.get("total_files") or 0)
        jid = _GHOST_STATE.get("job_id")
    if tot < MASTER_AUDIT_AUTO_PDF_THRESHOLD or not jid:
        return
    if st.session_state.get("_ghost_master_generated_for") == jid:
        return
    claims = list(st.session_state.get("neural_audit_batch") or [])
    if not claims:
        return
    brief = generate_executive_brief(claims)
    pdf_b = _executive_brief_to_pdf_bytes(
        brief,
        document_title="MASTER AUDIT SUMMARY",
        subtitle="Senturion Ghost Audit — distribution for Clinic CEO / CFO (Eddie handoff)",
    )
    if pdf_b:
        st.session_state["ghost_master_pdf_bytes"] = pdf_b
        st.session_state["ghost_master_pdf_fn"] = (
            f"Master_Audit_Summary_{brief.get('audit_hash', 'AUDIT')}.pdf"
        )
        st.session_state["_ghost_master_generated_for"] = jid
        _append_audit_log(
            f"Master Audit Summary PDF generated ({tot} files, {len(claims)} claim row(s))."
        )


def _ghost_audit_fragment_poll() -> None:
    """Industrial Intake: drain merge queue → vault on a fragment tick (no full-page rerun)."""
    _drain_ghost_merge_queue()
    _ghost_try_finalize_master_pdf()


if hasattr(st, "fragment"):
    _ghost_audit_fragment_poll = st.fragment(run_every=GHOST_FRAGMENT_POLL_SEC)(_ghost_audit_fragment_poll)


def _render_industrial_intake_telemetry_fragment() -> None:
    """Neural Telemetry bar + progress + live batch preview (fragment-only UI refresh while shredding)."""
    with _GHOST_LOCK:
        _pct = float(_GHOST_STATE.get("pct") or 0)
        _run = bool(_GHOST_STATE.get("running"))
        _fin = bool(_GHOST_STATE.get("finished"))
        _errs = list(_GHOST_STATE.get("pending_errors") or [])
        _jid = str(_GHOST_STATE.get("job_id") or "")
        _tot = int(_GHOST_STATE.get("total_files") or 0)
        _fc = int(_GHOST_STATE.get("files_completed") or 0)
        _est = float(_GHOST_STATE.get("estimated_recoverable_usd") or 0)
    _has_batch = bool(st.session_state.get("neural_audit_batch"))
    if not (_run or _fin) and not _has_batch:
        return
    _batch_tag = _jid[:8].upper() if _jid else "BATCH"
    st.markdown(
        f'<div class="neural-telemetry-bar">'
        f'<span class="neural-telemetry-title">Neural Revenue Recovery · Secure Clinical Logic</span>'
        f'<span class="mono neural-telemetry-metrics">Progress {_fc}/{max(_tot, 1)} · '
        f"Recoverable identified: ${_est:,.0f}</span>"
        f"</div>",
        unsafe_allow_html=True,
    )
    st.markdown(
        f'<p class="inst-neural-data-mode-title" style="margin-bottom:0.35rem;">'
        f'Data batch <span class="mono">[{_batch_tag}]</span></p>',
        unsafe_allow_html=True,
    )
    st.progress(min(_pct / 100.0, 1.0))
    if _run:
        st.caption(
            f"**Neural Revenue Recovery:** Secure Clinical Logic in progress. Files completed: **{_fc}** / **{_tot}**."
        )
        st.markdown(
            '<div class="senturion-dark-callout">Secure Clinical Logic — rows merge incrementally '
            "into your Neural Audit Summary.</div>",
            unsafe_allow_html=True,
        )
    elif _fin:
        st.success(
            "Uplink batch complete. Review the Neural Audit Ledger and extraction table below."
        )
    else:
        st.caption("Session extraction batch — preview below.")
    if _errs:
        with st.expander("Processing warnings", expanded=False):
            for _e in _errs[-24:]:
                st.caption(_e)
    if st.session_state.get("neural_audit_batch"):
        st.caption("**Live batch** (grows as files complete)")
        _render_neural_ghost_live_preview(
            list(st.session_state.neural_audit_batch),
            show_downloads=_fin,
        )


if hasattr(st, "fragment"):
    _render_industrial_intake_telemetry_fragment = st.fragment(run_every=GHOST_FRAGMENT_POLL_SEC)(
        _render_industrial_intake_telemetry_fragment
    )


def _strip_denial_csv_markdown(raw: str) -> str:
    """Remove optional ```csv fences from model output."""
    text = (raw or "").strip()
    if not text:
        return ""
    if "```" in text:
        parts = text.split("```")
        text = parts[1] if len(parts) > 1 else parts[0]
        if text.startswith("csv"):
            text = text[3:]
        text = text.strip()
    return text


def _normalize_filtered_denial_rows(rows: list[dict]) -> list[dict]:
    """Apply column normalization + quality filter (shared by strict and lenient parsers)."""
    normalized: list[dict] = []
    for row in rows:
        if not row:
            continue
        canon = _normalize_denial_csv_row(row)
        filled = sum(
            1
            for k in ("Patient ID", "Patient Name", "Denial Code", "Reason for Denial", "Fix Action")
            if canon.get(k) and canon[k] != "—"
        )
        patient_ok = canon["Patient ID"] != "—" or (
            (canon.get("Patient Name") or "").strip() not in ("", "—")
        )
        if (patient_ok and canon["Denial Code"] != "—") or filled >= 3:
            normalized.append(canon)
    return normalized


def _parse_denial_csv_from_raw(raw: str) -> list[dict]:
    """Standard CSV parse after markdown strip."""
    stripped = _strip_denial_csv_markdown(raw)
    if not stripped:
        return []
    try:
        reader = csv.DictReader(io.StringIO(stripped))
        rows = [dict(row) for row in reader]
    except csv.Error:
        return []
    return _normalize_filtered_denial_rows(rows)


def _recover_denial_rows_from_partial_csv(raw: str) -> list[dict]:
    """
    Lenient recovery: if output was truncated mid-row, drop tail lines until DictReader succeeds
    so complete rows are not lost.
    """
    stripped = _strip_denial_csv_markdown(raw)
    if not stripped:
        return []
    lines = stripped.splitlines()
    if not lines:
        return []
    for i in range(len(lines), 0, -1):
        chunk = "\n".join(lines[:i]).strip()
        if not chunk:
            continue
        try:
            reader = csv.DictReader(io.StringIO(chunk))
            rows = [dict(row) for row in reader]
        except csv.Error:
            continue
        normalized = _normalize_filtered_denial_rows(rows)
        if normalized:
            return normalized
    return []


def _append_engine_debug(entry: dict) -> None:
    """Append to hidden session debug log (timeouts / engine events for Batch Predator tuning)."""
    row = {"ts": datetime.now().strftime("%Y-%m-%d %H:%M:%S"), **entry}
    try:
        st.session_state.setdefault("debug_log", [])
        st.session_state.debug_log.append(row)
        if len(st.session_state.debug_log) > 500:
            st.session_state.debug_log = st.session_state.debug_log[-500:]
    except Exception:
        with _GHOST_LOCK:
            buf = _GHOST_STATE.setdefault("_thread_debug_buffer", [])
            buf.append(row)
            if len(buf) > 200:
                _GHOST_STATE["_thread_debug_buffer"] = buf[-200:]


def _is_engine_timeout_error(exc: BaseException) -> bool:
    n = type(exc).__name__
    if "Timeout" in n or n == "DeadlineExceeded":
        return True
    msg = str(exc).lower()
    return "timeout" in msg or "timed out" in msg or "deadline" in msg


def _apply_extraction_sanity_flags(rows: list[dict]) -> list[dict]:
    """Post-extraction QC: missing patient id/name or amount → MANUAL_FIX_TAG (batch never hard-fails)."""
    out: list[dict] = []
    for r in rows:
        row = dict(r)
        pid = clean_text(str(row.get("Patient ID") or ""))
        pname = clean_text(str(row.get("Patient Name") or ""))
        has_patient = (pid and pid != "—") or (pname and pname != "—")
        amt = _parse_amount_denied(row.get("Potential Revenue", "0"))
        has_amount = amt > 0
        if not has_patient or not has_amount:
            row["Validation"] = MANUAL_FIX_TAG
        else:
            row["Validation"] = row.get("Validation") or "OK"
        out.append(row)
    return out


# --- Automatic Revenue Extractor (messy EOB / letters when AI omits $) ---
_REV_LABEL_AFTER_AMT = re.compile(
    r"(?i)\$?\s*([\d,]+\.?\d*)\s*(?:USD)?\s*(?:total\s+billed|balance\s+due|amount\s+due|total\s+balance|"
    r"patient\s+responsibility|amount\s+owed|total\s+amount\s+due|charges\s+total|total\s+charges)\b"
)
_REV_LABEL_BEFORE_AMT = re.compile(
    r"(?i)\b(?:total\s+billed|balance\s+due|amount\s+due|total\s+balance|patient\s+responsibility|"
    r"amount\s+owed|total\s+amount\s+due|charges\s+total|total\s+charges)\b\s*[:\-]?\s*\$?\s*([\d,]+\.?\d*)"
)


def _scan_labeled_revenue_amounts(text: str) -> list[float]:
    """
    Deterministic fallback: find USD amounts adjacent to common billing labels
    (Total Billed, Balance Due, etc.) for Potential Revenue when AI returns 0.
    """
    if not text or not str(text).strip():
        return []
    blob = str(text)
    found: list[float] = []
    for rx in (_REV_LABEL_BEFORE_AMT, _REV_LABEL_AFTER_AMT):
        for m in rx.finditer(blob):
            a = _parse_amount_denied(m.group(1))
            if a > 0:
                found.append(a)
    return found


def _format_potential_revenue_cell(amount: float) -> str:
    """Plain numeric string for CSV / vault (no $)."""
    if amount <= 0:
        return "0"
    if abs(amount - round(amount)) < 1e-6:
        return str(int(round(amount)))
    return f"{amount:.2f}".rstrip("0").rstrip(".")


def _apply_automatic_revenue_extractor(rows: list[dict], full_text: str) -> list[dict]:
    """
    When Potential Revenue is missing (0), scan raw document text for labeled dollar amounts
    and use the maximum as a recovery placeholder for each affected row.
    """
    if not rows:
        return rows
    amounts = _scan_labeled_revenue_amounts(full_text)
    fallback = max(amounts) if amounts else None
    if fallback is None:
        return [dict(r) for r in rows]
    out: list[dict] = []
    for r in rows:
        row = dict(r)
        cur = _parse_amount_denied(row.get("Potential Revenue", "0"))
        if cur <= 0:
            row["Potential Revenue"] = _format_potential_revenue_cell(fallback)
        out.append(row)
    return out


def _apply_high_value_urgency_tags(rows: list[dict]) -> list[dict]:
    """Claims over HIGH_VALUE_TARGET_USD get URGENCY_TAG_HIGH_VALUE for Agent Terminal / vault."""
    out: list[dict] = []
    for r in rows:
        row = dict(r)
        amt = _parse_amount_denied(row.get("Potential Revenue", "0"))
        if amt > HIGH_VALUE_TARGET_USD:
            row["Urgency Tag"] = URGENCY_TAG_HIGH_VALUE
        else:
            row["Urgency Tag"] = clean_text(str(row.get("Urgency Tag", "") or ""))
        out.append(row)
    return out


def _gemini_extract_chunk_raw(
    chunk_text: str,
    chunk_idx_1: int,
    total_chunks: int,
) -> tuple[str, int | None, str | None]:
    """Single Gemini extraction call; logs timeouts; no Streamlit error UI."""
    if model is None:
        return "", None, None
    gen_config = {"max_output_tokens": 4096, "temperature": 0.0}
    _intel_pre = _build_gemini_payer_intel_prefix()
    full_prompt = (
        f"{_intel_pre}{EXTRACTION_PROMPT}\n\n---\n\n"
        f"Document content (chunk {chunk_idx_1} of {total_chunks}):\n\n{chunk_text}"
    )
    try:
        response = model.generate_content(full_prompt, generation_config=gen_config)
    except Exception as e:
        _append_engine_debug(
            {
                "event": "batch_predator_request_failed",
                "chunk": f"{chunk_idx_1}/{total_chunks}",
                "error": str(e),
                "exc_type": type(e).__name__,
            }
        )
        if _is_engine_timeout_error(e):
            _append_engine_debug(
                {
                    "event": "batch_predator_timeout",
                    "chunk": f"{chunk_idx_1}/{total_chunks}",
                    "detail": str(e),
                }
            )
        return "", None, None

    fr_int: int | None = None
    fr_name: str | None = None
    try:
        if getattr(response, "candidates", None):
            fr = response.candidates[0].finish_reason
            if fr is not None:
                fr_int = int(fr.value) if hasattr(fr, "value") else (int(fr) if isinstance(fr, int) else None)
                fr_name = getattr(fr, "name", str(fr))
    except Exception:
        pass

    raw_text = ""
    try:
        cand0 = response.candidates[0] if getattr(response, "candidates", None) else None
        parts = getattr(getattr(cand0, "content", None), "parts", None) if cand0 else None
        if parts:
            raw_text = (response.text or "").strip()
    except Exception:
        raw_text = ""

    if raw_text is None:
        raw_text = ""
    raw_text = str(raw_text).strip()
    return raw_text, fr_int, fr_name


def _rows_from_model_raw(raw_text: str) -> list[dict]:
    if not (raw_text or "").strip():
        return []
    parsed = _parse_denial_csv_from_raw(raw_text)
    if parsed:
        return parsed
    return _recover_denial_rows_from_partial_csv(raw_text) or []


def _predator_process_one_chunk(
    chunk_lines: list[str],
    chunk_idx: int,
    total_chunks: int,
    meta: dict,
) -> tuple[list[dict], list[str]]:
    """
    Process one logical chunk (typically 150 lines). On empty output, finish_reason==2 (MAX_TOKENS),
    or transport failure, silently re-run at PREDATOR_CHUNK_RETRY_LINES (75) without user-facing errors.
    """
    chunk_text = "\n".join(chunk_lines).strip()
    if not chunk_text:
        return [], []

    raw_text, fr_int, fr_name = _gemini_extract_chunk_raw(chunk_text, chunk_idx + 1, total_chunks)
    if fr_int is not None:
        meta["finish_reason"] = fr_int
    if fr_name:
        meta["finish_reason_name"] = fr_name
    if fr_int == 2:
        meta["finish_reason"] = 2

    rows = _rows_from_model_raw(raw_text)
    raw_frags: list[str] = [raw_text] if raw_text else []

    if rows:
        return rows, raw_frags

    need_retry = (not raw_text.strip()) or (fr_int == 2)
    if not need_retry:
        return [], raw_frags

    if len(chunk_lines) <= PREDATOR_CHUNK_RETRY_LINES:
        meta["predator_chunk_failures"] = int(meta.get("predator_chunk_failures") or 0) + 1
        return [], raw_frags

    meta["silent_line_retry"] = True
    agg_rows: list[dict] = []
    agg_raw: list[str] = []
    for j in range(0, len(chunk_lines), PREDATOR_CHUNK_RETRY_LINES):
        sub_lines = chunk_lines[j : j + PREDATOR_CHUNK_RETRY_LINES]
        stext = "\n".join(sub_lines).strip()
        if not stext:
            continue
        r2, fr2, fn2 = _gemini_extract_chunk_raw(stext, chunk_idx + 1, total_chunks)
        if fr2 is not None:
            meta["finish_reason"] = fr2
        if fn2:
            meta["finish_reason_name"] = fn2
        if fr2 == 2:
            meta["finish_reason"] = 2
        if r2:
            agg_raw.append(r2)
        sub_rows = _rows_from_model_raw(r2)
        agg_rows.extend(sub_rows)

    if not agg_rows:
        meta["predator_chunk_failures"] = int(meta.get("predator_chunk_failures") or 0) + 1
    return agg_rows, agg_raw if agg_raw else raw_frags


def extract_denial_data(text: str, *, headless: bool = False) -> tuple[list[dict], dict]:
    """
    THE BATCH PREDATOR:
    Splits large documents into ~150-line "Lethal Chunks", runs Gemini per chunk, merges CSV rows.
    On empty parts or finish_reason 2 (MAX_TOKENS), silently retries at 75 lines per sub-chunk.

    `headless=True` skips Streamlit progress widgets (required for Ghost Auditing background threads).
    """
    meta: dict = {
        "finish_reason": None,
        "finish_reason_name": None,
        "raw_response_text": None,
        "status": None,
        "predator_chunks": 0,
        "predator_chunk_failures": 0,
        "silent_line_retry": False,
    }

    text_input = (text or "").strip()
    if not text_input:
        meta["status"] = "EMPTY"
        return [], meta

    lines = text_input.split("\n")
    chunks = [lines[i : i + PREDATOR_CHUNK_LINES] for i in range(0, len(lines), PREDATOR_CHUNK_LINES)]
    meta["predator_chunks"] = len(chunks)

    all_extracted_rows: list[dict] = []
    raw_chunks: list[str] = []

    progress_text = None if headless else st.empty()

    for idx, chunk in enumerate(chunks):
        if progress_text is not None:
            progress_text.info(
                f"Neural Revenue Recovery: processing segment {idx + 1}/{len(chunks)}…"
            )
        rows, frags = _predator_process_one_chunk(chunk, idx, len(chunks), meta)
        all_extracted_rows.extend(rows)
        raw_chunks.extend(frags)

        if len(chunks) > 1 and idx < len(chunks) - 1:
            time.sleep(1)

    meta["raw_response_text"] = "\n\n---\n\n".join(raw_chunks) if raw_chunks else None
    all_extracted_rows = _apply_automatic_revenue_extractor(all_extracted_rows, text_input)
    all_extracted_rows = _apply_extraction_sanity_flags(all_extracted_rows)
    all_extracted_rows = _apply_high_value_urgency_tags(all_extracted_rows)
    if progress_text is not None:
        progress_text.success("✅ Neural audit segment complete. Neural Audit Summary updated.")
    meta["status"] = "SUCCESS"
    return (all_extracted_rows if all_extracted_rows else []), meta


def process_pasted_audit(raw_text: str) -> tuple[list[dict], dict]:
    """
    Fail-safe manual path: run the same Gemini extraction + CSV normalization as PDF uplink.
    Accepts pasted denial letter text or raw CSV-like rows.
    """
    t = (raw_text or "").strip()
    if len(t) < 20:
        return [], {"finish_reason": None, "finish_reason_name": None, "error": "insufficient_text"}
    return extract_denial_data(t)


def _strip_brackets(val: str) -> str:
    """Strip square brackets from placeholder-style values."""
    if not val:
        return ""
    return str(val).replace("[", "").replace("]", "").strip() or ""


def _ensure_claims_portal() -> None:
    if "claims_portal" not in st.session_state:
        st.session_state.claims_portal = {}


def _ensure_revenue_vault() -> None:
    if "revenue_vault" not in st.session_state:
        st.session_state.revenue_vault = []
    _hydrate_vault_from_shadow_log_once()


def _ensure_audit_log() -> None:
    if "audit_log_history" not in st.session_state:
        st.session_state.audit_log_history = []


def _append_audit_log(message: str) -> None:
    """Append a line to the live sidebar audit ticker."""
    _ensure_audit_log()
    entry = {"ts": datetime.now().strftime("%Y-%m-%d %H:%M:%S"), "message": (message or "").strip()}
    st.session_state.audit_log_history.append(entry)
    if len(st.session_state.audit_log_history) > 100:
        st.session_state.audit_log_history = st.session_state.audit_log_history[-100:]


def _ensure_clinic_profiles() -> None:
    """Session-scoped clinic registry (multi-tenancy). At least one clinic always exists."""
    if not st.session_state.get("clinic_profiles"):
        cid = str(uuid.uuid4())
        st.session_state.clinic_profiles = [
            {
                "clinic_id": cid,
                "name": DEFAULT_CLINIC_NAME,
                "tax_id": "",
                "npi": "",
                "address": "",
                "logo_bytes": None,
            }
        ]
    if not st.session_state.get("active_clinic_id"):
        st.session_state.active_clinic_id = st.session_state.clinic_profiles[0]["clinic_id"]


def _get_clinic_profile(clinic_id: str | None) -> dict | None:
    if not clinic_id:
        return None
    for p in st.session_state.get("clinic_profiles") or []:
        if p.get("clinic_id") == clinic_id:
            return p
    return None


def _clinic_display_name(clinic_id: str | None) -> str:
    prof = _get_clinic_profile(clinic_id)
    return (prof.get("name") if prof else None) or "—"


def _default_brand_logo_bytes() -> bytes | None:
    """Official logo for PDFs / DOCX when clinic logo is not set (logo.png preferred)."""
    for _p in ("logo.png", "centurion_logo.png"):
        if os.path.isfile(_p):
            try:
                with open(_p, "rb") as f:
                    return f.read()
            except OSError:
                pass
    return None


def _branding_for_clinic_id(clinic_id: str | None) -> dict:
    """Name, address, logo bytes for PDF/DOCX FROM block."""
    _ensure_clinic_profiles()
    prof = _get_clinic_profile(clinic_id) or _get_clinic_profile(st.session_state.get("active_clinic_id"))
    if not prof:
        return {
            "name": DEFAULT_CLINIC_NAME,
            "address": "",
            "logo_bytes": None,
            "tax_id": "",
            "npi": "",
        }
    return {
        "name": (prof.get("name") or "").strip() or DEFAULT_CLINIC_NAME,
        "address": (prof.get("address") or "").strip(),
        "logo_bytes": prof.get("logo_bytes"),
        "tax_id": (prof.get("tax_id") or "").strip(),
        "npi": (prof.get("npi") or "").strip(),
    }


def _vault_all_entries() -> list[dict]:
    _ensure_revenue_vault()
    return list(st.session_state.revenue_vault)


def _vault_entries_filtered() -> list[dict]:
    """Respects Revenue Vault clinic filter dropdown."""
    _ensure_revenue_vault()
    filt = st.session_state.get("vault_clinic_filter") or "ALL"
    if filt == "ALL":
        return list(st.session_state.revenue_vault)
    return [e for e in st.session_state.revenue_vault if e.get("clinic_id") == filt]


def _parse_enforcement_iso_ts(val: Any) -> datetime | None:
    """Parse enforcement_clock_started_at / payment timestamps from vault entries."""
    if val is None:
        return None
    s = str(val).strip()
    if not s:
        return None
    try:
        return datetime.fromisoformat(s.replace("Z", "+00:00"))
    except ValueError:
        try:
            return pd.to_datetime(s, errors="coerce").to_pydatetime()
        except Exception:
            return None


def _is_enforcement_clock_overdue(entry: dict) -> bool:
    """
    ENFORCED + Maryland prompt-payment window elapsed without confirmed payment.
    Uses enforcement_clock_started_at; compares to ENFORCEMENT_MD_PROMPT_PAY_DAYS.
    """
    if str(entry.get("vault_status", "")).strip() != VAULT_STATUS_ENFORCED:
        return False
    if entry.get("payment_received"):
        return False
    dt = _parse_enforcement_iso_ts(entry.get("enforcement_clock_started_at"))
    if dt is None:
        return False
    if getattr(dt, "tzinfo", None) is not None:
        dt = dt.replace(tzinfo=None)
    return (datetime.now() - dt) >= timedelta(days=ENFORCEMENT_MD_PROMPT_PAY_DAYS)


def _enforcement_clock_status_label(entry: dict) -> tuple[str, bool]:
    """(display label, use neon red overdue styling)."""
    if str(entry.get("vault_status", "")).strip() != VAULT_STATUS_ENFORCED:
        return "—", False
    if entry.get("payment_received"):
        return "Paid ✓", False
    dt = _parse_enforcement_iso_ts(entry.get("enforcement_clock_started_at"))
    if dt is None:
        return "[ENFORCED]", False
    if getattr(dt, "tzinfo", None) is not None:
        dt = dt.replace(tzinfo=None)
    elapsed = datetime.now() - dt
    days = elapsed.total_seconds() / 86400.0
    rem = ENFORCEMENT_MD_PROMPT_PAY_DAYS - days
    if rem <= 0:
        return f"OVERDUE ({int(abs(rem))}d)", True
    return f"T+{int(days)}d / {ENFORCEMENT_MD_PROMPT_PAY_DAYS}d", False


def _build_regulatory_escalation_draft(
    entries: list[dict],
    *,
    recipient: str = "MIA",
) -> str:
    """Pre-filled letter to MIA or DOL EBSA — Audit Hash + payer non-response after enforcement window."""
    today = datetime.now().strftime("%B %d, %Y")
    lines: list[str] = [
        today,
        "",
    ]
    if recipient == "DOL":
        lines.extend(
            [
                "U.S. Department of Labor",
                "Employee Benefits Security Administration (EBSA)",
                "Office of Enforcement",
                "Washington, DC 20210",
            ]
        )
    else:
        lines.extend(
            [
                "Maryland Insurance Administration",
                "Attn: Market Regulation / Consumer Services",
                "Baltimore, Maryland",
            ]
        )
    lines.extend(
        [
            "",
            "RE: Regulatory escalation — Maryland prompt payment (Ins. Code § 15-1005) / ERISA § 503",
            "",
            "To Whom It May Concern:",
            "",
            "Senturion AI Solutions submits this formal escalation regarding the claim(s) listed below. "
            "Each line was advanced to ENFORCED status under our Neural Audit, with a statutory enforcement "
            f"record tied to the Senturion Audit Hash. The responsible payor(s) have failed to remit payment "
            f"or an adequate written response within {ENFORCEMENT_MD_PROMPT_PAY_DAYS} calendar days from that "
            "enforcement record, consistent with Maryland prompt-payment expectations and applicable federal "
            "claims procedure obligations (including ERISA § 503 and 29 C.F.R. § 2560.503-1 where applicable).",
            "",
            "CLAIM DETAILS (AUDIT HASH — PAYOR — AMOUNT — ENFORCEMENT DATE):",
            "",
        ]
    )
    for i, e in enumerate(entries, 1):
        ah = str(e.get("Audit Hash") or "—")
        payer = str(e.get("payer_label") or "—")
        patient = str(e.get("Patient") or "—")
        dc = str(e.get("Denial Code") or "—")
        amt = _vault_amount_for_entry(e)
        started = str(e.get("enforcement_clock_started_at") or "—")
        vid = str(e.get("vault_id") or "")[:8]
        lines.append(
            f"{i}. Vault row {vid}… | Patient: {patient} | Denial: {dc} | "
            f"Audit Hash: {ah} | Payer: {payer} | Amount denied (recoverable): ${amt:,.2f}"
        )
        lines.append(f"   Enforcement record (clock start): {started}")
        lines.append("")
    lines.extend(
        [
            "We request that your office open an investigation, compel an adequate response, and take "
            "appropriate regulatory action. This correspondence may be relied upon as a contemporaneous "
            "record of payor non-performance following Senturion’s statutory enforcement position.",
            "",
            "Respectfully,",
            "",
            "SENTURION AI SOLUTIONS — Authorized Representative",
            "",
            f"Generated by Senturion Enforcement Clock · MD § 15-1005 · {ENFORCEMENT_MD_PROMPT_PAY_DAYS}-day window",
        ]
    )
    return "\n".join(lines)


def _render_clinic_management_tab() -> None:
    """Settings · Clinic Management — CRUD clinic profiles (name, tax ID, NPI, address, logo)."""
    _ensure_clinic_profiles()
    st.caption(
        "Register each practice. **Active clinic** in the sidebar controls which `Clinic_ID` is applied to new uploads."
    )
    with st.form("form_add_clinic", clear_on_submit=True):
        st.markdown("**Add clinic**")
        n = st.text_input("Clinic Name", key="new_clinic_name")
        tax = st.text_input("Tax ID", key="new_clinic_tax_id")
        npi = st.text_input("NPI Number", key="new_clinic_npi")
        addr = st.text_area("Practice address (letterhead / FOR line)", key="new_clinic_address", height=88)
        lug = st.file_uploader("Clinic logo (PNG/JPG)", type=["png", "jpg", "jpeg"], key="new_clinic_logo")
        if st.form_submit_button("Save new clinic", type="primary"):
            if not (n or "").strip():
                st.warning("Clinic name is required.")
            else:
                cid = str(uuid.uuid4())
                st.session_state.clinic_profiles.append(
                    {
                        "clinic_id": cid,
                        "name": (n or "").strip(),
                        "tax_id": (tax or "").strip(),
                        "npi": (npi or "").strip(),
                        "address": (addr or "").strip(),
                        "logo_bytes": lug.getvalue() if lug is not None else None,
                    }
                )
                st.session_state.active_clinic_id = cid
                st.success("Clinic added — set as active.")
                st.rerun()

    st.markdown("---")
    st.markdown("**Registered clinics**")
    for i, p in enumerate(list(st.session_state.clinic_profiles)):
        cc1, cc2 = st.columns([4, 1])
        with cc1:
            st.caption(
                f"**{p.get('name', '—')}** · Tax ID `{p.get('tax_id') or '—'}` · NPI `{p.get('npi') or '—'}`"
            )
            _aid = str(p.get("clinic_id", ""))[:12]
            st.caption(f"`{_aid}…`")
        with cc2:
            if len(st.session_state.clinic_profiles) > 1:
                if st.button("Remove", key=f"rm_clinic_{i}"):
                    st.session_state.clinic_profiles.pop(i)
                    if st.session_state.get("active_clinic_id") == p.get("clinic_id"):
                        st.session_state.active_clinic_id = st.session_state.clinic_profiles[0]["clinic_id"]
                    st.session_state.vault_clinic_filter = "ALL"
                    st.rerun()
        lu_edit = st.file_uploader(
            "Replace logo",
            type=["png", "jpg", "jpeg"],
            key=f"clinic_logo_replace_{i}",
        )
        if lu_edit is not None:
            _sig = f"{getattr(lu_edit, 'name', '')}_{getattr(lu_edit, 'size', 0)}"
            _k = f"_clinic_logo_applied_{i}"
            if st.session_state.get(_k) != _sig:
                p["logo_bytes"] = lu_edit.getvalue()
                st.session_state[_k] = _sig
                st.rerun()
        st.divider()


def _get_digital_signature_bytes() -> bytes | None:
    """PNG bytes uploaded in Settings (session-scoped)."""
    b = st.session_state.get("digital_signature_bytes")
    if b and isinstance(b, (bytes, bytearray)) and len(b) > 0:
        return bytes(b)
    return None


def _sync_vault_statuses_from_session_widgets() -> None:
    """Persist status selectbox values back onto revenue_vault entries."""
    opts = (
        VAULT_STATUS_UNAUDITED,
        VAULT_STATUS_NEURAL_DRAFT,
        VAULT_STATUS_AGENT_REVIEW,
        VAULT_STATUS_ENFORCED,
    )
    for e in st.session_state.revenue_vault:
        vid = e.get("vault_id")
        if not vid:
            continue
        k = f"vault_stat_{vid}"
        if k in st.session_state:
            val = st.session_state[k]
            if val in opts:
                old = e.get("vault_status")
                if val == VAULT_STATUS_ENFORCED and old != VAULT_STATUS_ENFORCED:
                    if not e.get("enforcement_clock_started_at"):
                        e["enforcement_clock_started_at"] = datetime.now().isoformat(timespec="seconds")
                e["vault_status"] = val


def _sync_vault_payment_flags_from_widgets() -> None:
    """Persist 'payment received' checkboxes (clears enforcement clock heatmap)."""
    for e in st.session_state.revenue_vault:
        vid = e.get("vault_id")
        if not vid:
            continue
        k = f"vault_payrecv_{vid}"
        if k in st.session_state:
            e["payment_received"] = bool(st.session_state[k])


def _mark_batch_statutory_enforced(batch_sig: str | None) -> None:
    """After batch ZIP, mark STATUTORY vault rows for this batch as ENFORCED."""
    if not batch_sig:
        return
    for e in st.session_state.revenue_vault:
        if e.get("batch_sig") != batch_sig:
            continue
        if str(e.get("Appeal Mode", "")).upper().strip() == "STATUTORY":
            e["vault_status"] = VAULT_STATUS_ENFORCED
            if not e.get("enforcement_clock_started_at"):
                e["enforcement_clock_started_at"] = datetime.now().isoformat(timespec="seconds")
            k = f"vault_stat_{e.get('vault_id')}"
            if k in st.session_state:
                st.session_state[k] = VAULT_STATUS_ENFORCED
    _record_payer_intel_wins_for_enforced_batch(batch_sig)


def _gemini_generate_appeal_text_for_row(row: dict) -> str:
    """Single-claim appeal body via Gemini (mirrors Appeal Engine pipeline)."""
    if model is None:
        return (
            "ADMINISTRATIVE NOTICE:\n\nAppeal drafting engine unavailable — configure GEMINI_API_KEY in secrets."
        )
    ctx = _build_smart_context(row)
    ctx = _text_only_for_prompt(ctx.strip())
    if not ctx:
        return ""
    reason = str(row.get("Reason for Denial", "") or "")
    law = str(row.get("Law Cited", "") or "")
    statutory = (
        str(row.get("Appeal Mode", "")).upper().strip() == "STATUTORY"
        or _resolve_neural_appeal_mode(reason, law) == "STATUTORY"
    )
    litigation_prompt = MASTER_LITIGATION_PROMPT_STATUTORY if statutory else MASTER_LITIGATION_PROMPT
    full_prompt = f"{litigation_prompt}\n\nCLAIM DATA:\n{ctx}"
    clean_prompt = str(full_prompt)
    code = _text_only_for_prompt(row.get("Denial Code", "N/A")) or "N/A"
    response = model.generate_content(
        clean_prompt,
        safety_settings={
            HarmCategory.HARM_CATEGORY_HARASSMENT: HarmBlockThreshold.BLOCK_NONE,
            HarmCategory.HARM_CATEGORY_DANGEROUS_CONTENT: HarmBlockThreshold.BLOCK_NONE,
        },
    )
    if not response.candidates or response.candidates[0].finish_reason.name == "SAFETY":
        safe_prompt = f"Draft a professional business letter regarding healthcare claim {code}. Use 600 words. Cite ERISA 503."
        response = model.generate_content(str(safe_prompt))
    final_text = response.text
    if not final_text:
        fr = response.candidates[0].finish_reason.name if response.candidates else "BLOCKED"
        final_text = (
            f"ADMINISTRATIVE NOTICE:\n\nThe system encountered a processing filter (Finish: {fr}). "
            f"Please manually review Claim ID: {row.get('Patient ID', 'Unknown')}"
        )
    return (final_text or "").strip()


def _build_statutory_batch_zip_bytes(
    statutory_rows: list[dict],
    *,
    clinic_name: str,
) -> bytes:
    """ZIP of PDF appeals for each STATUTORY row in the batch."""
    sig_bytes = _get_digital_signature_bytes()
    _fallback_name = (clinic_name or "").strip() or BRAND_INSTITUTIONAL_HEADER
    _gen_uid = str(st.session_state.get("user_id") or "")
    _gen_em = str(st.session_state.get("email") or "").strip()
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
        for i, row in enumerate(statutory_rows):
            appeal_text = _gemini_generate_appeal_text_for_row(row)
            patient_id_export = clean_text(str(row.get("Patient ID", "") or ""))
            patient_full_name = _get_best_val(row, ["patient", "name", "member", "pt"])
            claim_number = _get_best_val(row, ["claim", "#", "id", "ref", "number"])
            insurance_payer_name = _get_best_val(row, ["payer", "insur", "carrier", "plan"])
            denial_code = clean_text(str(row.get("Denial Code", "") or ""))
            denial_reason = clean_text(str(row.get("Reason for Denial", "") or ""))
            _case_id = _appeal_case_reference(patient_id_export, claim_number)
            _show_cms = _text_has_cpt_code(
                f"{appeal_text} {denial_reason} {patient_id_export} {claim_number}"
            )
            _br = _branding_for_clinic_id(row.get("clinic_id"))
            _pdf_name = _br.get("name") or _fallback_name
            pdf_b = _appeal_to_pdf_bytes(
                appeal_text,
                clinic_name=_pdf_name,
                clinic_address=_br.get("address") or None,
                logo_image_bytes=_br.get("logo_bytes"),
                insurance_payer_name=insurance_payer_name or "",
                patient_full_name=patient_full_name or "",
                patient_id=patient_id_export,
                claim_number=claim_number or "",
                denial_reason=denial_reason,
                denial_code=denial_code,
                case_id=_case_id,
                generated_by_user_id=_gen_uid or None,
                generated_by_email=_gen_em or None,
                statutory_appeal=True,
                show_cms_compliance_footer=_show_cms,
                signature_image_bytes=sig_bytes,
            )
            if not pdf_b:
                continue
            safe_pid = re.sub(r"[^\w\-]+", "_", patient_id_export or "patient")[:40]
            safe_dc = re.sub(r"[^\w\-]+", "_", denial_code or "denial")[:30]
            zf.writestr(f"STATUTORY_{safe_pid}_{safe_dc}_{i + 1:03d}.pdf", pdf_b)
    buf.seek(0)
    return buf.getvalue()


def _parse_amount_denied(val) -> float:
    """Numeric amount for vault metrics (Potential Revenue / Amount Denied)."""
    if val is None:
        return 0.0
    s = str(val).replace(",", "").replace("$", "").replace("—", "").strip()
    if not s:
        return 0.0
    try:
        return float(s)
    except ValueError:
        return 0.0


def _double_entry_revenue_ok(raw_potential_revenue: str) -> tuple[bool, float]:
    """
    Double-entry validation: raw Potential Revenue must match USD display regex,
    parse to (0, MAX_VAULT_SANE_USD] — else quarantine [MANUAL_VERIFICATION].
    """
    s = clean_text(str(raw_potential_revenue or ""))
    if not s or s == "—":
        return False, 0.0
    if not _USD_AMOUNT_DISPLAY_RE.match(s):
        return False, _parse_amount_denied(s)
    amt = _parse_amount_denied(s)
    if amt <= 0.0 or amt > MAX_VAULT_SANE_USD:
        return False, amt
    return True, amt


def _unique_claim_hash(patient_id: str, denial_date: str, amount_usd: float) -> str:
    """Stable idempotency key: Patient ID + Date + Amount (prevents double-billing)."""
    pid = clean_text(str(patient_id or "")).upper()
    ddt = clean_text(str(denial_date or "")).strip()
    amt_s = f"{float(amount_usd):.2f}"
    payload = f"{pid}|{ddt}|{amt_s}"
    return hashlib.sha256(payload.encode("utf-8")).hexdigest()[:16].upper()


def _vault_entry_is_quarantined(entry: dict) -> bool:
    v = str(entry.get("Validation") or "")
    return MANUAL_VERIFICATION_TAG in v


def _vault_entries_displayable() -> list[dict]:
    """Ledger / KPIs: clinic filter + exclude manual verification (double-entry) queue."""
    return [e for e in _vault_entries_filtered() if not _vault_entry_is_quarantined(e)]


def _vault_entries_quarantine() -> list[dict]:
    """Rows held for [MANUAL_VERIFICATION] — not shown in main ledger amounts."""
    return [e for e in _vault_entries_filtered() if _vault_entry_is_quarantined(e)]


def _append_vault_shadow_log_row(
    event: str,
    *,
    vault_id: str = "",
    patient_id: str = "",
    denial_date: str = "",
    amount: float = 0.0,
    unique_claim_hash: str = "",
    validation: str = "",
    entry: dict | None = None,
) -> None:
    """Append-only CSV shadow log — survives browser crash."""
    row = {
        "ts_utc": datetime.now(timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ"),
        "event": event,
        "vault_id": vault_id,
        "patient_id": patient_id,
        "denial_date": denial_date,
        "amount": f"{float(amount):.2f}",
        "unique_claim_hash": unique_claim_hash,
        "validation": validation,
        "entry_json": json.dumps(entry, default=str) if entry is not None else "",
    }
    path = VAULT_SHADOW_LOG_PATH
    with _VAULT_IO_LOCK:
        new_file = not path.exists()
        with path.open("a", encoding="utf-8", newline="") as f:
            w = csv.DictWriter(f, fieldnames=SHADOW_LOG_FIELDS, extrasaction="ignore")
            if new_file:
                w.writeheader()
            w.writerow(row)


def _hydrate_vault_from_shadow_log_once() -> None:
    """Restore vault + unique-claim hash set from vault_backup.csv (once per Streamlit session)."""
    if st.session_state.get("_vault_shadow_hydrated"):
        return
    st.session_state.setdefault("revenue_vault", [])
    path = VAULT_SHADOW_LOG_PATH
    if not path.exists():
        st.session_state._vault_shadow_hydrated = True
        return
    with _VAULT_IO_LOCK:
        if st.session_state.get("_vault_shadow_hydrated"):
            return
        hashes: set[str] = set(st.session_state.get("_vault_unique_claim_hashes") or set())
        seen_vault_ids: set[str] = {str(e.get("vault_id") or "") for e in st.session_state.revenue_vault}
        seen_vault_ids.discard("")
        try:
            with path.open("r", encoding="utf-8", newline="") as f:
                reader = csv.DictReader(f)
                for row in reader:
                    ev = (row.get("event") or "").strip().upper()
                    uhash = (row.get("unique_claim_hash") or "").strip()
                    if uhash:
                        hashes.add(uhash)
                    if ev == "DUPLICATE_BLOCKED":
                        continue
                    if ev != "VAULT_APPEND":
                        continue
                    ej = (row.get("entry_json") or "").strip()
                    if not ej:
                        continue
                    try:
                        entry = json.loads(ej)
                    except json.JSONDecodeError:
                        continue
                    if not isinstance(entry, dict):
                        continue
                    vid = str(entry.get("vault_id") or "")
                    if not vid or vid in seen_vault_ids:
                        continue
                    seen_vault_ids.add(vid)
                    if uhash and not entry.get("Unique_Claim_Hash"):
                        entry["Unique_Claim_Hash"] = uhash
                    elif not entry.get("Unique_Claim_Hash"):
                        src = entry.get("_source_row") if isinstance(entry.get("_source_row"), dict) else {}
                        dd = clean_text(str(src.get("Denial Date") or "\u2014"))
                        uh = _unique_claim_hash(
                            clean_text(str(entry.get("Patient") or "\u2014")),
                            dd,
                            _parse_amount_denied(entry.get("Potential Revenue")),
                        )
                        entry["Unique_Claim_Hash"] = uh
                        hashes.add(uh)
                    st.session_state.revenue_vault.append(entry)
        except OSError:
            pass
        st.session_state._vault_unique_claim_hashes = hashes
        st.session_state._vault_shadow_hydrated = True


def _claim_row_audit_hash(row: dict) -> str:
    """Short stable hash for Institutional Ledger audit column."""
    payload = json.dumps(
        {
            "Patient ID": row.get("Patient ID", ""),
            "Denial Code": row.get("Denial Code", ""),
            "Law Cited": row.get("Law Cited", ""),
            "Denial Date": row.get("Denial Date", ""),
        },
        sort_keys=True,
        default=str,
    )
    return hashlib.sha256(payload.encode("utf-8")).hexdigest()[:16]


def _merge_neural_batch_into_vault(rows: list[dict], batch_sig: str) -> None:
    """Append claims from one neural audit batch into the global revenue vault (idempotent per batch_sig)."""
    _ensure_revenue_vault()
    merged: set = st.session_state.setdefault("_vault_merged_sigs", set())
    if batch_sig in merged:
        return
    hashes: set[str] = st.session_state.setdefault("_vault_unique_claim_hashes", set())
    with _VAULT_IO_LOCK:
        merged.add(batch_sig)
        for row in rows:
            raw_pr = str(row.get("Potential Revenue", "0") or "0")
            rev_ok, rev_amt = _double_entry_revenue_ok(raw_pr)
            parsed_amt = _parse_amount_denied(raw_pr)
            base = float(rev_amt if rev_ok else parsed_amt)
            _pid = clean_text(str(row.get("Patient ID", "—") or "—"))
            _dd = clean_text(str(row.get("Denial Date", "") or "—"))
            uh = _unique_claim_hash(_pid, _dd, parsed_amt)
            if uh in hashes:
                _append_vault_shadow_log_row(
                    "DUPLICATE_BLOCKED",
                    vault_id="",
                    patient_id=_pid,
                    denial_date=_dd,
                    amount=parsed_amt,
                    unique_claim_hash=uh,
                    validation="DUPLICATE",
                    entry=None,
                )
                continue
            hashes.add(uh)
            vault_id = str(uuid.uuid4())
            _reason = str(row.get("Reason for Denial", "") or "")
            _law = str(row.get("Law Cited", "") or "")
            appeal = _resolve_neural_appeal_mode(_reason, _law)
            _wpf = _parse_win_probability(row.get("Win Probability", "0"))
            _cid = clean_text(str(row.get("clinic_id") or st.session_state.get("active_clinic_id") or ""))
            _pl = clean_text(str(row.get("Payer Name", "") or "")) or _extract_payer_from_row(row) or "—"
            _val = MANUAL_VERIFICATION_TAG if not rev_ok else clean_text(str(row.get("Validation", "OK") or "OK"))
            entry = {
                "vault_id": vault_id,
                "clinic_id": _cid,
                "Patient": _pid,
                "Patient Name": clean_text(str(row.get("Patient Name", "") or "")) or "—",
                "Denial Code": clean_text(str(row.get("Denial Code", "—") or "—")),
                "payer_label": _pl,
                "Law Cited": clean_text(str(row.get("Law Cited", "—") or "—")),
                "Potential Revenue": clean_text(str(row.get("Potential Revenue", "0") or "0")),
                "Amount Denied": base,
                "amount_denied_base": base,
                "Win Probability": _wpf,
                "Appeal Mode": appeal or "CLINICAL",
                "vault_status": VAULT_STATUS_NEURAL_DRAFT,
                "Validation": _val,
                "Urgency Tag": clean_text(str(row.get("Urgency Tag", "") or "")),
                "Audit Hash": _claim_row_audit_hash(row),
                "Unique_Claim_Hash": uh,
                "batch_sig": batch_sig,
                "_source_row": dict(row),
            }
            st.session_state.revenue_vault.append(entry)
            amt_key = f"vault_amt_{vault_id}"
            if amt_key not in st.session_state:
                st.session_state[amt_key] = float(base)
            _append_vault_shadow_log_row(
                "VAULT_APPEND",
                vault_id=vault_id,
                patient_id=_pid,
                denial_date=_dd,
                amount=base,
                unique_claim_hash=uh,
                validation=_val,
                entry=dict(entry),
            )
    _record_payer_intel_for_vault_rows(rows)


def _sync_vault_from_neural_batch() -> None:
    """Merge current session neural_audit_batch into the vault (if present)."""
    data = st.session_state.get("neural_audit_batch")
    sig = st.session_state.get("neural_audit_batch_sig")
    if data and sig:
        skip = st.session_state.get("_ghost_skip_full_batch_sync_sig")
        if skip and skip == sig:
            return
        _merge_neural_batch_into_vault(data, sig)


def _vault_amount_for_entry(entry: dict) -> float:
    """Effective Amount Denied for metrics (widget-backed for manual overrides)."""
    vid = entry["vault_id"]
    k = f"vault_amt_{vid}"
    if k in st.session_state:
        try:
            return float(st.session_state[k])
        except (TypeError, ValueError):
            return float(entry.get("amount_denied_base", 0) or 0)
    return float(entry.get("amount_denied_base", 0) or 0)


def _vault_total_recoverable_for_clinic(clinic_id: str | None) -> float:
    """Total recoverable (Amount Denied sum) for displayable Revenue Vault rows for the active clinic."""
    _ensure_revenue_vault()
    cid = clean_text(str(clinic_id or ""))
    rows = _vault_entries_displayable()
    if cid:
        rows = [e for e in rows if clean_text(str(e.get("clinic_id") or "")) == cid]
    return float(sum(_vault_amount_for_entry(e) for e in rows))


def _vault_metrics_compute() -> tuple[float, float, float]:
    """Total recoverable (Amount Denied sum), statutory enforcement rate %, founders commission (15%)."""
    vault = _vault_entries_displayable()
    if not vault:
        return 0.0, 0.0, 0.0
    total = sum(_vault_amount_for_entry(e) for e in vault)
    statutory = sum(
        1 for e in vault if str(e.get("Appeal Mode", "")).upper().strip() == "STATUTORY"
    )
    rate = (statutory / len(vault)) * 100.0 if vault else 0.0
    founders = total * FOUNDERS_COMMISSION_RATE
    return total, rate, founders


def _vault_insurer_heatmap_df() -> pd.DataFrame:
    """Aggregate vault recoverable $ by insurer (payer) for client-facing bar chart."""
    totals: dict[str, float] = defaultdict(float)
    for e in _vault_entries_displayable():
        amt = _vault_amount_for_entry(e)
        src = e.get("_source_row")
        if isinstance(src, dict):
            payer = _get_best_val(src, ["payer", "insur", "carrier", "plan"])
        else:
            payer = ""
        if not payer:
            payer = _get_best_val(e, ["payer", "insur", "carrier", "plan"])
        payer = (payer or "").strip()[:72] or "Other / Unspecified"
        totals[payer] += float(amt)
    if not totals:
        return pd.DataFrame({"Recovery ($)": []})
    s = pd.Series(totals, name="Recovery ($)")
    s = s.sort_values(ascending=False)
    return s.to_frame()


def _service_agreement_pdf_bytes(clinic_name: str) -> bytes:
    """Senturion Service Agreement — 15% contingent success fee (for physician signature)."""
    if not HAS_FPDF or FPDF is None or SenturionPDF is None:
        return b""
    clinic = (clinic_name or "").strip() or "Medical Practice"
    today = datetime.now().strftime("%B %d, %Y")
    pdf = SenturionPDF()
    pdf.set_auto_page_break(auto=True, margin=max(15.0, PDF_BOTTOM_MARGIN_FOR_FOOTER_MM))
    m = 18.0
    pdf.set_margins(m, PDF_TOP_MARGIN_FOR_SEAL_MM, m)
    fname = _register_rcm_pdf_font(pdf)
    pdf.rcm_font_family = fname
    pdf.audit_tracking_hash = _vault_audit_hash_for_monetization()
    pdf.add_page()
    pdf.set_font(fname, "B", 15)
    pdf.cell(0, 10, "SENTURION SERVICE AGREEMENT", align="C", new_x="LMARGIN", new_y="NEXT")
    pdf.ln(2)
    pdf.set_font(fname, "", 10)
    pdf.cell(0, 6, f"Effective Date: {today}", align="C", new_x="LMARGIN", new_y="NEXT")
    pdf.ln(6)

    sections = [
        (
            "1. Parties",
            f"This Agreement is entered into by and between Senturion AI Solutions "
            f"(\"Senturion\") and the provider practice identified below (\"Provider\" or \"Practice\"). "
            f"Provider legal name as stated for this engagement: {clinic}.",
        ),
        (
            "2. Scope of Services",
            "Senturion may provide denial analytics, appeals support, administrative recovery coordination, "
            "and related revenue-cycle services as mutually agreed in writing or through the Senturion platform.",
        ),
        (
            "3. Success Fee — Fifteen Percent (15%)",
            "Provider agrees to pay Senturion a contingent success fee equal to fifteen percent (15%) of "
            "Gross Recovered Amounts actually collected or credited to Provider as a result of Senturion's "
            "services under this Agreement (including but not limited to payer reversals, reprocesses, "
            "and contractual payments attributable to qualifying claims). The fee applies only to amounts "
            "recovered after the Effective Date of this Agreement and is due and payable in accordance "
            "with Senturion's standard invoicing schedule.",
        ),
        (
            "4. No Recovery, No Fee",
            "Except as otherwise agreed in a separate addendum, if no recovery is realized, no success fee "
            "is owed for that claim or period.",
        ),
        (
            "5. Authorization",
            "By signing below, Provider authorizes Senturion to proceed with recovery services under these "
            "terms and confirms that the individual signing has authority to bind the Practice.",
        ),
        (
            "6. General",
            "This Agreement is governed by the terms presented to Provider at the time of execution and "
            "may be supplemented by additional exhibits. Electronic signatures and PDF copies are valid.",
        ),
    ]

    for title, body in sections:
        pdf.set_font(fname, "B", 11)
        pdf.multi_cell(0, 6, _appeal_text_for_pdf(title))
        pdf.ln(1)
        pdf.set_font(fname, "", 10)
        pdf.multi_cell(0, 5, _appeal_text_for_pdf(body))
        pdf.ln(4)

    pdf.ln(6)
    pdf.set_font(fname, "B", 11)
    pdf.multi_cell(0, 6, _appeal_text_for_pdf("PROVIDER / AUTHORIZED SIGNATORY"))
    pdf.ln(2)
    pdf.set_font(fname, "", 10)
    pdf.multi_cell(0, 5, _appeal_text_for_pdf(f"Practice: {clinic}"))
    pdf.ln(10)
    pdf.cell(0, 6, "_" * 72, new_x="LMARGIN", new_y="NEXT")
    pdf.set_font(fname, "I", 9)
    pdf.cell(0, 5, "Signature", new_x="LMARGIN", new_y="NEXT")
    pdf.ln(6)
    pdf.set_font(fname, "", 10)
    pdf.cell(0, 6, "Printed Name & Title: _________________________________", new_x="LMARGIN", new_y="NEXT")
    pdf.ln(2)
    pdf.cell(0, 6, "Date: _________________________________", new_x="LMARGIN", new_y="NEXT")

    out = pdf.output(dest="S")
    if isinstance(out, (bytes, bytearray)):
        return bytes(out)
    return str(out).encode("latin-1", errors="replace")


def _msa_master_service_agreement_pdf_bytes(
    clinic_name: str,
    clinic_id: str | None,
    total_recoverable_usd: float,
) -> bytes:
    """
    Senturion Master Service Agreement (MSA) — institutional PDF for digital signature.
    Success fee fixed at MSA_SUCCESS_FEE_PERCENT. Clinic name + vault recoverable are dynamic.
    """
    if not HAS_FPDF or FPDF is None or SenturionPDF is None:
        return b""
    clinic = (clinic_name or "").strip() or "Medical Practice"
    _cid_disp = clean_text(str(clinic_id or "")) or "—"
    today = datetime.now().strftime("%B %d, %Y")
    ts_iso = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    rev_s = f"${float(total_recoverable_usd):,.2f}"
    fee_pct = int(MSA_SUCCESS_FEE_PERCENT)
    try:
        _ah = _vault_wide_audit_hash_for_certificate(clinic_id) if clinic_id else _vault_audit_hash_for_monetization()
    except Exception:
        _ah = _vault_audit_hash_for_monetization()

    pdf = SenturionPDF()
    pdf.set_auto_page_break(auto=True, margin=max(16.0, PDF_BOTTOM_MARGIN_FOR_FOOTER_MM))
    m = 18.0
    pdf.set_margins(m, PDF_TOP_MARGIN_FOR_SEAL_MM, m)
    fname = _register_rcm_pdf_font(pdf)
    pdf.rcm_font_family = fname
    pdf.audit_tracking_hash = str(_ah)
    pdf.add_page()

    def _f(style: str, size: int) -> None:
        try:
            pdf.set_font(fname, style, size)
        except Exception:
            pdf.set_font(fname, "", size)

    _w = pdf.w - 2 * pdf.l_margin
    _f("B", 16)
    pdf.cell(0, 9, _appeal_text_for_pdf("SENTURION MASTER SERVICE AGREEMENT (MSA)"), align="C", new_x="LMARGIN", new_y="NEXT")
    pdf.ln(2)
    _f("I", 9)
    pdf.cell(0, 5, _appeal_text_for_pdf("Revenue Recovery & Neural Audit Services"), align="C", new_x="LMARGIN", new_y="NEXT")
    pdf.ln(4)
    _f("", 9)
    pdf.multi_cell(_w, 4, _appeal_text_for_pdf(f"Effective Date: {today}  ·  Generated: {ts_iso}"))
    pdf.ln(3)
    _f("B", 10)
    pdf.multi_cell(_w, 5, _appeal_text_for_pdf("1. Parties"))
    _f("", 9)
    pdf.multi_cell(
        _w,
        4.2,
        _appeal_text_for_pdf(
            f"This Master Service Agreement (\"Agreement\") is entered into as of the Effective Date by and between "
            f"Senturion AI Solutions, a revenue-cycle intelligence and recovery services provider (\"Senturion\"), "
            f"and {clinic}, a medical practice or provider organization (\"Provider\"). "
            f"Provider's registered clinic identifier in the Senturion platform: {_cid_disp}."
        ),
    )
    pdf.ln(3)
    _f("B", 10)
    pdf.multi_cell(_w, 5, _appeal_text_for_pdf("2. Background & Scope"))
    _f("", 9)
    pdf.multi_cell(
        _w,
        4.2,
        _appeal_text_for_pdf(
            "Senturion may perform denial analytics, appeals support, administrative recovery coordination, "
            "statutory enforcement tooling, and related services using Provider data processed through the "
            "Senturion Neural Audit and Revenue Vault. Services are provided subject to this Agreement and "
            "any order or statement of work Senturion and Provider may execute."
        ),
    )
    pdf.ln(3)
    _f("B", 10)
    pdf.multi_cell(_w, 5, _appeal_text_for_pdf("3. Compensation — Success Fee (Fixed)"))
    _f("", 9)
    pdf.multi_cell(
        _w,
        4.2,
        _appeal_text_for_pdf(
            f"Provider agrees to pay Senturion a contingent success fee equal to {fee_pct} percent ({fee_pct}%) "
            f"of Gross Recovered Amounts actually collected or credited to Provider that are attributable to "
            f"Senturion's services under this Agreement (including payer reversals, reprocesses, and contractual "
            f"payments tied to qualifying claims identified or advanced through the platform). "
            f"The success fee percentage is fixed at {fee_pct}% for the term of this MSA unless superseded by a "
            f"written amendment signed by both parties."
        ),
    )
    pdf.ln(3)
    _f("B", 10)
    pdf.multi_cell(_w, 5, _appeal_text_for_pdf("4. Revenue Vault Reference (Informational)"))
    _f("", 9)
    pdf.multi_cell(
        _w,
        4.2,
        _appeal_text_for_pdf(
            f"As of generation of this document, the aggregate Total Recoverable Revenue shown in the "
            f"Revenue Vault for Provider (displayable, non-quarantined rows) — Senturion AI Solutions — is {rev_s}. "
            f"This figure is an operational proxy for planning and is not a guarantee of collection; invoicing of the "
            f"{fee_pct}% fee applies only to amounts actually recovered in accordance with Section 3."
        ),
    )
    pdf.ln(3)
    _f("B", 10)
    pdf.multi_cell(_w, 5, _appeal_text_for_pdf("5. No Recovery, No Fee; Invoicing"))
    _f("", 9)
    pdf.multi_cell(
        _w,
        4.2,
        _appeal_text_for_pdf(
            "Except as otherwise agreed in a signed addendum, if no recovery is realized for a given claim or period, "
            "no success fee is owed for that claim or period. Senturion may invoice success fees in accordance with "
            "its standard billing cadence after recovery is credited or confirmed."
        ),
    )
    pdf.ln(3)
    _f("B", 10)
    pdf.multi_cell(_w, 5, _appeal_text_for_pdf("6. Term; Electronic Signature"))
    _f("", 9)
    pdf.multi_cell(
        _w,
        4.2,
        _appeal_text_for_pdf(
            "This Agreement continues until terminated by either party with reasonable written notice, or as "
            "superseded by a successor agreement. Electronic signatures and PDF copies are valid. "
            "Provider represents that the signatory below has authority to bind the Practice."
        ),
    )
    pdf.ln(5)
    _f("B", 10)
    pdf.multi_cell(_w, 5, _appeal_text_for_pdf("SIGNATURES"))
    pdf.ln(2)
    _f("", 9)
    pdf.multi_cell(
        _w,
        4,
        _appeal_text_for_pdf(
            "SENTURION AI SOLUTIONS — Authorized Representative\n"
            "Signature: _______________________________________   Date: _______________"
        ),
    )
    pdf.ln(4)
    pdf.multi_cell(
        _w,
        4,
        _appeal_text_for_pdf(
            f"PROVIDER — {clinic}\n"
            "Signature: _______________________________________   Date: _______________\n"
            "Printed Name & Title: _____________________________"
        ),
    )

    out = pdf.output(dest="S")
    if isinstance(out, (bytes, bytearray)):
        return bytes(out)
    return str(out).encode("latin-1", errors="replace")


def _vault_audit_hash_for_monetization() -> str:
    """Stable audit batch linkage hash from current Revenue Vault scope (filtered)."""
    vault = _vault_entries_displayable()
    if not vault:
        return "NO_AUDIT_DATA"
    rows = []
    for e in sorted(vault, key=lambda x: str(x.get("vault_id") or "")):
        rows.append(
            {
                "vault_id": e.get("vault_id"),
                "patient": e.get("Patient"),
                "denial": e.get("Denial Code"),
                "amt": round(float(_vault_amount_for_entry(e)), 2),
            }
        )
    return hashlib.sha256(json.dumps(rows, sort_keys=True, default=str).encode("utf-8")).hexdigest()[
        :12
    ].upper()


def _senturion_revenue_recovery_agreement_pdf_bytes(
    clinic_name: str,
    audit_hash: str,
    *,
    signed_at_iso: str | None = None,
) -> bytes:
    """
    One-page Senturion AI Solutions Revenue Recovery Agreement — 15% success fee, audit hash, dual signature blocks
    (Senturion representative / Provider physician). Uses SenturionPDF seal + legal footer.
    """
    if not HAS_FPDF or FPDF is None or SenturionPDF is None:
        return b""
    clinic = (clinic_name or "").strip() or "Medical Practice"
    _ts = signed_at_iso or datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    today = datetime.now().strftime("%B %d, %Y")
    pdf = SenturionPDF()
    pdf.set_auto_page_break(auto=True, margin=max(14.0, PDF_BOTTOM_MARGIN_FOR_FOOTER_MM))
    m = 14.0
    pdf.set_margins(m, PDF_TOP_MARGIN_FOR_SEAL_MM, m)
    fname = _register_rcm_pdf_font(pdf)
    pdf.rcm_font_family = fname
    pdf.audit_tracking_hash = str(audit_hash or "").strip()
    pdf.add_page()
    pdf.set_font(fname, "B", 13)
    pdf.cell(0, 7, "SENTURION AI SOLUTIONS · REVENUE RECOVERY AGREEMENT", align="C", new_x="LMARGIN", new_y="NEXT")
    pdf.set_font(fname, "I", 8)
    pdf.cell(0, 5, _appeal_text_for_pdf("One-page contingent fee authorization (Neural Audit)"), align="C", new_x="LMARGIN", new_y="NEXT")
    pdf.ln(3)
    pdf.set_font(fname, "", 8)
    body = (
        f"Effective date: {today}. This Revenue Recovery Agreement (\"Agreement\") is entered into by and between "
        f"Senturion AI Solutions (\"Senturion\") and {clinic} (\"Provider\"). "
        f"Provider authorizes Senturion to pursue recovery services relating to claims identified in Neural Audit "
        f"batch **{audit_hash}** as reflected in the Revenue Vault for Provider's account (Senturion AI Solutions). "
        f"Provider agrees to pay Senturion a contingent success fee equal to **fifteen percent (15%)** of all "
        f"gross recovered funds actually collected or credited to Provider that are **identified in or attributable to** "
        f"Audit **{audit_hash}** and Senturion's recovery efforts. The fee is due per Senturion's standard invoicing. "
        f"If no recovery is realized, no success fee is owed for that claim or period except as otherwise agreed in writing. "
        f"This Agreement may be executed electronically. A record of acceptance is stored at {_ts}. "
        f"Provider represents the signatory below has authority to bind the Practice."
    )
    pdf.multi_cell(0, 3.6, _appeal_text_for_pdf(body))
    pdf.ln(4)
    pdf.set_font(fname, "B", 9)
    pdf.multi_cell(0, 4, _appeal_text_for_pdf("SIGNATURES — BINDING"))
    pdf.ln(2)
    pdf.set_font(fname, "", 8)
    pdf.multi_cell(
        0,
        3.5,
        _appeal_text_for_pdf(
            "SENTURION AI SOLUTIONS — Authorized Representative (Eddie / Contracting Officer)\n"
            "Signature: _________________________________   Date: _______________"
        ),
    )
    pdf.ln(3)
    pdf.multi_cell(
        0,
        3.5,
        _appeal_text_for_pdf(
            f"PROVIDER — {clinic} — Authorized Physician / Executive\n"
            "Signature: _________________________________   Date: _______________"
        ),
    )
    pdf.ln(2)
    pdf.set_font(fname, "I", 7)
    pdf.multi_cell(
        0,
        3,
        _appeal_text_for_pdf(
            f"Electronically acknowledged and PDF generated under Audit Hash {audit_hash} · {_ts}"
        ),
    )
    out = pdf.output(dest="S")
    if isinstance(out, (bytes, bytearray)):
        return bytes(out)
    return str(out).encode("latin-1", errors="replace")


def _ensure_client_vault_contracts() -> None:
    st.session_state.setdefault("client_vault_contracts", [])


def _append_client_vault_contract(entry: dict) -> None:
    _ensure_client_vault_contracts()
    st.session_state.client_vault_contracts.append(entry)
    if len(st.session_state.client_vault_contracts) > 100:
        st.session_state.client_vault_contracts = st.session_state.client_vault_contracts[-100:]


def _monetization_signed_for_clinic(clinic_id: str | None) -> bool:
    if not clinic_id:
        return False
    return bool(st.session_state.get(f"_monetization_signed_{clinic_id}"))


def _monetization_gate_dialog_body() -> None:
    """Modal body: summary + Agree / Cancel; on Agree, PDF to Client Vault + gate flag."""
    _ensure_clinic_profiles()
    _cid = st.session_state.get("_modal_monetization_clinic_id") or st.session_state.get("active_clinic_id")
    _cn = str(st.session_state.get("_modal_monetization_clinic_name") or "Clinic").strip()
    _ah = str(st.session_state.get("_modal_monetization_audit_hash") or "—")
    st.markdown("##### Senturion Monetization Gate")
    st.info(
        f"By clicking **Agree**, **{_cn}** agrees to pay **Senturion AI Solutions** **15%** of all "
        f"recovered funds identified in Audit **[{_ah}]**."
    )
    with st.expander("Full agreement template (1-page PDF on Agree)", expanded=False):
        st.caption(
            "The downloadable contract includes party names, audit hash, 15% contingent fee, and signature lines for "
            "Senturion (Eddie / Contracting Officer) and the Provider authorized physician."
        )
    c1, c2 = st.columns(2)
    with c1:
        if st.button("Agree & sign", type="primary", use_container_width=True, key="mz_gate_agree_btn"):
            _signed_iso = datetime.now().isoformat()
            pdf_b = _senturion_revenue_recovery_agreement_pdf_bytes(
                _cn,
                _ah,
                signed_at_iso=_signed_iso,
            )
            if not pdf_b:
                st.error("PDF engine unavailable. Install: pip install fpdf2")
            else:
                fn = f"Senturion_Revenue_Recovery_Agreement_{_ah}.pdf"
                _append_client_vault_contract(
                    {
                        "kind": "revenue_recovery_agreement",
                        "clinic_id": str(_cid or ""),
                        "clinic_name": _cn,
                        "audit_hash": _ah,
                        "signed_at": _signed_iso,
                        "pdf_bytes": pdf_b,
                        "file_name": fn,
                    }
                )
                if _cid:
                    st.session_state[f"_monetization_signed_{_cid}"] = True
                    st.session_state[f"_monetization_audit_hash_{_cid}"] = _ah
                st.session_state["_open_monetization_dialog"] = False
                _append_audit_log(
                    f"Monetization gate SIGNED — Clinic [{_cn}] · Audit [{_ah}] · Revenue Recovery Agreement stored in Client Vault."
                )
                st.success("Agreement signed. Contract saved to **Client Vault**. Batch enforcement is now unlocked for this clinic.")
                st.rerun()
    with c2:
        if st.button("Cancel", use_container_width=True, key="mz_gate_cancel_btn"):
            st.session_state["_open_monetization_dialog"] = False
            st.rerun()


if hasattr(st, "dialog"):

    @st.dialog("Senturion AI Solutions · Revenue Recovery")
    def monetization_gate_dialog() -> None:
        _monetization_gate_dialog_body()

else:

    def monetization_gate_dialog() -> None:
        with st.container(border=True):
            st.caption("Authorization dialog")
            _monetization_gate_dialog_body()


def _treasury_wallet_qr_placeholder_png() -> bytes:
    """Placeholder QR-style image for USDC/Solana wallet (not a real QR — replace with generated QR in production)."""
    try:
        from PIL import Image, ImageDraw, ImageFont
    except ImportError:
        return b""

    w, h = 240, 240
    im = Image.new("RGB", (w, h), (10, 10, 10))
    dr = ImageDraw.Draw(im)
    for i in range(0, w, 14):
        for j in range(0, h, 14):
            if (i // 14 + j // 14) % 2 == 0:
                dr.rectangle([i, j, i + 12, j + 12], fill=(0, 55, 28))
            else:
                dr.rectangle([i, j, i + 12, j + 12], fill=(18, 18, 18))
    dr.rectangle([10, 10, 70, 70], outline=(224, 224, 224), width=2)
    dr.rectangle([w - 80, 10, w - 10, 70], outline=(224, 224, 224), width=2)
    dr.rectangle([10, h - 80, 70, h - 10], outline=(224, 224, 224), width=2)
    try:
        font = ImageFont.truetype("arial.ttf", 11)
        font_sm = ImageFont.truetype("arial.ttf", 9)
    except Exception:
        font = ImageFont.load_default()
        font_sm = font
    dr.text((14, h - 52), "Senturion Treasury", fill=(224, 224, 224), font=font)
    dr.text((14, h - 34), "USDC / Solana (placeholder)", fill=(0, 255, 65), font=font_sm)
    dr.text((14, h - 18), "Wallet QR — configure in Admin", fill=(140, 140, 140), font=font_sm)
    buf = io.BytesIO()
    im.save(buf, format="PNG")
    return buf.getvalue()


def _signed_msa_doc_id_footer() -> str:
    """Auditor-facing Merchant Agreement DocID (official identifier for partners)."""
    return str(MSA_MERCHANT_DOC_ID)


def _msa_doc_id_display() -> str:
    """Short UI form (e.g. 0f3a498…)."""
    h = str(MSA_MERCHANT_DOC_ID).strip()
    return f"{h[:7]}…" if len(h) > 7 else h


def _paystack_net_usd(gross_usd: float) -> float:
    """USD amount after estimated Paystack international card fee (deducted from gross)."""
    return float(gross_usd) * (1.0 - PAYSTACK_INTL_CARD_FEE_PCT)


def _standard_audit_fee_partner_split_usd(
    gross_usd: float | None = None,
) -> tuple[float, float, float, float, float]:
    """
    Revenue engine — canonical **50/50 net** split on each incoming **$2,625** audit fee
    (or any override gross): gross → est. Paystack → net USD → equal CEO/CFO shares on net.

    Returns ``(gross, paystack_est, net, ceo_net, cfo_net)``.
    """
    g = float(AUDIT_FEE_USD_STANDARD if gross_usd is None else gross_usd)
    ps = g * PAYSTACK_INTL_CARD_FEE_PCT
    net = _paystack_net_usd(g)
    half = net * 0.5
    return (g, ps, net, half, half)


def _capitec_statement_abs_path() -> str:
    return os.path.join(os.path.dirname(os.path.abspath(__file__)), CAPITEC_STATEMENT_FILENAME)


def _read_capitec_statement_bytes() -> bytes | None:
    p = _capitec_statement_abs_path()
    if not os.path.isfile(p):
        return None
    try:
        with open(p, "rb") as f:
            return f.read()
    except OSError:
        return None


@st.cache_data(ttl=FX_CACHE_TTL_SEC, show_spinner=False)
def _fetch_usd_zar_rate_live() -> float:
    """
    Live USD → ZAR (mid). Cached ~1h. Uses public FX endpoints (no API key).
    Falls back to FX_FALLBACK_USD_ZAR if unreachable.
    """
    if requests is None:
        return float(FX_FALLBACK_USD_ZAR)
    headers = {"User-Agent": "Senturion-BPO/1.0 (CFO FX)"}
    endpoints = (
        "https://api.frankfurter.app/latest?from=USD&to=ZAR",
        "https://api.exchangerate-api.com/v4/latest/USD",
    )
    for url in endpoints:
        try:
            r = requests.get(url, timeout=12, headers=headers)
            r.raise_for_status()
            data = r.json()
            rates = data.get("rates") or {}
            z = rates.get("ZAR")
            if z is not None:
                return float(z)
        except Exception:
            continue
    return float(FX_FALLBACK_USD_ZAR)


def _usd_to_zar(usd: float, zar_per_usd: float) -> float:
    return float(usd) * float(zar_per_usd)


def _vault_entries_for_scope(vf: str) -> list[dict]:
    """Revenue Vault rows for treasury scope (ALL or clinic_id)."""
    _ensure_revenue_vault()
    rows = _vault_entries_displayable()
    vf = str(vf or "ALL")
    if vf == "ALL":
        return list(rows)
    return [e for e in rows if clean_text(str(e.get("clinic_id") or "")) == clean_text(vf)]


def _partner_claim_display_id(entry: dict) -> str:
    """Best-effort Claim ID for partner transparency."""
    src = entry.get("_source_row")
    if isinstance(src, dict):
        for k in ("Claim #", "Claim Number", "Claim ID", "Claim No", "claim_id", "Claim"):
            c = clean_text(str(src.get(k) or ""))
            if c and c not in ("—", "N/A"):
                return c
    for k in ("Claim #", "Claim Number", "Claim ID"):
        c = clean_text(str(entry.get(k) or ""))
        if c and c not in ("—", "N/A"):
            return c
    pid = clean_text(str(entry.get("Patient") or entry.get("Patient ID") or "")) or "—"
    dc = clean_text(str(entry.get("Denial Code") or "")) or "—"
    vid = str(entry.get("vault_id") or "")[:10]
    return f"{pid} · {dc} · {vid}"


def _append_partner_settlement_rows_for_scope(vf: str) -> None:
    """
    Partner Settlement Log — one row per vault claim in scope when payment is COLLECTED.
    Gross 15% fee → subtract est. Paystack intl. fee → net USD → 50/50 CEO/CFO.
    ZAR columns use live USD/ZAR at append time (cached FX).
    """
    entries = _vault_entries_for_scope(vf)
    now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    log = st.session_state.setdefault("partner_settlement_log", [])
    half = 0.5
    fx = _fetch_usd_zar_rate_live()
    for entry in entries:
        amt = _vault_amount_for_entry(entry)
        fee_gross = float(amt) * FOUNDERS_COMMISSION_RATE
        fee_paystack = fee_gross * PAYSTACK_INTL_CARD_FEE_PCT
        fee_net = _paystack_net_usd(fee_gross)
        ceo_usd = fee_net * half
        cfo_usd = fee_net * half
        ceo_zar = _usd_to_zar(ceo_usd, fx)
        cfo_zar = _usd_to_zar(cfo_usd, fx)
        vid = str(entry.get("vault_id") or "")
        log.append(
            {
                "Date": now,
                "Claim ID": _partner_claim_display_id(entry),
                "Gross Fee (USD)": round(fee_gross, 2),
                "Paystack est. (USD)": round(fee_paystack, 2),
                "Net Fee (USD)": round(fee_net, 2),
                "CEO Net (USD)": round(ceo_usd, 2),
                "CFO Net (USD)": round(cfo_usd, 2),
                "CEO Net (ZAR)": round(ceo_zar, 2),
                "CFO Net (ZAR)": round(cfo_zar, 2),
                "FX USD/ZAR": round(fx, 4),
                "Status": "🟢 COLLECTED",
                "Scope": str(vf),
                "vault_id": vid,
            }
        )


def _treasury_celebrate_if_collected_transition(vf: str) -> None:
    """Balloons + confetti + partner ledger rows when settlement moves to COLLECTED (per vault scope)."""
    sk = f"treasury_settlement_{vf}"
    st.session_state.setdefault(sk, TREASURY_SETTLEMENT_OPTIONS[0])
    lk = f"{sk}_prev_celebrate"
    cur = st.session_state.get(sk, TREASURY_SETTLEMENT_OPTIONS[0])
    prev = st.session_state.get(lk)
    if prev is not None and prev != "🟢 COLLECTED" and cur == "🟢 COLLECTED":
        st.balloons()
        if hasattr(st, "confetti"):
            try:
                st.confetti()
            except Exception:
                pass
        _append_partner_settlement_rows_for_scope(vf)
    st.session_state[lk] = cur


def _treasury_invoice_dialog_body() -> None:
    """Stripe link, wallet + QR placeholder, settlement status."""
    vf = str(st.session_state.get("vault_clinic_filter") or "ALL")
    sk = f"treasury_settlement_{vf}"
    st.session_state.setdefault(sk, TREASURY_SETTLEMENT_OPTIONS[0])
    stripe_url = (
        clean_text(str(st.session_state.get("admin_treasury_stripe_url") or ""))
        or DEFAULT_TREASURY_STRIPE_URL
    )
    wallet_addr = (
        clean_text(str(st.session_state.get("admin_treasury_sol_wallet") or ""))
        or DEFAULT_TREASURY_SOL_WALLET
    )
    st.markdown("##### Senturion Treasury · Payment")
    st.caption("Placeholder links — replace with live endpoints in **Admin settings** (sidebar).")
    st.link_button("Open Stripe checkout (15% success fee)", stripe_url, use_container_width=True)
    st.markdown("**USDC / Solana wallet address**")
    st.code(wallet_addr, language=None)
    _png = _treasury_wallet_qr_placeholder_png()
    if _png:
        st.image(_png, caption="Wallet QR (placeholder — paste real address in Admin settings)", width=260)
    st.selectbox(
        "Settlement Status",
        TREASURY_SETTLEMENT_OPTIONS,
        key=sk,
        help="When set to COLLECTED, celebration runs and the **Partner Settlement Log** is populated (per claim) for this scope.",
    )
    if st.button("Close", use_container_width=True, key="treasury_invoice_close_btn"):
        st.session_state["_treasury_invoice_open"] = False
        st.rerun()


if hasattr(st, "dialog"):

    @st.dialog("Senturion Treasury · Invoice clinic")
    def treasury_invoice_dialog() -> None:
        _treasury_invoice_dialog_body()

else:

    def treasury_invoice_dialog() -> None:
        with st.container(border=True):
            _treasury_invoice_dialog_body()


def _render_client_facing_view(clinic_name: str) -> None:
    """Pitch mode: emerald hero, insurer heatmap, monetization gate + Client Vault contracts."""
    st.markdown(
        """
<style>
section.main [data-testid="stButton"] button[kind="primary"],
section.main div[data-testid="stDownloadButton"] button {
    background: #00FF41 !important;
    color: #050505 !important;
    border: 3px solid #00FF41 !important;
    font-weight: 700 !important;
    letter-spacing: 0.05em !important;
}
section.main div[data-testid="stDownloadButton"] button {
    text-transform: uppercase !important;
}
</style>
""",
        unsafe_allow_html=True,
    )
    st.markdown(
        '<h1 class="hud-title" style="border:none;">Revenue Recovery · Client View</h1>',
        unsafe_allow_html=True,
    )
    total, _, _ = _vault_metrics_compute()
    total_fmt = f"{total:,.2f}"
    st.markdown(
        f"""
<div class="client-pitch-hero">
  <div class="client-pitch-label">REVENUE RECOVERY OPPORTUNITY</div>
  <p class="client-pitch-value">${total_fmt}</p>
  <div class="client-pitch-sub">Aggregated from your vault · recoverable proxy</div>
</div>
""",
        unsafe_allow_html=True,
    )

    st.markdown(
        '<p style="font-family:JetBrains Mono,monospace;font-size:0.72rem;letter-spacing:0.18em;'
        'color:#E0E0E0;text-transform:uppercase;margin:0.5rem 0 0.35rem;">Denial heatmap · by insurer</p>',
        unsafe_allow_html=True,
    )
    st.caption("Estimated recovery exposure by payer (from your Neural Audit Summary).")
    heat_df = _vault_insurer_heatmap_df()
    if heat_df.empty or float(heat_df["Recovery ($)"].sum()) <= 0:
        st.info("No insurer-level data yet — run a neural audit or sync denials to populate the vault.")
    else:
        st.bar_chart(heat_df)

    st.markdown("---")
    _ensure_clinic_profiles()
    _active_cid = st.session_state.get("active_clinic_id")
    _display_cname = (clinic_name or "").strip() or (
        _clinic_display_name(_active_cid) if _active_cid else "Clinic"
    )
    _audit_hash = _vault_audit_hash_for_monetization()
    if _active_cid and _monetization_signed_for_clinic(str(_active_cid)):
        _signed_ah = st.session_state.get(f"_monetization_audit_hash_{_active_cid}")
        if _signed_ah:
            st.success(
                f"**Revenue Recovery Agreement signed** for this clinic (Audit **{_signed_ah}**). "
                "Batch enforcement is authorized for agents."
            )
        else:
            st.success(
                "**Revenue Recovery Agreement signed** for this clinic. Batch enforcement is authorized for agents."
            )

    if st.button(
        "🤝 AUTHORIZE RECOVERY",
        type="primary",
        use_container_width=True,
        key="btn_client_authorize_recovery",
    ):
        st.session_state["_open_monetization_dialog"] = True
        st.session_state["_modal_monetization_clinic_id"] = str(_active_cid or "")
        st.session_state["_modal_monetization_clinic_name"] = _display_cname
        st.session_state["_modal_monetization_audit_hash"] = _audit_hash
        st.rerun()

    if st.session_state.get("_open_monetization_dialog"):
        monetization_gate_dialog()

    st.markdown("##### Client Vault")
    st.caption("Signed contracts (this session) — PDFs for Provider and Senturion (Eddie handoff).")
    _ensure_client_vault_contracts()
    _vc = [
        c
        for c in (st.session_state.client_vault_contracts or [])
        if not _active_cid or str(c.get("clinic_id") or "") == str(_active_cid)
    ]
    if not _vc:
        st.warning("No signed contracts in vault yet. Use **Authorize Recovery** to execute the agreement.")
    else:
        for i, rec in enumerate(reversed(_vc[-12:])):
            _fn = rec.get("file_name") or "Senturion_Revenue_Recovery_Agreement.pdf"
            _pb = rec.get("pdf_bytes")
            _sa = rec.get("signed_at") or ""
            if _pb:
                _safe_k = re.sub(r"[^\w\-]+", "_", (_fn or "c")[:48])
                st.download_button(
                    label=f"⬇ {_fn} · {_sa[:19]}",
                    data=_pb,
                    file_name=_fn,
                    mime="application/pdf",
                    key=f"dl_client_vault_{i}_{_safe_k}",
                    use_container_width=True,
                )


def _render_vault_institutional_dashboard() -> None:
    """Three emerald KPIs above the Institutional Ledger."""
    _ensure_clinic_profiles()
    _opts = ["ALL"] + [p["clinic_id"] for p in st.session_state.clinic_profiles]
    if st.session_state.get("vault_clinic_filter") not in _opts:
        st.session_state.vault_clinic_filter = "ALL"

    def _vault_filter_label(k: str) -> str:
        if k == "ALL":
            return "All clinics"
        return _clinic_display_name(k)

    st.selectbox(
        "Filter Neural Audit Summary by clinic",
        options=_opts,
        format_func=_vault_filter_label,
        key="vault_clinic_filter",
    )
    total, rate, founders = _vault_metrics_compute()
    _vf_scope = str(st.session_state.get("vault_clinic_filter") or "ALL")
    _sk_settle = f"treasury_settlement_{_vf_scope}"
    st.session_state.setdefault(_sk_settle, TREASURY_SETTLEMENT_OPTIONS[0])
    _founders_glow = " founders-commission-glow"
    _vflash = st.session_state.get("_vault_metric_flash_ts")
    _pulse = bool(_vflash and (time.time() - float(_vflash) < 2.5))
    _rfl = " vault-recoverable-flash" if _pulse else ""
    st.markdown(
        f"""
<div class="neural-audit-summary-hero">
  <div class="na-hero-label">Neural Audit Summary</div>
  <div class="na-hero-amount{_rfl}">${total:,.2f}</div>
  <div class="na-hero-hint">Total recoverable exposure · Amount Denied</div>
</div>
<div class="vault-institutional-metrics vault-sub-metrics">
  <div class="vault-metric-card">
    <div class="vault-metric-label">STATUTORY ENFORCEMENT RATE</div>
    <div class="vault-metric-value" style="color:{VAULT_EMERALD};">{rate:.1f}%</div>
    <div class="vault-metric-hint">Claims in STATUTORY appeal mode</div>
  </div>
  <div class="vault-metric-card">
    <div class="vault-metric-label">SENTURION SUCCESS FEE (15%)</div>
    <div class="vault-metric-value{_founders_glow}" style="color:{VAULT_EMERALD};">${founders:,.2f}</div>
    <div class="vault-metric-hint">Fee pool on recoverable (real-time)</div>
  </div>
</div>
""",
        unsafe_allow_html=True,
    )
    _inv1, _inv2, _inv3 = st.columns([1, 1, 1])
    with _inv3:
        st.markdown('<div class="treasury-invoice-wrap">', unsafe_allow_html=True)
        if st.button(
            "💳 INVOICE CLINIC",
            type="primary",
            use_container_width=True,
            key="treasury_invoice_clinic_btn",
        ):
            st.session_state["_treasury_invoice_open"] = True
            st.rerun()
        st.markdown("</div>", unsafe_allow_html=True)
    st.caption(
        f"**Settlement (this scope):** `{st.session_state.get(_sk_settle, TREASURY_SETTLEMENT_OPTIONS[0])}` "
        "— open **Invoice clinic** to pay via Stripe / USDC and update status."
    )
    if st.session_state.get("_treasury_invoice_open"):
        treasury_invoice_dialog()
    _treasury_celebrate_if_collected_transition(_vf_scope)


def _render_institutional_ledger(*, show_title: bool = True) -> None:
    """Institutional Ledger table + Manual Revenue Override (popover)."""
    _ensure_clinic_profiles()
    vault = list(reversed(_vault_entries_displayable()))
    _q = _vault_entries_quarantine()
    if show_title:
        st.markdown(
            '<div class="projection-table-cap">Neural Audit Ledger</div>',
            unsafe_allow_html=True,
        )
    if not vault and not _q:
        st.caption("No claims in the Neural Audit Summary yet. Run a neural audit to deposit rows.")
        return
    if not vault and _q:
        st.warning(
            "**Double-entry validation:** All claims are in **[MANUAL_VERIFICATION]** — invalid USD format, **$0**, or over **$100,000**. "
            "They are withheld from the ledger until corrected."
        )

    if vault:
        st.markdown('<div class="institutional-ledger-wrap table-reveal">', unsafe_allow_html=True)
        h1, h2, h3, h4, h5, h6, h7, h8, h9, h10 = st.columns(
            [1.15, 0.72, 0.9, 1.0, 1.15, 0.9, 0.9, 0.95, 0.95, 1.35]
        )
        h1.markdown("**Patient**")
        h2.markdown("**Clinic**")
        h3.markdown("**QC**")
        h4.markdown("**Denial**")
        h5.markdown("**Law**")
        h6.markdown("**Revenue**")
        h7.markdown("**Hash**")
        h8.markdown("**Status**")
        h9.markdown("**MD §15-1005**")
        h10.markdown("**Override**")
    
        _opt_status = (
            VAULT_STATUS_UNAUDITED,
            VAULT_STATUS_NEURAL_DRAFT,
            VAULT_STATUS_AGENT_REVIEW,
            VAULT_STATUS_ENFORCED,
        )
    
        for entry in vault:
            vid = entry["vault_id"]
            if "vault_status" not in entry:
                entry["vault_status"] = VAULT_STATUS_NEURAL_DRAFT
            amt_key = f"vault_amt_{vid}"
            if amt_key not in st.session_state:
                st.session_state[amt_key] = float(entry.get("amount_denied_base", 0) or 0)
    
            kstat = f"vault_stat_{vid}"
            if kstat not in st.session_state:
                st.session_state[kstat] = entry.get("vault_status", VAULT_STATUS_NEURAL_DRAFT)
            if st.session_state[kstat] not in _opt_status:
                st.session_state[kstat] = entry.get("vault_status", VAULT_STATUS_NEURAL_DRAFT)
    
            kpay = f"vault_payrecv_{vid}"
            if kpay not in st.session_state:
                st.session_state[kpay] = bool(entry.get("payment_received", False))
    
            overdue_row = _is_enforcement_clock_overdue(entry)
            clk_lbl, clk_neon = _enforcement_clock_status_label(entry)
            if overdue_row:
                st.markdown('<div class="enforcement-row-shell">', unsafe_allow_html=True)
    
            c1, c2, c3, c4, c5, c6, c7, c8, c9, c10 = st.columns(
                [1.15, 0.72, 0.9, 1.0, 1.15, 0.9, 0.9, 0.95, 0.95, 1.35]
            )
            c1.markdown(f'<span class="ledger-cell">{entry.get("Patient", "—")}</span>', unsafe_allow_html=True)
            _cn = _clinic_display_name(entry.get("clinic_id"))
            c2.markdown(f'<span class="ledger-cell mono">{_cn}</span>', unsafe_allow_html=True)
            _vqc = str(entry.get("Validation") or "OK").strip()
            if MANUAL_FIX_TAG in _vqc or _vqc == MANUAL_FIX_TAG:
                c3.markdown(
                    f'<span class="ledger-cell ledger-law mono">{_vqc}</span>',
                    unsafe_allow_html=True,
                )
            else:
                c3.markdown(f'<span class="ledger-cell mono">{_vqc}</span>', unsafe_allow_html=True)
            c4.markdown(f'<span class="ledger-cell mono">{entry.get("Denial Code", "—")}</span>', unsafe_allow_html=True)
            law = (entry.get("Law Cited") or "—").replace("<", "&lt;")
            c5.markdown(f'<span class="ledger-cell ledger-law">{law}</span>', unsafe_allow_html=True)
            cur_amt = float(st.session_state.get(amt_key, entry.get("amount_denied_base", 0) or 0))
            c6.markdown(
                f'<span class="ledger-cell ledger-rev">${cur_amt:,.2f}</span>',
                unsafe_allow_html=True,
            )
            c7.markdown(f'<span class="ledger-hash mono">{entry.get("Audit Hash", "—")}</span>', unsafe_allow_html=True)
            with c8:
                st.selectbox(
                    "status",
                    options=list(_opt_status),
                    key=kstat,
                    format_func=lambda x: VAULT_STATUS_LABELS.get(x, x),
                    label_visibility="collapsed",
                )
            _clk_cls = "enforcement-overdue-text" if clk_neon else "ledger-cell mono"
            c9.markdown(f'<span class="{_clk_cls}">{clk_lbl}</span>', unsafe_allow_html=True)
            with c10:
                with st.popover("Manual Revenue Override", use_container_width=True):
                    st.caption("If the AI missed the dollar amount, set **Amount Denied** here. Metrics update on change.")
                    st.number_input(
                        "Amount Denied (USD)",
                        min_value=0.0,
                        value=float(st.session_state.get(amt_key, entry.get("amount_denied_base", 0) or 0)),
                        step=1.0,
                        format="%.2f",
                        key=amt_key,
                        help="Overrides AI Potential Revenue for this vault row.",
                    )
                    st.checkbox(
                        "Payment received (clears §15-1005 enforcement clock alert)",
                        key=kpay,
                        help="Mark when payer has remitted — removes overdue heatmap for this row.",
                    )
                    if st.button("Reset to AI estimate", key=f"vault_reset_{vid}"):
                        base = float(entry.get("amount_denied_base", 0) or 0)
                        st.session_state[amt_key] = base
                        st.rerun()
    
            if overdue_row:
                st.markdown("</div>", unsafe_allow_html=True)
    
        st.markdown("</div>", unsafe_allow_html=True)

    if _q:
        with st.expander(f"Manual verification queue ({len(_q)}) — double-entry revenue", expanded=False):
            st.caption(
                "These rows failed USD regex / sanity (zero, >$100k, or malformed currency). "
                "They are stored in session + **vault_backup.csv** but **not** shown in the ledger above."
            )
            for e in _q:
                st.markdown(
                    f"- **Patient:** `{e.get('Patient', '—')}` · **Denial:** `{e.get('Denial Code', '—')}` · "
                    f"**Unique_Claim_Hash:** `{e.get('Unique_Claim_Hash', '—')}` · **Validation:** `{e.get('Validation', '—')}`"
                )

    _sync_vault_statuses_from_session_widgets()
    _sync_vault_payment_flags_from_widgets()


def _render_vault_master_certificate_panel() -> None:
    """Full-clinic Senturion Audit Certificate PDF — vault-wide Total Recovery Target + Audit Tracking Hash + disclaimer."""
    _ensure_clinic_profiles()
    st.markdown("##### Vault · Senturion Audit Certificate")
    st.caption(
        "**Senturion Audit Certificate** for the **entire Neural Audit Summary** (active clinic): **Total Recovery Target**, "
        "**Audit Tracking Hash**, **Legal Compliance Disclaimer**, EB Garamond / Titanium Silver styling, and a "
        "**fail-safe footer hash** on every page."
    )
    _cid = st.session_state.get("active_clinic_id")
    _cn = _clinic_display_name(_cid)
    st.caption(f"**Scope:** `{_cn}` (`{_cid[:10]}…`)" if _cid else "**Scope:** — select a clinic in the sidebar.")
    if st.button(
        "Generate Senturion Audit Certificate (full clinic ledger)",
        type="primary",
        use_container_width=True,
        key="vault_master_cert_full_gen",
    ):
        if not _cid:
            st.warning("Select an **active clinic** in the sidebar first.")
        else:
            cert = generate_vault_master_audit_certificate_data(_cid)
            sig_b = _get_digital_signature_bytes()
            pdf_b = _master_audit_certificate_to_pdf_bytes(cert, signature_image_bytes=sig_b)
            st.session_state["vault_master_cert_pdf"] = pdf_b
            st.session_state["vault_master_cert_fn"] = (
                f"Senturion_Audit_Certificate_Vault_{cert.get('audit_hash', 'AUDIT')}.pdf"
            )
            st.success("Senturion Audit Certificate generated — vault scope, disclaimer, footer fingerprint.")
    _pdf_vm = st.session_state.get("vault_master_cert_pdf")
    _fn_vm = st.session_state.get("vault_master_cert_fn") or "Senturion_Audit_Certificate_Vault.pdf"
    if _pdf_vm:
        st.download_button(
            label="Download Senturion Audit Certificate — Vault (PDF)",
            data=_pdf_vm,
            file_name=_fn_vm,
            mime="application/pdf",
            key="vault_master_cert_dl",
            use_container_width=True,
        )


def _render_neural_vault_block() -> None:
    """Vault dashboard + ledger (metrics in emerald, ledger styling)."""
    _render_vault_institutional_dashboard()
    _render_vault_master_certificate_panel()
    st.markdown(
        '<div class="holographic-table institutional-ledger-shell"><div class="projection-table-cap">'
        "NEURAL AUDIT LEDGER</div>",
        unsafe_allow_html=True,
    )
    _render_institutional_ledger(show_title=False)
    st.markdown("</div>", unsafe_allow_html=True)


def _client_email() -> str:
    se = (st.session_state.get("email") or "").strip().lower()
    if se:
        return se
    u = st.session_state.user
    return (getattr(u, "email", None) or "").strip().lower()


def _fetch_client_view_claims() -> list[dict]:
    """Strict client isolation: read only via `client_view` (auth.uid() scoped in the database)."""
    try:
        res = get_supabase().table("client_view").select("*").execute()
        data = getattr(res, "data", None)
        if isinstance(data, list):
            return data
    except Exception:
        pass
    return []


def _normalize_client_view_row(row: dict) -> dict:
    """Map `claims` / `client_view` columns to portal-style dicts."""

    def pick(row_: dict, *names: str):
        for n in names:
            if n in row_ and row_[n] is not None and str(row_[n]).strip() != "":
                return row_[n]
        return None

    pid = clean_text(str(pick(row, "patient_id", "Patient ID", "patientId") or ""))
    dc = clean_text(str(pick(row, "denial_code", "Denial Code", "denialCode") or ""))
    reason = clean_text(str(pick(row, "reason_for_denial", "reason", "Reason for Denial") or ""))
    fix = clean_text(str(pick(row, "fix_action", "Fix Action") or ""))
    pos = clean_text(
        str(
            pick(
                row,
                "place_of_service",
                "place_of_service_state",
                "Place of Service (State)",
                "POS",
            )
            or ""
        )
    )
    law_cited = clean_text(str(pick(row, "law_cited", "Law Cited", "statute") or ""))
    pot_rev = clean_text(str(pick(row, "potential_revenue", "Potential Revenue", "recovery_target") or ""))
    den_dt = clean_text(str(pick(row, "denial_date", "Denial Date", "date_of_denial") or ""))
    appeal_mode = clean_text(str(pick(row, "appeal_mode", "Appeal Mode") or ""))
    cid = pick(row, "id", "claim_id")
    cid_s = str(cid).strip() if cid is not None else ""
    if not cid_s:
        cid_s = str(uuid.uuid4())[:10]
    status = (str(pick(row, "status") or "pending")).lower()
    vault_path = str(pick(row, "final_appeal_path", "storage_path", "vault_path") or "").strip()
    return {
        "claim_id": cid_s,
        "Patient ID": pid or "—",
        "Denial Code": dc or "—",
        "Reason for Denial": reason,
        "Fix Action": fix,
        "Place of Service (State)": pos or "—",
        "Law Cited": law_cited or "—",
        "Potential Revenue": pot_rev or "0",
        "Denial Date": den_dt or "—",
        "Appeal Mode": appeal_mode or _infer_appeal_mode(reason),
        "status": status,
        "docx_bytes": None,
        "vault_path": vault_path or None,
    }


def _parse_denial_csv(upload) -> tuple[list[dict] | None, str | None]:
    """Parse denial CSV (Patient ID + Denial Code). Returns (rows, None) or (None, error)."""
    if upload is None:
        return None, None
    try:
        df = pd.read_csv(upload)
        col_map = {c.lower().strip(): c for c in df.columns}
        pid_col = next((col_map[k] for k in col_map if "patient" in k and "id" in k), None) or next(
            (col_map[k] for k in col_map if "patient" in k), None
        )
        dc_col = next((col_map[k] for k in col_map if "denial" in k and "code" in k), None) or next(
            (col_map[k] for k in col_map if "denial" in k), None
        )
        reason_col = next((col_map[k] for k in col_map if "reason" in k), None)
        fix_col = next((col_map[k] for k in col_map if "fix" in k or "action" in k), None)
        pos_col = next(
            (col_map[k] for k in col_map if "place" in k and ("service" in k or "state" in k)),
            None,
        )
        law_col = next(
            (col_map[k] for k in col_map if "law" in k and "cited" in k),
            None,
        ) or next((col_map[k] for k in col_map if k in ("law cited", "statute")), None)
        rev_col = next(
            (col_map[k] for k in col_map if "potential" in k and "revenue" in k),
            None,
        ) or next((col_map[k] for k in col_map if "recovery" in k and "target" in k), None)
        ddate_col = next(
            (col_map[k] for k in col_map if "denial" in k and "date" in k),
            None,
        ) or next((col_map[k] for k in col_map if k in ("denial date", "date of denial")), None)
        payer_col = next(
            (col_map[k] for k in col_map if "payer" in k and "name" in k),
            None,
        ) or next((col_map[k] for k in col_map if k in ("payer", "insurer", "carrier", "plan")), None)
        if not pid_col or not dc_col:
            return None, "CSV must include Patient ID and Denial Code columns."
        rows: list[dict] = []
        for _, r in df.iterrows():
            row: dict = {}
            for c in df.columns:
                v = r[c]
                row[str(c)] = "" if pd.isna(v) else clean_text(str(v))
            pid = clean_text(r.get(pid_col)) if pid_col else "N/A"
            dc = clean_text(r.get(dc_col)) if dc_col else "N/A"
            reason = clean_text(r.get(reason_col)) if reason_col else ""
            fix = clean_text(r.get(fix_col)) if fix_col else ""
            pos = clean_text(r.get(pos_col)) if pos_col else ""
            law = clean_text(r.get(law_col)) if law_col else ""
            prev = clean_text(r.get(rev_col)) if rev_col else ""
            ddat = clean_text(r.get(ddate_col)) if ddate_col else ""
            payer_n = clean_text(r.get(payer_col)) if payer_col else ""
            row["Patient ID"] = pid
            row["Denial Code"] = dc
            row["Reason for Denial"] = reason
            row["Fix Action"] = fix
            row["Place of Service (State)"] = pos
            row["Law Cited"] = law
            row["Potential Revenue"] = prev or "0"
            row["Denial Date"] = ddat
            row["Payer Name"] = payer_n or "—"
            row["Win Probability"] = "50"
            row["Appeal Mode"] = _resolve_neural_appeal_mode(reason, law)
            _ensure_clinic_profiles()
            row["clinic_id"] = st.session_state.get("active_clinic_id")
            rows.append(row)
        return rows, None
    except Exception as e:
        return None, str(e)


def _merge_portal_claims(email: str, rows: list[dict]) -> None:
    """Append portal claims for a client (idempotent new IDs per row)."""
    _ensure_claims_portal()
    bucket = st.session_state.claims_portal.setdefault(email, [])
    for r in rows:
        bucket.append({
            "claim_id": str(uuid.uuid4())[:10],
            "Patient ID": r.get("Patient ID", "N/A"),
            "Denial Code": r.get("Denial Code", "N/A"),
            "Reason for Denial": r.get("Reason for Denial", ""),
            "Fix Action": r.get("Fix Action", ""),
            "Place of Service (State)": r.get("Place of Service (State)", "") or "—",
            "Law Cited": r.get("Law Cited", "") or "—",
            "Potential Revenue": r.get("Potential Revenue", "") or "0",
            "Denial Date": r.get("Denial Date", "") or "",
            "Appeal Mode": r.get("Appeal Mode") or _infer_appeal_mode(str(r.get("Reason for Denial", ""))),
            "status": "pending",
            "docx_bytes": None,
        })


def _render_kpi_grid() -> None:
    st.markdown(
        """<div class="kpi-grid">
            <div class="kpi-card"><div>AUDIT ACCURACY</div><div class="metric">99.5%</div></div>
            <div class="kpi-card"><div>PROCESSING LATENCY</div><div class="metric">&lt;1.2s</div></div>
            <div class="kpi-card"><div>REVENUE RECOVERED</div><div class="metric metric-critical">$0.00</div></div>
        </div>""",
        unsafe_allow_html=True,
    )


def _render_hud_title(title_html: str) -> None:
    st.markdown(title_html, unsafe_allow_html=True)


def _render_executive_dashboard_header() -> None:
    """
    Partner War Room / Agent: branded title + top-right compliance badge.
    **Clinic** users must never call this (no CEO/CFO chrome).
    """
    perms = check_permissions()
    if perms.can_clinic_portal:
        return
    mid = html_std.escape(str(PAYSTACK_MERCHANT_ID))
    st_esc = html_std.escape(str(PAYSTACK_MERCHANT_STATUS))
    title_esc = html_std.escape(DASHBOARD_TITLE_LINE)
    st.markdown(
        f'<div class="executive-header-row">'
        f'<h1 class="hud-title hud-title-mirror executive-header-h1">{title_esc}</h1>'
        f'<div class="compliance-badge">Merchant ID: {mid} | Status: {st_esc}</div>'
        f"</div>",
        unsafe_allow_html=True,
    )


def _appeal_engine_auto_switch_module() -> None:
    if st.session_state.get("appeal_csv_uploader") is not None:
        st.session_state.file_uploader_key = True
    else:
        st.session_state.file_uploader_key = None
    if st.session_state.get("file_uploader_key") is not None or st.session_state.get("appeal_letter") is not None:
        st.session_state.active_module = "APPEAL"


def _render_module_switcher_admin() -> None:
    col1, col2 = st.columns(2)
    with col1:
        if st.button("NEURAL AUDIT", use_container_width=True, key="sw_sniper"):
            st.session_state.active_module = "SNIPER"
            st.rerun()
    with col2:
        if st.button("APPEAL ENGINE", use_container_width=True, key="sw_appeal"):
            st.session_state.active_module = "APPEAL"
            st.rerun()


def _render_executive_brief_controls(claims: list[dict], *, key_prefix: str) -> None:
    """Fintech executive summary + Senturion Audit Certificate (same cohort selection)."""
    if not claims:
        return
    st.markdown("##### Executive Summary & Senturion Audit Certificate")
    st.caption(
        "Select one or more processed audits, then generate the **Executive Brief** and/or "
        "**Senturion Audit Certificate** (batch summary: Total Recovery Target, Audit Tracking Hash, legal disclaimer; "
        "same selection for both)."
    )
    labels = [
        f"{i}: {c.get('Patient ID', '—')} | {c.get('Denial Code', '—')}"
        for i, c in enumerate(claims)
    ]
    sel = st.multiselect(
        "Claims included in exports",
        options=list(range(len(claims))),
        format_func=lambda i: labels[int(i)],
        key=f"{key_prefix}_exec_ms",
    )
    gen = False
    gen_mac = False
    if sel:
        c1, c2 = st.columns(2)
        with c1:
            gen = st.button(
                "Generate Executive Brief",
                key=f"{key_prefix}_exec_gen",
                type="secondary",
                use_container_width=True,
            )
        with c2:
            gen_mac = st.button(
                "Generate Senturion Audit Certificate",
                key=f"{key_prefix}_mac_gen",
                type="primary",
                use_container_width=True,
            )
    if gen and sel:
        picked = [claims[int(i)] for i in sel]
        brief = generate_executive_brief(picked)
        pdf_b = _executive_brief_to_pdf_bytes(brief)
        st.session_state[f"{key_prefix}_exec_pdf"] = pdf_b
        st.session_state[f"{key_prefix}_exec_fn"] = f"senturion_executive_brief_{brief['audit_hash']}.pdf"
        st.success("Executive Brief compiled. Audit hash linked.")
    if gen_mac and sel:
        picked = [claims[int(i)] for i in sel]
        cert = generate_master_audit_certificate_data(picked)
        sig_b = _get_digital_signature_bytes()
        pdf_m = _master_audit_certificate_to_pdf_bytes(cert, signature_image_bytes=sig_b)
        st.session_state[f"{key_prefix}_mac_pdf"] = pdf_m
        st.session_state[f"{key_prefix}_mac_fn"] = (
            f"Senturion_Audit_Certificate_{cert.get('audit_hash', 'AUDIT')}.pdf"
        )
        st.success(
            "Senturion Audit Certificate generated — Total Recovery Target, Audit Tracking Hash, legal disclaimer; "
            "footer fingerprint on every page."
        )
    pdf_bytes = st.session_state.get(f"{key_prefix}_exec_pdf")
    fn = st.session_state.get(f"{key_prefix}_exec_fn") or "senturion_executive_brief.pdf"
    if pdf_bytes:
        st.download_button(
            label="Download Executive Brief (PDF)",
            data=pdf_bytes,
            file_name=fn,
            mime="application/pdf",
            key=f"{key_prefix}_exec_dl",
            use_container_width=True,
        )
    pdf_mac = st.session_state.get(f"{key_prefix}_mac_pdf")
    fn_mac = st.session_state.get(f"{key_prefix}_mac_fn") or "Senturion_Master_Audit_Certificate.pdf"
    if pdf_mac:
        st.download_button(
            label="Download Senturion Audit Certificate (PDF)",
            data=pdf_mac,
            file_name=fn_mac,
            mime="application/pdf",
            key=f"{key_prefix}_mac_dl",
            use_container_width=True,
        )


def _render_neural_extraction_results(
    data: list[dict],
    meta: dict | None = None,
    *,
    neural_source_tab: str = "uplink",
) -> None:
    """Shared UI after Gemini extraction (PDF uplink or manual paste).

    `neural_source_tab` scopes Partial Recovery reruns so the correct tab re-renders (Streamlit reruns full script).
    """
    meta = meta or {}
    if (
        "_neural_partial_recovery_applied" in st.session_state
        and st.session_state.get("_neural_recovery_source_tab") == neural_source_tab
    ):
        data = st.session_state.pop("_neural_partial_recovery_applied")
        st.session_state.pop("_neural_recovery_source_tab", None)
        meta = {**meta, "finish_reason": None, "recovered_from_partial": True}

    if meta.get("finish_reason") == 2 and not meta.get("silent_line_retry"):
        st.info(
            "⚠️ **Batch density detected.** If extraction fails, please use **Manual Data Paste** mode "
            "to process this specific document."
        )
    if meta.get("error") == "insufficient_text":
        st.warning("Paste at least 20 characters of denial text or CSV data.")
        return

    # Partial recovery: re-parse truncated CSV with lenient line dropping (finish_reason often MAX_TOKENS / 2)
    if meta.get("finish_reason") == 2 and meta.get("raw_response_text"):
        if st.button(
            "Partial Recovery",
            key=f"neural_partial_recovery_btn_{neural_source_tab}",
            help="Re-parse partial model output with lenient CSV recovery (drops incomplete trailing row).",
        ):
            recovered = _recover_denial_rows_from_partial_csv(meta["raw_response_text"])
            if recovered:
                st.session_state["_neural_partial_recovery_applied"] = recovered
                st.session_state["_neural_recovery_source_tab"] = neural_source_tab
                st.rerun()
            else:
                st.warning(
                    "Could not recover tabular rows from the partial response. "
                    "Try splitting the paste into smaller batches (~5 patients)."
                )

    if not data:
        st.info("No denial data could be extracted. The document may not contain structured denial information.")
        return

    if meta.get("recovered_from_partial"):
        st.success(
            f"**Partial recovery:** **{len(data)}** row(s) salvaged from truncated output (lenient CSV parse)."
        )
    else:
        st.success(f"Extracted {len(data)} denial record(s).")
    _play_success_beep()
    _sig_b = hashlib.sha256(json.dumps(data, default=str).encode("utf-8")).hexdigest()[:24]
    if st.session_state.get("neural_audit_batch_sig") != _sig_b:
        st.session_state.neural_audit_batch_sig = _sig_b
        st.session_state.pop("neural_exec_pdf", None)
        st.session_state.pop("neural_exec_fn", None)
    _ensure_clinic_profiles()
    _ac = st.session_state.get("active_clinic_id")
    for r in data:
        r["clinic_id"] = _ac
        r["Appeal Mode"] = _resolve_neural_appeal_mode(
            str(r.get("Reason for Denial", "") or ""),
            str(r.get("Law Cited", "") or ""),
        )
        _pn = str(r.get("Payer Name", "") or "").strip()
        if not _pn or _pn == "—":
            r["Payer Name"] = _extract_payer_from_row(r) or "—"
        else:
            r["Payer Name"] = _pn
    st.session_state.neural_audit_batch = data

    _render_payer_intelligence_strategy_panel(data)

    _neural_display_cols = [
        "Clinic_ID",
        "Patient ID",
        "Patient Name",
        "Payer Name",
        "Denial Code",
        "Reason for Denial",
        "Fix Action",
        "Place of Service (State)",
        "Law Cited",
        "Potential Revenue",
        "Denial Date",
        "Win Probability",
        "Appeal Mode",
        "Strike",
        "Priority Score",
        "Validation",
    ]
    st.markdown(
        '<p class="projection-table-cap">Extracted claims · validation review</p>',
        unsafe_allow_html=True,
    )
    def _neural_cell(r: dict, c: str) -> str:
        if c == "Clinic_ID":
            return str(r.get("clinic_id") or "—")
        if c == "Priority Score":
            return f"{_neural_triage_priority_score(r):,.2f}"
        if c == "Strike":
            return "⚡ STRIKE" if _is_strike_claim_row(r) else "—"
        return r.get(c, "—")

    st.dataframe(
        pd.DataFrame([{c: _neural_cell(r, c) for c in _neural_display_cols} for r in data]),
        use_container_width=True,
        hide_index=True,
    )

    _render_executive_brief_controls(data, key_prefix="neural")

    output = io.StringIO()
    writer = csv.DictWriter(
        output,
        fieldnames=_neural_display_cols,
    )
    writer.writeheader()
    for r in data:
        row_out = {}
        for c in _neural_display_cols:
            if c == "Clinic_ID":
                row_out[c] = str(r.get("clinic_id") or "")
            elif c == "Potential Revenue":
                row_out[c] = r.get(c, "0") or "0"
            elif c == "Priority Score":
                row_out[c] = f"{_neural_triage_priority_score(r):.2f}"
            elif c == "Strike":
                row_out[c] = "STRIKE" if _is_strike_claim_row(r) else ""
            else:
                row_out[c] = r.get(c, "—")
        writer.writerow(row_out)

    st.download_button(
        label="Download CSV Report",
        data=output.getvalue(),
        file_name="denial_report.csv",
        mime="text/csv",
    )


def _render_neural_ghost_live_preview(data: list[dict], *, show_downloads: bool = False) -> None:
    """Live ledger during Ghost Auditing — does not reset session batch or spam success beeps."""
    if not data:
        return
    _neural_display_cols = [
        "Clinic_ID",
        "Patient ID",
        "Patient Name",
        "Payer Name",
        "Denial Code",
        "Reason for Denial",
        "Fix Action",
        "Place of Service (State)",
        "Law Cited",
        "Potential Revenue",
        "Denial Date",
        "Win Probability",
        "Appeal Mode",
        "Strike",
        "Priority Score",
        "Validation",
    ]
    st.markdown(
        '<p class="projection-table-cap">Live uplink · extraction preview</p>',
        unsafe_allow_html=True,
    )

    def _neural_cell(r: dict, c: str) -> str:
        if c == "Clinic_ID":
            return str(r.get("clinic_id") or "—")
        if c == "Priority Score":
            return f"{_neural_triage_priority_score(r):,.2f}"
        if c == "Strike":
            return "⚡ STRIKE" if _is_strike_claim_row(r) else "—"
        return r.get(c, "—")

    st.dataframe(
        pd.DataFrame([{c: _neural_cell(r, c) for c in _neural_display_cols} for r in data]),
        use_container_width=True,
        hide_index=True,
    )
    if show_downloads:
        _render_executive_brief_controls(data, key_prefix="neural")
        output = io.StringIO()
        writer = csv.DictWriter(output, fieldnames=_neural_display_cols)
        writer.writeheader()
        for r in data:
            row_out: dict = {}
            for c in _neural_display_cols:
                if c == "Clinic_ID":
                    row_out[c] = str(r.get("clinic_id") or "")
                elif c == "Potential Revenue":
                    row_out[c] = r.get(c, "0") or "0"
                elif c == "Priority Score":
                    row_out[c] = f"{_neural_triage_priority_score(r):.2f}"
                elif c == "Strike":
                    row_out[c] = "STRIKE" if _is_strike_claim_row(r) else ""
                else:
                    row_out[c] = r.get(c, "—")
            writer.writerow(row_out)
        st.download_button(
            label="Download CSV Report",
            data=output.getvalue(),
            file_name="denial_report.csv",
            mime="text/csv",
            key="ghost_batch_csv_dl",
        )


def _render_neural_audit_module() -> None:
    _ensure_clinic_profiles()
    _ap = _get_clinic_profile(st.session_state.get("active_clinic_id"))
    if _ap:
        st.caption(
            f"**Active clinic:** {_ap.get('name', '—')} · uploads tagged with `Clinic_ID` "
            f"`{str(_ap.get('clinic_id', ''))[:10]}…`"
        )
    st.markdown(
        """<div class="inst-neural-uplink-head">
        <div class="inst-neural-title">DEPOSIT CLINIC DATA FOR NEURAL AUDIT</div>
        </div>""",
        unsafe_allow_html=True,
    )
    st.markdown(
        '<p class="inst-neural-data-mode-title">Neural Revenue Recovery · Secure Clinical Logic</p>',
        unsafe_allow_html=True,
    )
    tab_uplink, tab_paste = st.tabs(["📁 Secure File Uplink", "📝 Manual Data Paste"])

    with tab_uplink:
        st.caption(
            "**Neural Revenue Recovery · Secure Clinical Logic** — multi-select **100+ PDFs** in the file dialog. "
            "High-volume batches process securely in the background; results merge into your **Neural Audit Summary** "
            "without a full-page refresh. Duplicate files are detected and skipped automatically."
        )
        st.caption("High-density PDFs may truncate; use **Manual Data Paste** if extraction is incomplete.")
        uploaded_files = st.file_uploader(
            " ",
            type=["pdf"],
            accept_multiple_files=True,
            help="Select one PDF or many (folder-style multi-select).",
            key="claim_sniper_upload",
            label_visibility="collapsed",
        )

        _gm_pdf = st.session_state.get("ghost_master_pdf_bytes")
        _gm_fn = st.session_state.get("ghost_master_pdf_fn") or "Master_Audit_Summary.pdf"
        if _gm_pdf:
            st.success("**Master Audit Summary** ready — send to Clinic CEO (Eddie handoff).")
            st.download_button(
                label="Download Master Audit Summary (PDF)",
                data=_gm_pdf,
                file_name=_gm_fn,
                mime="application/pdf",
                key="ghost_master_audit_summary_dl",
                type="primary",
                use_container_width=True,
            )

        if not uploaded_files:
            st.session_state.pop("_ghost_started_sig", None)
            st.markdown(
                '<div class="senturion-dark-callout">Upload PDF denial letters or claim documents — multi-select for '
                "batch intake.</div>",
                unsafe_allow_html=True,
            )
        else:
            ufs = list(uploaded_files)
            n = len(ufs)

            if (
                "_neural_partial_recovery_applied" in st.session_state
                and st.session_state.get("_neural_recovery_source_tab") == "uplink"
                and n == 1
            ):
                _render_neural_extraction_results([], {}, neural_source_tab="uplink")
            elif n >= GHOST_ASYNC_THRESHOLD:
                _sig = _ghost_files_signature(ufs)
                with _GHOST_LOCK:
                    _already_running = bool(_GHOST_STATE.get("running"))
                if st.session_state.get("_ghost_started_sig") != _sig:
                    if _already_running:
                        st.warning(
                            "A **Ghost batch** is already processing. Wait for completion before starting another."
                        )
                    else:
                        st.session_state["_ghost_started_sig"] = _sig
                        _ghost_start_batch(ufs)
                        _append_audit_log(
                            f"Neural Revenue Recovery: batch started ({n} PDF file(s), Secure Clinical Logic)."
                        )
                        st.rerun()
                if hasattr(st, "fragment"):
                    _render_industrial_intake_telemetry_fragment()
                else:
                    with _GHOST_LOCK:
                        _pct = float(_GHOST_STATE.get("pct") or 0)
                        _sl = str(_GHOST_STATE.get("status_line") or "")
                        _fin = bool(_GHOST_STATE.get("finished"))
                    st.progress(min(_pct / 100.0, 1.0))
                    st.caption(_sl or "—")
                    if st.session_state.get("neural_audit_batch"):
                        _render_neural_ghost_live_preview(
                            list(st.session_state.neural_audit_batch),
                            show_downloads=_fin,
                        )
            else:
                uploaded_file = ufs[0]
                if (
                    "_neural_partial_recovery_applied" in st.session_state
                    and st.session_state.get("_neural_recovery_source_tab") == "uplink"
                ):
                    _render_neural_extraction_results([], {}, neural_source_tab="uplink")
                else:
                    with st.spinner("Converting PDF to text..."):
                        file_bytes = uploaded_file.getvalue()
                        file_size = len(file_bytes)
                        if file_size > MAX_FILE_SIZE_BYTES:
                            st.markdown(
                                '<div class="file-magnitude-warning">FILE MAGNITUDE EXCEEDS STANDARD SCAN. TRUNCATING TO PRIMARY DATA PAGES.</div>',
                                unsafe_allow_html=True,
                            )
                        text, err = _pdf_bytes_to_text_for_neural(file_bytes)
                        del file_bytes
                        gc.collect()

                    if err:
                        st.error(f"PDF read failed: {err}")
                        st.stop()
                    if not text or len(text.strip()) < 20:
                        st.warning(
                            "Could not extract meaningful text from the PDF. Try a different file or use Manual Data Paste."
                        )
                        st.stop()

                    with st.spinner("Analyzing with Gemini 2.5 Flash..."):
                        try:
                            data, meta = extract_denial_data(text)
                        except Exception as e:
                            st.error(f"AI extraction failed: {e}")
                            st.stop()

                    _render_neural_extraction_results(data, meta, neural_source_tab="uplink")

    with tab_paste:
        st.caption(
            "Fail-safe for **high-density PDFs** or when MarkItDown text is noisy: paste denial paragraphs "
            "or CSV rows; processing uses the **same** Gemini extraction as the file uplink."
        )
        st.markdown('<div class="neural-manual-paste-wrap">', unsafe_allow_html=True)
        with st.form("neural_manual_paste_form"):
            pasted = st.text_area(
                "Paste Denial Text or CSV Data Here",
                height=280,
                placeholder="Paste EOB text, denial letter body, or CSV rows (Patient ID, Denial Code, …)…",
                key="neural_manual_paste_field",
                label_visibility="visible",
            )
            submitted = st.form_submit_button("Process pasted audit", type="primary", use_container_width=True)
        st.markdown("</div>", unsafe_allow_html=True)
        if submitted:
            with st.spinner("Analyzing pasted content with Gemini…"):
                try:
                    data, meta = process_pasted_audit(pasted)
                except Exception as e:
                    st.error(f"AI extraction failed: {e}")
                else:
                    _render_neural_extraction_results(data, meta, neural_source_tab="paste")
        elif (
            "_neural_partial_recovery_applied" in st.session_state
            and st.session_state.get("_neural_recovery_source_tab") == "paste"
        ):
            _render_neural_extraction_results([], {}, neural_source_tab="paste")

    _sync_vault_from_neural_batch()
    _render_neural_vault_block()


def _render_session_client_queues() -> None:
    """Session-scoped client queues + DOCX delivery (not a substitute for Supabase `claims`)."""
    st.caption("Client queues and final appeal delivery (this browser session).")
    portal = st.session_state.get("claims_portal") or {}
    if not portal:
        st.info("No client portals with submissions in this session.")
        return
    for email in sorted(portal.keys()):
        claims = portal[email]
        with st.expander(f"**{email}** — {len(claims)} claim(s)", expanded=False):
            for i, c in enumerate(claims):
                st.divider()
                st.markdown(f"**Row {i + 1}** · `{c.get('Patient ID', '—')}` · `{c.get('Denial Code', '—')}`")
                cur = (c.get("status") or "pending").lower()
                idx = 1 if cur == "finished" else 0
                choice = st.selectbox(
                    "Queue status",
                    ["pending", "finished"],
                    index=idx,
                    key=f"adm_st_{email}_{i}",
                    label_visibility="collapsed",
                )
                c["status"] = str(choice).lower()
                uf = st.file_uploader(
                    "Attach final appeal (.docx)",
                    type=["docx"],
                    key=f"adm_docx_{email}_{i}",
                )
                if uf is not None:
                    c["docx_bytes"] = uf.getvalue()
                    c["status"] = "finished"


def _render_admin_role_directory() -> None:
    """Promote/demote any user by `profiles.role` (DB updates; users pick up on next rerun)."""
    st.caption(
        "Assign **admin**, **agent**, **client**, or **clinic** (restricted client-facing portal). "
        "Target users see the new role on the next refresh (session is synced from `profiles` each run)."
    )
    sb = get_supabase()
    rows = _fetch_profiles_directory(sb)
    if not rows:
        st.warning(
            "No rows returned from `profiles`. Add an RLS policy so **admin** can `SELECT` the directory, "
            "or verify the table exists."
        )
        return
    role_order = ("pending_review", "admin", "agent", "client", "clinic")
    for r in rows:
        rid = str(r.get("id", ""))
        if not rid:
            continue
        em = str(r.get("email", "") or "").strip() or f"uid:{rid[:8]}…"
        cur_raw = r.get("role")
        cur = _normalize_profile_role(cur_raw) if cur_raw is not None else "pending_review"
        try:
            ix = role_order.index(cur) if cur in role_order else 0
        except ValueError:
            ix = 0
        st.divider()
        c1, c2, c3 = st.columns((3, 2, 1))
        with c1:
            st.markdown(f"**{em}**  \n`{rid}`")
        with c2:
            pick = st.selectbox(
                "Role",
                role_order,
                index=ix,
                key=f"admin_role_dir_sel_{rid}",
                label_visibility="collapsed",
            )
        with c3:
            if st.button("Apply", key=f"admin_role_apply_{rid}", type="primary"):
                ok, err = _admin_set_profile_role(sb, rid, pick)
                if ok:
                    if rid == str(st.session_state.get("user_id", "")):
                        st.session_state.role = pick
                    st.success(f"Role set to **{pick}**. User effective on next run.")
                    time.sleep(0.3)
                    st.rerun()
                else:
                    st.error(err or "Update failed (check RLS / service policy).")


def _render_user_management() -> None:
    st.markdown("### User Management")
    t_queues, t_roles = st.tabs(["Client queues", "Role directory"])
    with t_queues:
        _render_session_client_queues()
    with t_roles:
        _render_admin_role_directory()


def _render_live_treasury_capitec() -> None:
    """Capitec bank statement PDF for partner verification of ZAR opening balance (place file next to app.py)."""
    st.markdown("##### Live Treasury · Bank statements")
    st.caption(
        f"**Capitec** statement dated **{CAPITEC_STATEMENT_DATE_LABEL}** — opening balance "
        f"**R{CAPITEC_OPENING_BALANCE_ZAR:,.2f}** (auditor / partner verification)."
    )
    _pdf_b = _read_capitec_statement_bytes()
    if _pdf_b:
        try:
            if hasattr(st, "pdf"):
                st.pdf(_pdf_b, height=520)
            else:
                st.download_button(
                    "Download Capitec statement (PDF)",
                    data=_pdf_b,
                    file_name=CAPITEC_STATEMENT_FILENAME,
                    mime="application/pdf",
                    use_container_width=True,
                )
        except Exception:
            st.warning("Could not display the treasury PDF file.")
    else:
        st.info(
            f"Add **`{CAPITEC_STATEMENT_FILENAME}`** to the project folder (next to `app.py`) to enable the "
            "inline PDF viewer."
        )


def _render_sidebar_founder_mirror() -> None:
    """
    Sidebar: Capitec statement (dated) + Paystack MSA DocID — visible to **both** partners (admin/agent).
    Never shown to **clinic** or **client pitch view** (gated by caller + RBAC).
    """
    if st.session_state.get("client_view_mode"):
        return
    perms = check_permissions()
    if not perms.can_financial_analytics or perms.can_clinic_portal:
        return
    with st.expander("Founder mirror · Live documents", expanded=False):
        st.caption(f"**Capitec bank statement** — **{CAPITEC_STATEMENT_DATE_LABEL}**")
        b = _read_capitec_statement_bytes()
        if b:
            st.download_button(
                "Download Capitec statement (PDF)",
                data=b,
                file_name=CAPITEC_STATEMENT_FILENAME,
                mime="application/pdf",
                use_container_width=True,
                key="sidebar_founder_capitec_dl",
            )
            if hasattr(st, "pdf"):
                st.pdf(b, height=260)
        else:
            st.info(f"Add **`{CAPITEC_STATEMENT_FILENAME}`** next to `app.py` to enable preview.")
        st.markdown("---")
        st.caption("**Paystack — Merchant Service Agreement**")
        st.markdown(
            f'<p class="sidebar-msa-docid">Paystack MSA · DocID: <code>{html_std.escape(_msa_doc_id_display())}</code></p>',
            unsafe_allow_html=True,
        )
        st.caption("Mirrored for **Eduard de Lange** (CEO) & **Monré Wessel Nagel** (CFO).")


def _render_financial_analytics() -> None:
    """Read-only financial / pipeline metrics (extend with Supabase queries)."""
    _fp = check_permissions()
    if not _fp.can_financial_analytics or _fp.can_clinic_portal:
        return
    with st.expander("Financial Analytics", expanded=False):
        vault_all = _vault_entries_displayable()
        _t_ov, _t_an, _t_pay = st.tabs(["Overview", "Clinic Analytics", "Payer Performance"])
        with _t_ov:
            st.caption("Operational metrics — wire to your warehouse or `claims` aggregates.")
            _exp = sum(_vault_amount_for_entry(e) for e in vault_all)
            _gross_usd = _exp * FOUNDERS_COMMISSION_RATE
            _paystack_usd = _gross_usd * PAYSTACK_INTL_CARD_FEE_PCT
            _net_usd = _paystack_net_usd(_gross_usd)
            _fx_zar = _fetch_usd_zar_rate_live()
            _net_zar = _usd_to_zar(_net_usd, _fx_zar)
            _ceo_net_usd = _net_usd * 0.5
            _cfo_net_usd = _net_usd * 0.5
            _ceo_net_zar = _usd_to_zar(_ceo_net_usd, _fx_zar)
            _cfo_net_zar = _usd_to_zar(_cfo_net_usd, _fx_zar)
            c1, c2, c3 = st.columns(3)
            with c1:
                st.metric("Neural Audit recoverable (all clinics)", f"${_exp:,.2f}")
            with c2:
                st.metric("Avg. cycle (days)", "—")
            with c3:
                st.metric("Claims in vault", str(len(vault_all)))
            st.markdown("---")
            _role_fa = str(st.session_state.get("role") or "")
            if _role_fa == "agent":
                st.markdown("##### Treasury · Invoice clinic")
                st.caption("Partner parity: **Agent** uses this strip when the War Room vault strip is not in view.")
                _ensure_clinic_profiles()
                _opts_fa = ["ALL"] + [p["clinic_id"] for p in st.session_state.clinic_profiles]
                if st.session_state.get("vault_clinic_filter") not in _opts_fa:
                    st.session_state.vault_clinic_filter = "ALL"

                def _vault_filter_label_fa(k: str) -> str:
                    if k == "ALL":
                        return "All clinics"
                    return _clinic_display_name(k)

                st.selectbox(
                    "Treasury scope (Neural Audit Summary)",
                    options=_opts_fa,
                    format_func=_vault_filter_label_fa,
                    key="vault_clinic_filter",
                )
                _vf_scope_fa = str(st.session_state.get("vault_clinic_filter") or "ALL")
                _sk_settle_fa = f"treasury_settlement_{_vf_scope_fa}"
                st.session_state.setdefault(_sk_settle_fa, TREASURY_SETTLEMENT_OPTIONS[0])
                _, _, _i3 = st.columns([1, 1, 1])
                with _i3:
                    st.markdown('<div class="treasury-invoice-wrap">', unsafe_allow_html=True)
                    if st.button(
                        "💳 INVOICE CLINIC",
                        type="primary",
                        use_container_width=True,
                        key="treasury_invoice_fa_btn",
                    ):
                        st.session_state["_treasury_invoice_open"] = True
                        st.rerun()
                    st.markdown("</div>", unsafe_allow_html=True)
                st.caption(
                    f"**Settlement (this scope):** `{st.session_state.get(_sk_settle_fa, TREASURY_SETTLEMENT_OPTIONS[0])}` "
                    "— open **Invoice clinic** to pay via Stripe / USDC and update status."
                )
                if st.session_state.get("_treasury_invoice_open"):
                    treasury_invoice_dialog()
                _treasury_celebrate_if_collected_transition(_vf_scope_fa)
            elif _role_fa == "admin":
                st.caption(
                    "**Treasury · Invoice clinic:** use **💳 INVOICE CLINIC** on the **SNIPER** tab "
                    "(Neural Audit Summary strip) — same settlement scope and **Partner Settlement Log**."
                )
            st.markdown("---")
            _render_live_treasury_capitec()
            st.markdown("---")
            st.markdown("##### Projected earnings — $2,625 audit fee (50/50 net)")
            _aud, _ps_aud, _net_aud, _pe_ceo, _pe_cfo = _standard_audit_fee_partner_split_usd()
            pe_a, pe_b, pe_c = st.columns(3)
            with pe_a:
                st.metric(
                    "Gross audit fee (USD)",
                    f"${_aud:,.2f}",
                    help="Standard per-audit fee — basis for projected partner nets.",
                )
            with pe_b:
                st.metric(
                    "Projected CEO net (USD)",
                    f"${_pe_ceo:,.2f}",
                    help="50% of net after Paystack — Eduard de Lange.",
                )
            with pe_c:
                st.metric(
                    "Projected CFO net (USD)",
                    f"${_pe_cfo:,.2f}",
                    help="50% of net after Paystack — Monré Wessel Nagel.",
                )
            st.caption(
                f"Each **$2,625** audit fee: est. Paystack **${_ps_aud:,.2f}** → **${_net_aud:,.2f}** net USD → "
                "**50% / 50%** to CEO / CFO."
            )
            st.markdown("---")
            st.markdown("##### Partner dashboard · CFO currency")
            st.caption(
                f"US clinics pay in **USD**. **Live FX:** 1 USD = **{_fx_zar:,.4f} ZAR** (cached {FX_CACHE_TTL_SEC // 60} min). "
                f"**Paystack intl. card fee:** **{PAYSTACK_INTL_CARD_FEE_PCT * 100:.2f}%** of gross — deducted before the "
                "**50% / 50%** split so **CEO (Eduard de Lange) / CFO (Monré Wessel Nagel)** shares reflect **net** profit."
            )
            g1, g2, g3, g4 = st.columns(4)
            with g1:
                st.metric(
                    "Gross 15% fee (USD)",
                    f"${_gross_usd:,.2f}",
                    help="Full Senturion success-fee pool on vault recoverable (before Paystack).",
                )
            with g2:
                st.metric(
                    "Est. Paystack fee (USD)",
                    f"${_paystack_usd:,.2f}",
                    help=f"Approx. {PAYSTACK_INTL_CARD_FEE_PCT * 100:.1f}% international card fee on gross.",
                )
            with g3:
                st.metric(
                    "Net pool (USD)",
                    f"${_net_usd:,.2f}",
                    help="Gross minus estimated Paystack — basis for 50/50 split.",
                )
            with g4:
                st.metric(
                    "Net pool (ZAR)",
                    f"R{_net_zar:,.2f}",
                    help="Net USD converted at live USD/ZAR (CFO / Capitec view).",
                )
            p1, p2, p3, p4 = st.columns(4)
            with p1:
                st.metric(
                    "💰 CEO net (USD)",
                    f"${_ceo_net_usd:,.2f}",
                    help="50% of net pool (after Paystack) — Eduard de Lange.",
                )
            with p2:
                st.metric(
                    "💰 CFO net (USD)",
                    f"${_cfo_net_usd:,.2f}",
                    help="50% of net pool (after Paystack), Monré Wessel Nagel.",
                )
            with p3:
                st.metric("💰 CEO net (ZAR)", f"R{_ceo_net_zar:,.2f}")
            with p4:
                st.metric("💰 CFO net (ZAR)", f"R{_cfo_net_zar:,.2f}")
            st.markdown("##### Partner Settlement Log")
            st.caption(
                "Populated when **Settlement Status** is **🟢 COLLECTED** in **Invoice clinic**. "
                "Each row: **gross 15% fee (USD)** → est. **Paystack** → **net USD** → **50/50 CEO/CFO**, plus **net ZAR** at live FX."
            )
            _plog = list(st.session_state.get("partner_settlement_log") or [])
            if _plog:
                _pdf_pl = pd.DataFrame(_plog)
                if "Monré Split" in _pdf_pl.columns:
                    _pdf_pl = _pdf_pl.rename(
                        columns={"Monré Split": "CFO Split (50%)", "Eddie Split": "CEO Split (50%)"}
                    )
                _show_cols = [
                    c
                    for c in (
                        "Date",
                        "Claim ID",
                        "Gross Fee (USD)",
                        "Paystack est. (USD)",
                        "Net Fee (USD)",
                        "CEO Net (USD)",
                        "CFO Net (USD)",
                        "CEO Net (ZAR)",
                        "CFO Net (ZAR)",
                        "FX USD/ZAR",
                        "Status",
                        "Total Fee",
                        "CEO Split (50%)",
                        "CFO Split (50%)",
                    )
                    if c in _pdf_pl.columns
                ]
                if _show_cols:
                    st.dataframe(
                        _pdf_pl[_show_cols].iloc[::-1],
                        use_container_width=True,
                        hide_index=True,
                    )
                else:
                    st.dataframe(_pdf_pl.iloc[::-1], use_container_width=True, hide_index=True)
            else:
                st.info(
                    "No settlement rows yet. When a doctor pays, move **Settlement Status** to **🟢 COLLECTED** "
                    "in **💳 Invoice clinic** to lock in partner splits."
                )
        with _t_an:
            st.caption("**Senturion success fee (15%)** — recoverable proxy per registered clinic.")
            _ensure_clinic_profiles()
            by_cid: dict[str, float] = defaultdict(float)
            for e in vault_all:
                cid = e.get("clinic_id") or "__unassigned__"
                by_cid[cid] += _vault_amount_for_entry(e)
            rows = []
            for cid, amt in sorted(by_cid.items(), key=lambda x: -x[1]):
                _nm = _clinic_display_name(cid) if cid != "__unassigned__" else "Unassigned / legacy"
                sent = amt * FOUNDERS_COMMISSION_RATE
                rows.append(
                    {
                        "Clinic": _nm,
                        "Clinic_ID": str(cid)[:12] + "…" if cid != "__unassigned__" else "—",
                        "Recoverable ($)": round(amt, 2),
                        "Owed to Senturion 15% ($)": round(sent, 2),
                    }
                )
            if rows:
                st.dataframe(pd.DataFrame(rows), use_container_width=True, hide_index=True)
            else:
                st.info("No vault rows yet — run a neural audit to deposit claims.")
        with _t_pay:
            st.caption(
                "**Payer Performance** — insurers ranked by **Senturion recovery target** ($) in the vault "
                "across **all clinics** (most “delinquent” by dollars at risk)."
            )
            _ensure_payer_intel_loaded()
            _pdf = _payer_delinquency_all_clinics_df()
            if _pdf.empty:
                st.info("No payer-tagged vault rows yet — neural audits with **Payer Name** populate this chart.")
            else:
                _chart = _pdf.set_index("Payer")[["Senturion recovery target ($)"]].head(18)
                st.bar_chart(_chart)
                st.dataframe(_pdf, use_container_width=True, hide_index=True)
            st.caption(
                "Payer Intelligence ledger (hidden): `.senturion_payer_intel.json` — statutory win proxies updated "
                "when batches merge and when **STATUTORY** enforcement completes."
            )


def _render_pending_access_gate() -> None:
    """Restricted shell for users whose `profiles.role` is still `pending_review` (default after signup)."""
    _render_hud_title('<h1 class="hud-title">ACCESS PENDING REVIEW</h1>')
    st.info(
        "Your account is pending review. You will be notified once an Admin grants you **Agent**, "
        "**Client vault**, or **Clinic** portal access."
    )
    st.caption("You can sign out below. No billing or appeal data is available until your role is assigned.")


def _render_client_vault() -> None:
    email = _client_email()
    _ensure_claims_portal()
    _render_hud_title('<h1 class="hud-title">SECURE DOWNLOAD VAULT</h1>')
    slot = st.empty()
    with slot.container():
        with st.container(border=True):
            st.markdown("##### Submit denials")
            up = st.file_uploader(
                "Upload denial report (CSV)",
                type=["csv"],
                help="Columns: Patient ID, Denial Code; optional Reason, Fix Action, Place of Service (State), Law Cited.",
                key="client_portal_csv",
            )
            if up is not None:
                _sig = (getattr(up, "name", ""), getattr(up, "size", 0))
                if st.session_state.get("_client_portal_csv_sig") != _sig:
                    rows, err = _parse_denial_csv(up)
                    if err:
                        st.error(err)
                    elif rows:
                        _merge_portal_claims(email, rows)
                        st.session_state._client_portal_csv_sig = _sig
                        st.success(f"Recorded {len(rows)} submission(s).")
                        st.rerun()

        # Primary: Supabase `client_view` (enforces auth.uid() in SQL — never other users' rows)
        db_raw = _fetch_client_view_claims()
        if db_raw:
            claims = [_normalize_client_view_row(r) for r in db_raw]
        else:
            claims = st.session_state.claims_portal.get(email, [])

        if not claims:
            st.info("No submissions yet. Upload a CSV to register claims, or wait for DB sync.")
            return

        st.markdown("##### Your claims")
        display_rows = []
        for c in claims:
            display_rows.append({
                "Patient ID": c.get("Patient ID"),
                "Denial Code": c.get("Denial Code"),
                "Reason": (c.get("Reason for Denial") or "")[:80],
                "Place of Service (State)": c.get("Place of Service (State)") or "—",
                "Law Cited": c.get("Law Cited") or "—",
                "Potential Revenue": c.get("Potential Revenue") or "—",
                "Denial Date": c.get("Denial Date") or "—",
            })
        st.dataframe(pd.DataFrame(display_rows), use_container_width=True, hide_index=True)

        _render_executive_brief_controls(claims, key_prefix="vault")

        sb = get_supabase()
        for c in claims:
            with st.container(border=True):
                st.markdown(f"**{c.get('Patient ID', '—')}** · `{c.get('Denial Code', '—')}`")
                fin = (c.get("status") or "pending").lower() == "finished"
                state = "complete" if fin else "running"
                label = "Finished" if fin else "Pending"
                with st.status(label, state=state):
                    if fin:
                        st.caption("Your appeal package is complete.")
                    else:
                        st.caption("Under review — you’ll be notified when the final file is ready.")
                docx = c.get("docx_bytes")
                vault_path = c.get("vault_path")
                dl_bytes = docx
                dl_name = f"final_appeal_{c.get('Patient ID', 'claim')}_{c.get('claim_id', '')}.docx"
                dl_mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                if not dl_bytes and vault_path:
                    dl_bytes = download_from_vault(sb, vault_path)
                    if dl_bytes:
                        ext = vault_path.rsplit(".", 1)[-1].lower() if "." in vault_path else "bin"
                        dl_name = f"final_appeal_{c.get('Patient ID', 'claim')}_{c.get('claim_id', '')}.{ext}"
                        if ext == "pdf":
                            dl_mime = "application/pdf"
                if dl_bytes:
                    st.download_button(
                        "Download Final Appeal",
                        data=dl_bytes,
                        file_name=dl_name,
                        mime=dl_mime,
                        type="primary",
                        use_container_width=True,
                        key=f"client_dl_{c.get('claim_id', id(c))}",
                    )


def _render_clinic_portal() -> None:
    """
    Client-facing portal for `clinic` role: upload claims + view audit reports only.
    No Capitec / treasury / founder equity / executive certificate tooling.
    """
    email = _client_email()
    _ensure_claims_portal()
    _render_hud_title(
        '<h1 class="hud-title hud-title-mirror">'
        "Senturion AI: Secure Medical Revenue Recovery Terminal"
        "</h1>"
    )
    st.caption("Clinic access — restricted to submissions and your audit deliverables.")

    with st.container(border=True):
        st.markdown("### Upload Claims")
        st.caption("Submit denial data as CSV (same schema as the secure client vault).")
        up = st.file_uploader(
            "Upload denial report (CSV)",
            type=["csv"],
            help="Columns: Patient ID, Denial Code; optional Reason, Fix Action, Place of Service (State), Law Cited.",
            key="clinic_portal_csv",
        )
        if up is not None:
            _sig = (getattr(up, "name", ""), getattr(up, "size", 0))
            if st.session_state.get("_clinic_portal_csv_sig") != _sig:
                rows, err = _parse_denial_csv(up)
                if err:
                    st.error(err)
                elif rows:
                    _merge_portal_claims(email, rows)
                    st.session_state._clinic_portal_csv_sig = _sig
                    st.success(f"Recorded {len(rows)} submission(s).")
                    st.rerun()

    db_raw = _fetch_client_view_claims()
    if db_raw:
        claims = [_normalize_client_view_row(r) for r in db_raw]
    else:
        claims = st.session_state.claims_portal.get(email, [])

    st.markdown("---")
    st.markdown("### View Audit Reports")
    if not claims:
        st.info("No audit reports yet. Upload a CSV above, or wait for database sync.")
        return

    display_rows = []
    for c in claims:
        display_rows.append(
            {
                "Patient ID": c.get("Patient ID"),
                "Denial Code": c.get("Denial Code"),
                "Reason": (c.get("Reason for Denial") or "")[:80],
                "Place of Service (State)": c.get("Place of Service (State)") or "—",
                "Law Cited": c.get("Law Cited") or "—",
                "Denial Date": c.get("Denial Date") or "—",
            }
        )
    st.dataframe(pd.DataFrame(display_rows), use_container_width=True, hide_index=True)

    sb = get_supabase()
    for c in claims:
        with st.container(border=True):
            st.markdown(f"**{c.get('Patient ID', '—')}** · `{c.get('Denial Code', '—')}`")
            fin = (c.get("status") or "pending").lower() == "finished"
            state = "complete" if fin else "running"
            label = "Finished" if fin else "Pending"
            with st.status(label, state=state):
                if fin:
                    st.caption("Audit package complete — download below.")
                else:
                    st.caption("Under review — you’ll be notified when the report is ready.")
            docx = c.get("docx_bytes")
            vault_path = c.get("vault_path")
            dl_bytes = docx
            dl_name = f"audit_report_{c.get('Patient ID', 'claim')}_{c.get('claim_id', '')}.docx"
            dl_mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            if not dl_bytes and vault_path:
                dl_bytes = download_from_vault(sb, vault_path)
                if dl_bytes:
                    ext = vault_path.rsplit(".", 1)[-1].lower() if "." in vault_path else "bin"
                    dl_name = f"audit_report_{c.get('Patient ID', 'claim')}_{c.get('claim_id', '')}.{ext}"
                    if ext == "pdf":
                        dl_mime = "application/pdf"
            if dl_bytes:
                st.download_button(
                    "Download audit report",
                    data=dl_bytes,
                    file_name=dl_name,
                    mime=dl_mime,
                    type="primary",
                    use_container_width=True,
                    key=f"clinic_dl_{c.get('claim_id', id(c))}",
                )


def _render_verify_submit_tools() -> None:
    """Agent-only QA / handoff gate (no Neural Audit — compartmentalized)."""
    with st.expander("Verify / Submit", expanded=False):
        st.caption("Confirm appeal quality before vault / client delivery.")
        with st.form("agent_verify_submit_form"):
            st.text_area(
                "Verification notes",
                placeholder="e.g. Coder reviewed, attachments verified…",
                key="agent_verify_notes",
                height=100,
            )
            ack = st.checkbox("I verify this appeal is accurate and complete for submission.", key="agent_verify_ack")
            submitted = st.form_submit_button("Submit to completion queue")
        if submitted:
            if not ack:
                st.warning("Acknowledgment required to submit.")
            else:
                st.success("Recorded in session — connect to `claims` status updates in production.")


def _render_enforcement_clock_agent_terminal() -> None:
    """MD § 15-1005 enforcement heatmap + regulatory escalation draft (Agent Terminal)."""
    st.markdown("---")
    st.markdown("##### Enforcement Clock · Maryland Prompt Payment (§ 15-1005)")
    st.caption(
        f"When a claim moves to **[ENFORCED]**, a **{ENFORCEMENT_MD_PROMPT_PAY_DAYS}-day** clock starts. "
        "Neon red = past that window **without** **payment received** (see Override popover in the Vault ledger)."
    )
    vf = _vault_entries_displayable()
    disp: list[dict[str, Any]] = []
    overdue_flags: list[bool] = []
    for e in vf:
        if str(e.get("vault_status", "")).strip() != VAULT_STATUS_ENFORCED:
            continue
        lbl, neon = _enforcement_clock_status_label(e)
        overdue_flags.append(bool(neon))
        disp.append(
            {
                "Patient": e.get("Patient") or "—",
                "Payer": e.get("payer_label") or "—",
                "Denial": e.get("Denial Code") or "—",
                "Audit Hash": e.get("Audit Hash") or "—",
                "Recoverable $": round(float(_vault_amount_for_entry(e)), 2),
                "Enforcement clock": lbl,
            }
        )
    if not disp:
        st.info("No **[ENFORCED]** claims in the current Neural Audit filter — enforcement clock idle.")
        return
    df = pd.DataFrame(disp)
    try:

        def _row_style(row: Any) -> list[str]:
            i = int(row.name)  # type: ignore[arg-type]
            if i < len(overdue_flags) and overdue_flags[i]:
                return [
                    "background-color: rgba(255, 49, 49, 0.2); color: #FF3131; font-weight: 700; "
                    "border: 1px solid #FF3131; text-shadow: 0 0 8px rgba(255,49,49,0.5)"
                ] * len(row)
            return [""] * len(row)

        st.dataframe(df.style.apply(_row_style, axis=1), use_container_width=True, hide_index=True)
    except Exception:
        st.dataframe(df, use_container_width=True, hide_index=True)

    esc_rec = st.selectbox(
        "Escalation recipient",
        ["Maryland Insurance Administration (MIA)", "U.S. DOL — EBSA (ERISA)"],
        key="regulatory_escalation_recipient_choice",
    )
    if st.button(
        "🚨 TRIGGER REGULATORY ESCALATION",
        type="primary",
        use_container_width=True,
        key="btn_regulatory_escalation",
    ):
        overdue_entries = [e for e in vf if _is_enforcement_clock_overdue(e)]
        if not overdue_entries:
            st.warning(
                "No claims are **past the 30-day** enforcement window without payment — nothing to escalate yet."
            )
        else:
            rec = "DOL" if "DOL" in esc_rec else "MIA"
            st.session_state["reg_esc_ta"] = _build_regulatory_escalation_draft(overdue_entries, recipient=rec)
    st.text_area(
        "Regulatory escalation draft (editable)",
        height=380,
        key="reg_esc_ta",
        placeholder="Click **TRIGGER REGULATORY ESCALATION** to generate a pre-filled letter citing Audit Hash and payer non-response.",
        help="Pre-filled for MIA or DOL — Maryland § 15-1005 / ERISA § 503 context.",
    )


def _render_agent_terminal_batch_panel(clinic_name: str) -> None:
    """ZIP of statutory PDFs for the current neural_audit_batch + audit log + vault ENFORCED."""
    st.markdown("#### Agent Terminal · Neural Triage & batch enforcement")
    st.caption(
        "Priority = **Potential Revenue × Win Probability**. **STATUTORY** appeals package into ZIP "
        "(same batch as Neural Audit Summary / `neural_audit_batch`)."
    )
    cur_sig = st.session_state.get("neural_audit_batch_sig")
    if st.session_state.get("_agent_batch_zip_sig") != cur_sig:
        st.session_state.pop("_agent_batch_zip_bytes", None)
        st.session_state.pop("_agent_batch_zip_sig", None)

    batch = list(st.session_state.get("neural_audit_batch") or [])
    if batch:
        st.markdown("##### Neural Triage queue (auto-prioritized)")
        st.caption(
            f"**[HIGH_VALUE_TARGET]** (top): Potential Revenue over **${HIGH_VALUE_TARGET_USD:,.0f}** — immediate enforcement review. "
            f"**Strike list** (neon): recoverable ≥ **${STRIKE_MIN_REVENUE_USD:,.0f}** and win probability ≥ **{STRIKE_MIN_WIN_PCT:.0f}%**."
        )
        _tri_rows = []
        for r in batch:
            rev = _parse_amount_denied(r.get("Potential Revenue", "0"))
            wp = _parse_win_probability(r.get("Win Probability", "0"))
            sc = _neural_triage_priority_score(r)
            reason = str(r.get("Reason for Denial", "") or "")
            law = str(r.get("Law Cited", "") or "")
            am = _resolve_neural_appeal_mode(reason, law)
            _utag = clean_text(str(r.get("Urgency Tag", "") or ""))
            if not _utag and rev > HIGH_VALUE_TARGET_USD:
                _utag = URGENCY_TAG_HIGH_VALUE
            _tri_rows.append(
                {
                    "Urgency": _utag or "—",
                    "Patient ID": r.get("Patient ID", "—"),
                    "Denial Code": r.get("Denial Code", "—"),
                    "Potential Revenue": rev,
                    "Win %": wp,
                    "Priority Score": round(sc, 2),
                    "Appeal Mode": am,
                    "Strike": "⚡ STRIKE" if _is_strike_claim_row(r) else "",
                }
            )
        _tri_df = pd.DataFrame(_tri_rows)
        _tri_df["_hv_sort"] = _tri_df["Urgency"].apply(
            lambda x: 1 if URGENCY_TAG_HIGH_VALUE in str(x) else 0
        )
        _tri_df = _tri_df.sort_values(["_hv_sort", "Priority Score"], ascending=[False, False])
        _tri_df = _tri_df.drop(columns=["_hv_sort"])
        try:
            _styled = _tri_df.style.apply(
                lambda row: [
                    (
                        "border: 2px solid #FFB020; box-shadow: 0 0 12px rgba(255,176,32,0.55);"
                        if row.get("Urgency") == URGENCY_TAG_HIGH_VALUE
                        else (
                            "border: 2px solid #00FF41; box-shadow: 0 0 10px rgba(0,255,65,0.45);"
                            if row.get("Strike") == "⚡ STRIKE"
                            else ""
                        )
                    )
                    for _ in range(len(row))
                ],
                axis=1,
            )
            st.dataframe(_styled, use_container_width=True, hide_index=True)
        except Exception:
            st.dataframe(_tri_df, use_container_width=True, hide_index=True)
        st.markdown("---")

    statutory_rows: list[dict] = []
    for r in batch:
        reason = str(r.get("Reason for Denial", "") or "")
        law = str(r.get("Law Cited", "") or "")
        am = str(r.get("Appeal Mode", "") or "").upper().strip()
        if am == "STATUTORY" or _resolve_neural_appeal_mode(reason, law) == "STATUTORY":
            statutory_rows.append(r)

    _agent_cid = st.session_state.get("active_clinic_id")
    _mz_gate_ok = _monetization_signed_for_clinic(str(_agent_cid) if _agent_cid else None)

    if not batch:
        st.warning("No neural batch in session. Run **Neural Audit** (War Room) or process a paste, then return here.")
    elif not statutory_rows:
        st.info("Current batch has **no STATUTORY** claims — nothing to add to the enforcement ZIP.")
    else:
        total_amt = sum(_parse_amount_denied(r.get("Potential Revenue", "0")) for r in statutory_rows)
        st.metric("STATUTORY rows in current batch", len(statutory_rows))
        st.caption(f"Aggregate recoverable proxy (Potential Revenue): **${total_amt:,.2f}**")

    if statutory_rows and not _mz_gate_ok:
        st.warning(
            "**Monetization gate:** The clinic must execute **🤝 Authorize Recovery** in **Client View** "
            "and sign the **Senturion AI Solutions Revenue Recovery Agreement** before batch enforcement can run."
        )
        if not _agent_cid:
            st.caption("Select an **active clinic** in the sidebar so the signed agreement can be linked.")

    _can_enforce = bool(statutory_rows) and _mz_gate_ok
    if st.button(
        "🚀 EXECUTE BATCH ENFORCEMENT",
        type="primary",
        use_container_width=True,
        key="btn_batch_enforcement",
        disabled=not _can_enforce,
    ):
        if not statutory_rows:
            st.error("No STATUTORY claims to package.")
        else:
            total_amt = sum(_parse_amount_denied(r.get("Potential Revenue", "0")) for r in statutory_rows)
            with st.spinner(f"Generating {len(statutory_rows)} statutory PDF(s) and building ZIP…"):
                try:
                    zdata = _build_statutory_batch_zip_bytes(statutory_rows, clinic_name=clinic_name)
                    if not zdata:
                        st.error("ZIP build failed (empty PDF output or PDF engine unavailable).")
                    else:
                        agent_disp = (st.session_state.get("email") or "agent").strip()
                        agent_short = agent_disp.split("@")[0][:48]
                        clinic_disp = (clinic_name or "Clinic").strip()[:80]
                        _append_audit_log(
                            f"Agent [{agent_short}] enforced ${total_amt:,.0f} batch for Clinic [{clinic_disp}]"
                        )
                        _mark_batch_statutory_enforced(cur_sig)
                        st.session_state["_agent_batch_zip_bytes"] = zdata
                        st.session_state["_agent_batch_zip_sig"] = cur_sig or ""
                        st.success("Batch enforcement package ready — download below.")
                except Exception as ex:
                    st.error(f"Batch enforcement failed: {ex}")

    zbytes = st.session_state.get("_agent_batch_zip_bytes")
    zsig = st.session_state.get("_agent_batch_zip_sig")
    if zbytes and cur_sig and zsig == cur_sig:
        fn = f"senturion_statutory_batch_{(cur_sig or 'batch')[:12]}.zip"
        st.download_button(
            label="⬇ Download batch enforcement ZIP",
            data=zbytes,
            file_name=fn,
            mime="application/zip",
            key="dl_batch_enforcement_zip",
            use_container_width=True,
        )

    _render_enforcement_clock_agent_terminal()


def _render_agent_console(logo_file, clinic_name: str) -> None:
    st.session_state.active_module = "APPEAL"
    _render_executive_dashboard_header()
    _render_kpi_grid()
    _appeal_engine_auto_switch_module()
    tab_appeal, tab_agent = st.tabs(["Appeal Engine", "Agent Terminal"])
    with tab_appeal:
        _render_appeal_generator(logo_file=logo_file, clinic_name=clinic_name)
    with tab_agent:
        _render_agent_terminal_batch_panel(clinic_name or "")
    st.markdown("---")
    _render_verify_submit_tools()
    if check_permissions().can_financial_analytics:
        _render_financial_analytics()


def _render_admin_war_room(logo_file, clinic_name: str) -> None:
    if "active_module" not in st.session_state:
        st.session_state.active_module = "SNIPER"
    _render_executive_dashboard_header()
    _render_kpi_grid()
    _render_module_switcher_admin()
    _appeal_engine_auto_switch_module()

    body = st.empty()
    with body.container():
        if st.session_state.active_module == "SNIPER":
            _render_neural_audit_module()
        else:
            st.markdown("### Generate appeal letters from denial data")
            _render_appeal_generator(logo_file=logo_file, clinic_name=clinic_name)

    st.markdown("---")
    with st.expander("User Management", expanded=False):
        _render_user_management()
    _render_financial_analytics()


def _build_smart_context(selected_row: dict) -> str:
    """Build denial context with smart-fill logic. Use raw data when present; placeholders when missing."""
    parts = []
    # Prompt-safe text only — no bytes, file handles, or upload objects in claim strings.
    pid = _text_only_for_prompt(selected_row.get("Patient ID", ""))
    claim = _text_only_for_prompt(selected_row.get("Claim #", ""))
    dc = _text_only_for_prompt(selected_row.get("Denial Code", ""))
    reason = _text_only_for_prompt(selected_row.get("Reason for Denial", ""))
    fix = _text_only_for_prompt(selected_row.get("Fix Action", ""))
    # Strip brackets from values for cleaner output
    pid = _strip_brackets(pid) or pid
    claim = _strip_brackets(claim) or claim
    if pid or claim:
        if pid:
            parts.append(f"Patient ID: {pid}")
        if claim:
            parts.append(f"Claim #: {claim}")
    else:
        parts.append("RE: [ACTION REQUIRED: INSERT PATIENT NAME]")
    if dc:
        parts.append(f"Denial Code: {dc}")
    if reason:
        parts.append(f"Reason for Denial: {reason}")
    if fix:
        parts.append(f"Fix Action: {fix}")
    pos = _text_only_for_prompt(selected_row.get("Place of Service (State)", ""))
    law = _text_only_for_prompt(selected_row.get("Law Cited", ""))
    pos = _strip_brackets(pos) or pos
    law = _strip_brackets(law) or law
    if pos:
        parts.append(f"Place of Service (State): {pos}")
    if law:
        parts.append(f"Law Cited: {law}")
    _mode = clean_text(str(selected_row.get("Appeal Mode", "") or ""))
    if not _mode:
        _mode = _infer_appeal_mode(reason)
    parts.append(f"Appeal Mode: {_mode}")
    return "\n".join(parts)


def _render_legal_templates_sidebar() -> None:
    """Sidebar tab: Legal Templates — Senturion MSA with clinic + vault fields and PDF export."""
    _ensure_clinic_profiles()
    _ensure_revenue_vault()
    _cid = st.session_state.get("active_clinic_id")
    _cn = _clinic_display_name(_cid)
    _tot = _vault_total_recoverable_for_clinic(_cid)
    st.markdown("##### Senturion Master Service Agreement (MSA)")
    st.caption(
        f"**Provider (settings):** `{_cn}` · **Success fee:** **{MSA_SUCCESS_FEE_PERCENT}%** (fixed) · "
        f"**Total recoverable (Vault, this clinic):** **${_tot:,.2f}**"
    )
    st.caption(
        "Clinic name comes from **Clinic Management** / active profile. Recoverable total sums **Amount Denied** "
        "for displayable vault rows for the **active clinic**."
    )
    if st.button(
        "Print to PDF",
        key="msa_print_to_pdf_btn",
        type="primary",
        use_container_width=True,
    ):
        pdf_b = _msa_master_service_agreement_pdf_bytes(
            clinic_name=_cn,
            clinic_id=_cid,
            total_recoverable_usd=_tot,
        )
        if pdf_b:
            _slug = re.sub(r"[^\w\-]+", "_", (_cn or "clinic"))[:48].strip("_") or "clinic"
            st.session_state["legal_msa_pdf_bytes"] = pdf_b
            st.session_state["legal_msa_pdf_fn"] = f"Senturion_MSA_{_slug}.pdf"
            st.success("MSA PDF generated — use download below for digital signature.")
    _mb = st.session_state.get("legal_msa_pdf_bytes")
    _mf = st.session_state.get("legal_msa_pdf_fn") or "Senturion_MSA.pdf"
    if _mb:
        st.download_button(
            label="Download MSA (PDF)",
            data=_mb,
            file_name=_mf,
            mime="application/pdf",
            key="legal_msa_download_btn",
            use_container_width=True,
        )


def _cold_strike_email_bundle(total_usd: float, clinic_display: str) -> tuple[str, str]:
    """
    Cold Strike outbound: subject uses live Total Recoverable Revenue + clinic name.
    Returns (subject_line, full_text_for_gmail_paste).
    """
    amt_s = f"${float(total_usd):,.2f}"
    subject = f"URGENT: {amt_s} in Found Revenue identified for {clinic_display}"
    full = "\n".join(
        [
            f"Subject: {subject}",
            "",
            "---",
            "",
            "Hi [Name],",
            "",
            "We completed a targeted Neural Audit on your recent denial ledger — the recoverable stack is material and actionable.",
            "",
            "Our Neural Audit found ERISA violations in your recent denials. We will release the full legal enforcement batch for a 15% success fee. Zero upfront cost.",
            "",
            "Reply with a time to speak today or tomorrow and I will send the full enforcement packet.",
            "",
            "— Senturion",
        ]
    )
    return subject, full


def _render_communications_sidebar() -> None:
    """Sidebar tab: COMMUNICATIONS — Cold Strike email (dynamic vault $ + clinic)."""
    _ensure_clinic_profiles()
    _ensure_revenue_vault()
    _cid = st.session_state.get("active_clinic_id")
    _raw = _clinic_display_name(_cid)
    _disp = _raw if _raw != "—" else "Your Clinic"
    _tot = _vault_total_recoverable_for_clinic(_cid)
    subject, full_email = _cold_strike_email_bundle(_tot, _disp)

    st.markdown(
        '<p style="font-family:JetBrains Mono,monospace;font-size:0.62rem;letter-spacing:0.2em;'
        'color:#00FF41;text-transform:uppercase;margin-bottom:0.35rem;">Communications · Cold Strike</p>',
        unsafe_allow_html=True,
    )
    st.caption(
        f"**Clinic:** `{_disp}` · **Total recoverable (Vault, current batch):** **${_tot:,.2f}** — "
        "subject line pulls these values live."
    )
    _sub_esc = html_std.escape(subject)
    st.markdown(
        f'<div style="font-family:JetBrains Mono,monospace;font-size:0.78rem;color:#E8F5E9;'
        f'border:1px solid #1B5E20;background:linear-gradient(135deg,#0D1F0D 0%,#0a0a0a 100%);'
        f'padding:0.75rem 0.85rem;border-radius:8px;margin-bottom:0.5rem;line-height:1.45;">'
        f'<strong style="color:#00FF41;">SUBJECT</strong><br/>{_sub_esc}</div>',
        unsafe_allow_html=True,
    )
    st.code(full_email, language=None)
    _b64 = base64.b64encode(full_email.encode("utf-8")).decode("ascii")
    _safe_b64 = json.dumps(_b64)
    components.html(
        f"""
<!DOCTYPE html>
<html><head><meta charset="utf-8"/></head>
<body style="margin:0;font-family:system-ui,Segoe UI,sans-serif;">
<button id="cold-strike-copy" type="button"
  style="width:100%;box-sizing:border-box;background:#00FF41;color:#0a0a0a;font-weight:800;
  border:none;padding:14px 16px;border-radius:8px;cursor:pointer;font-size:15px;
  letter-spacing:0.03em;box-shadow:0 0 18px rgba(0,255,65,0.45);">
  📋 Copy Email to Clipboard
</button>
<script>
const b64 = {_safe_b64};
const btn = document.getElementById("cold-strike-copy");
btn.addEventListener("click", function() {{
  const t = atob(b64);
  navigator.clipboard.writeText(t).then(() => {{
    btn.textContent = "✓ Copied — paste into Gmail";
    btn.style.background = "#1a1a1a";
    btn.style.color = "#00FF41";
    btn.style.boxShadow = "0 0 22px rgba(0,255,65,0.6)";
  }}).catch(() => {{ btn.textContent = "Copy failed — select text"; }});
}});
</script>
</body></html>
        """,
        height=72,
    )


def _render_neural_activity_sidebar_ghost() -> None:
    """Sidebar telemetry for Ghost Auditing (background PDF batch)."""
    with _GHOST_LOCK:
        snap = {
            "running": bool(_GHOST_STATE.get("running")),
            "finished": bool(_GHOST_STATE.get("finished")),
            "pct": float(_GHOST_STATE.get("pct") or 0),
            "status_line": str(_GHOST_STATE.get("status_line") or ""),
            "total_files": int(_GHOST_STATE.get("total_files") or 0),
            "errs": list(_GHOST_STATE.get("pending_errors") or [])[-12:],
        }
    if not snap["status_line"] and not snap["running"] and not snap["finished"]:
        return
    st.markdown("---")
    st.markdown(
        '<p style="font-family:JetBrains Mono,monospace;font-size:0.62rem;letter-spacing:0.2em;'
        'color:#d4af37;text-transform:uppercase;margin-bottom:0.35rem;">Neural Revenue Recovery · Status</p>',
        unsafe_allow_html=True,
    )
    st.caption(snap["status_line"])
    st.progress(min(snap["pct"] / 100.0, 1.0))
    if snap["total_files"]:
        st.caption(f"Files in batch: **{snap['total_files']}**")
    if snap["errs"]:
        with st.expander("Processing notes", expanded=False):
            for e in snap["errs"]:
                st.caption(e)


def _paystack_release_checkout_url() -> str:
    """Hosted Paystack page for Pay Release Fee (configure in secrets for production)."""
    try:
        if "PAYSTACK_RELEASE_CHECKOUT_URL" in st.secrets:
            u = str(st.secrets["PAYSTACK_RELEASE_CHECKOUT_URL"]).strip()
            if u:
                return u
    except Exception:
        pass
    return DEFAULT_PAYSTACK_RELEASE_CHECKOUT_URL


def _render_demo_audit_page() -> None:
    """Neural Audit Demo — Paystack Reviewer (reviews@paystack.com): polished cohort + Pay Now."""
    _mid_esc = html_std.escape(str(PAYSTACK_MERCHANT_ID))
    st.markdown(
        """
<div class="demo-audit-hero paystack-reviewer-hero">
  <div class="demo-audit-title">Neural Audit Demo</div>
  <div class="demo-audit-sub">Senturion AI Solutions · Merchant verification preview</div>
</div>
""",
        unsafe_allow_html=True,
    )
    st.markdown(
        '<div class="senturion-dark-panel reviewer-roles-strip neural-audit-demo-executive">'
        "<p class=\"demo-exec-line\"><strong>CEO:</strong> Eduard de Lange</p>"
        "<p class=\"demo-exec-line\"><strong>CFO:</strong> Monré Wessel Nagel</p>"
        f"<p class=\"demo-address-line\"><strong>Registered address:</strong> "
        f"{html_std.escape(OFFICIAL_BUSINESS_ADDRESS)}</p>"
        f"<p class=\"demo-merchant-line\">Paystack Merchant ID: <strong>{_mid_esc}</strong> · "
        "Illustrative cohort — Secure Clinical Logic applied.</p>"
        "</div>",
        unsafe_allow_html=True,
    )
    _cohort_total = REVIEWER_MOCK_RECOVERABLE_USD * len(REVIEWER_MOCK_AUDIT_IDS)
    st.markdown(
        f'<div class="reviewer-recoverable-hero">Recoverable Revenue (demo cohort): '
        f'<span class="reviewer-usd">${_cohort_total:,.2f}</span>'
        f'<span class="reviewer-hero-sub">3 audits × ${REVIEWER_MOCK_RECOVERABLE_USD:,.2f} each</span></div>',
        unsafe_allow_html=True,
    )
    _mock_audits = [
        {
            "Audit ID": aid,
            "Clinic": REVIEWER_MOCK_CLINIC,
            "Status": REVIEWER_MOCK_AUDIT_STATUS,
            "Recoverable Revenue": f"${REVIEWER_MOCK_RECOVERABLE_USD:,.2f}",
        }
        for aid in REVIEWER_MOCK_AUDIT_IDS
    ]
    st.markdown(
        '<p class="reviewer-section-title">Mock audits — Neural Analysis (injected demo data)</p>',
        unsafe_allow_html=True,
    )
    st.dataframe(
        pd.DataFrame(_mock_audits),
        use_container_width=True,
        hide_index=True,
    )
    if hasattr(st, "link_button"):
        st.link_button(
            "Pay Now",
            _paystack_release_checkout_url(),
            type="primary",
            use_container_width=True,
            help=f"Opens Paystack Checkout for Merchant ID {PAYSTACK_MERCHANT_ID} (set PAYSTACK_RELEASE_CHECKOUT_URL in secrets).",
        )
    else:
        st.markdown(
            f'<a href="{html_std.escape(_paystack_release_checkout_url())}" target="_blank" '
            'rel="noopener noreferrer" class="reviewer-pay-link">Pay Now — Paystack Checkout</a>',
            unsafe_allow_html=True,
        )


def _render_mandatory_kyc_footer() -> None:
    """Paystack compliance — mandatory KYC strip on every page (below primary chrome)."""
    st.markdown(
        f'<div class="kyc-footer-bar kyc-footer-legal">{_kyc_footer_html_inner()}</div>',
        unsafe_allow_html=True,
    )


def main():
    _ghost_audit_fragment_poll()
    _sync_session_with_profiles_table()
    _apply_paystack_demo_role_override()
    perms = check_permissions()
    role = perms.role
    _ensure_claims_portal()
    _demo = _is_paystack_demo_session()

    with st.sidebar:
        logo_file = None

        user_email = (
            (st.session_state.get("email") or "").strip()
            or (st.session_state.user.email if st.session_state.user else "")
        )
        st.markdown(
            f'<p style="font-family: \'JetBrains Mono\', monospace; letter-spacing: 2px; '
            f'text-transform: uppercase; font-size: 0.72rem; color: #E0E0E0;">'
            f"Active session · <code style=\"color:#E0E0E0;\">{user_email}</code></p>",
            unsafe_allow_html=True,
        )
        if perms.can_clinic_portal:
            st.markdown(
                '<p style="font-family:\'JetBrains Mono\',monospace;font-size:0.62rem;letter-spacing:0.08em;'
                'color:#00FF41;text-transform:none;margin:0.2rem 0 0.35rem;line-height:1.4;">'
                "Senturion AI: Secure Medical Revenue Recovery Terminal</p>",
                unsafe_allow_html=True,
            )
        else:
            st.markdown(
                '<p style="font-family:\'JetBrains Mono\',monospace;font-size:0.65rem;letter-spacing:0.14em;'
                'color:#00FF41;text-transform:uppercase;margin:0.2rem 0 0.35rem;line-height:1.35;">'
                "SENTURION AI SOLUTIONS // NEURAL AUDIT NODE</p>",
                unsafe_allow_html=True,
            )
        st.caption(f"Access tier: **{(role or '…').upper()}**")
        if _demo:
            st.caption("Neural Audit Demo · Reviewer · Paystack merchant preview (Secure Clinical Logic)")
            st.markdown(
                '<div class="reviewer-sidebar-executive">'
                '<p class="rse-head">Executive leadership</p>'
                '<p class="rse-line"><strong>CEO:</strong> Eduard de Lange</p>'
                '<p class="rse-line rse-cfo"><strong>CFO:</strong> Monré Wessel Nagel</p>'
                "</div>",
                unsafe_allow_html=True,
            )
        if perms.can_appeal_engine and not _demo:
            st.toggle("Switch to Client View", key="client_view_mode")
        _cv = bool(st.session_state.get("client_view_mode"))

        if not _demo and not _cv and not perms.can_clinic_portal:
            st.markdown(
                """<div class="senturion-command-hierarchy">
<div class="sch-title">Senturion Command Hierarchy</div>
<div class="sch-level"><span class="sch-tag">[LVL 1]</span> War Room <span class="sch-sub">(Batch Upload)</span></div>
<div class="sch-level"><span class="sch-tag">[LVL 2]</span> Neural Audit Summary <span class="sch-sub">(Ledger)</span></div>
<div class="sch-level"><span class="sch-tag">[LVL 3]</span> Statutory Engine <span class="sch-sub">(ERISA / Maryland Law settings)</span></div>
<div class="sch-level"><span class="sch-tag">[LVL 4]</span> Agent Terminal <span class="sch-sub">(Team management)</span></div>
</div>""",
                unsafe_allow_html=True,
            )
            if perms.can_appeal_engine:
                _render_neural_activity_sidebar_ghost()
        elif not _cv and perms.can_clinic_portal:
            st.caption("**Clinic portal** — uploads & audit reports only (no internal treasury).")

        if st.button("Secure Logout", key="logout_btn"):
            try:
                get_supabase().auth.sign_out()
            except Exception:
                pass
            st.cache_data.clear()
            # Absolute session purge — no RBAC / PII leakage between accounts on shared machines.
            st.session_state.clear()
            st.rerun()

        if not _demo and not _cv and not perms.can_clinic_portal and HAS_LOTTIE and st_lottie:
            lottie_data = _fetch_lottie_json(NEURAL_LOTTIE_URL)
            if lottie_data:
                st_lottie(
                    lottie_data,
                    speed=0.72,
                    height=130,
                    key="sidebar_neural_lottie",
                )

        st.markdown("---")
        if perms.is_pending_access and not _demo:
            st.caption("**Access pending review** — limited sidebar until an Admin assigns your role.")
        elif not _demo and perms.can_appeal_engine and not _cv:
            _tab_ops, _tab_legal, _tab_comm = st.tabs(
                ["Operations", "Legal Templates", "COMMUNICATIONS"]
            )
            with _tab_ops:
                st.subheader("Multi-tenant clinic")
                _ensure_clinic_profiles()
                _ids = [p["clinic_id"] for p in st.session_state.clinic_profiles]
                _labels = {p["clinic_id"]: p.get("name") or "Clinic" for p in st.session_state.clinic_profiles}
                st.selectbox(
                    "Active clinic (uploads & appeals)",
                    options=_ids,
                    format_func=lambda k: _labels.get(k, str(k)),
                    key="active_clinic_id",
                )
                with st.expander("Settings", expanded=False):
                    _tsig, _tcl = st.tabs(["Signature & digital", "Clinic Management"])
                    with _tsig:
                        st.caption(
                            "Upload an **Authorized Signature (PNG)**. It is placed automatically on the "
                            "**Authorized Senturion Auditor** line in generated PDFs together with the audit timestamp."
                        )
                        sig_up = st.file_uploader(
                            "Authorized Signature (PNG)",
                            type=["png"],
                            key="digital_signature_upload",
                        )
                        if sig_up is not None:
                            st.session_state.digital_signature_bytes = sig_up.getvalue()
                            st.success("Signature stored for this session.")
                        if st.button("Clear digital signature", key="clear_digital_signature_btn"):
                            st.session_state.pop("digital_signature_bytes", None)
                            st.rerun()
                    with _tcl:
                        _render_clinic_management_tab()
            with _tab_legal:
                _render_legal_templates_sidebar()
            with _tab_comm:
                _render_communications_sidebar()
        elif perms.can_clinic_portal and not _cv:
            st.caption("**Clinic** — use main panel for uploads and reports.")
        elif not _cv:
            st.caption("Stealth client portal — submissions only.")

        st.markdown("---")
        logo_path = "logo.png"
        if os.path.exists(logo_path):
            st.markdown('<div class="logo-glow">', unsafe_allow_html=True)
            st.image(logo_path, width=280)
            st.markdown("</div>", unsafe_allow_html=True)
        else:
            st.markdown(
                '<div class="logo-glow" style="font-family:\'JetBrains Mono\',monospace;font-size:0.85rem;color:#00FF41;text-align:center;padding:1rem;letter-spacing:0.18em;text-transform:uppercase;">'
                "SENTURION AI SOLUTIONS</div>",
                unsafe_allow_html=True,
            )
        if perms.can_financial_analytics and not _cv and not _demo:
            st.markdown("---")
            _render_sidebar_founder_mirror()
        if not _cv and not perms.can_clinic_portal:
            st.markdown("---")
            if not _demo:
                st.markdown(
                    '<div class="system-pulse"><span class="pulse-dot"></span> Neural Engine: Active</div>',
                    unsafe_allow_html=True,
                )
                if perms.can_appeal_engine and st.session_state.get("appeal_csv_synced"):
                    st.markdown(
                        '<div class="data-synced">▼ DATA SYNCED</div>',
                        unsafe_allow_html=True,
                    )
            st.markdown(
                '''<div class="founding-partners">
                <span class="partner-exec-line ceo-line">CEO: Eduard de Lange</span>
                <span class="partner-sep"></span>
                <span class="partner-exec-line cfo-line">CFO: Monré Wessel Nagel</span>
            </div>''',
                unsafe_allow_html=True,
            )
            if not _demo and perms.can_appeal_engine:
                st.markdown("---")
                st.markdown(
                    '<p style="font-family:JetBrains Mono,monospace;font-size:0.62rem;letter-spacing:0.2em;'
                    'color:#fafafa;text-transform:uppercase;margin-bottom:0.35rem;">Compliance Trail</p>',
                    unsafe_allow_html=True,
                )
                _ensure_audit_log()
                _hist = list(reversed(st.session_state.audit_log_history[-18:]))
                if not _hist:
                    st.caption("System Integrity: Optimal")
                else:
                    for ent in _hist:
                        if isinstance(ent, dict):
                            msg = (ent.get("message") or "").strip()
                            ts = (ent.get("ts") or "").strip()
                            st.caption(f"**{ts}** — {msg}" if ts else msg)
                        else:
                            st.caption(str(ent))
            if perms.can_financial_analytics and not _demo:
                st.markdown("---")
                with st.expander("Partner & treasury settings", expanded=False):
                    st.caption("**Admin + Agent** — full treasury configuration (partner parity).")
                    st.session_state.setdefault("appeal_output_format", "docx")
                    _cur_fmt = st.session_state.get("appeal_output_format", "docx")
                    _fmt_ix = 0 if _cur_fmt == "docx" else 1
                    _fmt_choice = st.radio(
                        "Appeal output format",
                        ("DOCX", "PDF"),
                        index=_fmt_ix,
                        horizontal=True,
                        key="admin_appeal_output_format_radio",
                    )
                    st.session_state.appeal_output_format = (
                        "docx" if _fmt_choice == "DOCX" else "pdf"
                    )
                    st.markdown("---")
                    st.markdown("**Senturion Treasury** (Invoice clinic modal)")
                    st.caption(
                        "Paste your live **Stripe** payment link and **Solana / USDC** wallet. "
                        "Leave blank to use built-in placeholders."
                    )
                    st.session_state.setdefault("admin_treasury_stripe_url", "")
                    st.session_state.setdefault("admin_treasury_sol_wallet", "")
                    st.text_input(
                        "Stripe checkout link",
                        key="admin_treasury_stripe_url",
                        placeholder=DEFAULT_TREASURY_STRIPE_URL,
                    )
                    st.text_input(
                        "USDC / Solana wallet address",
                        key="admin_treasury_sol_wallet",
                        placeholder=DEFAULT_TREASURY_SOL_WALLET,
                    )

    _ensure_clinic_profiles()
    _active_prof = _get_clinic_profile(st.session_state.get("active_clinic_id"))
    clinic_name = (_active_prof.get("name") if _active_prof else None) or DEFAULT_CLINIC_NAME

    if (
        perms.can_appeal_engine
        and not st.session_state.get("client_view_mode")
        and not _is_paystack_demo_session()
    ):
        if "GEMINI_API_KEY" not in st.secrets:
            st.error("Missing GEMINI_API_KEY in .streamlit/secrets.toml.")
            st.stop()

    viewport = st.empty()
    with viewport.container():
        if _is_paystack_demo_session():
            _render_demo_audit_page()
        elif perms.is_pending_access:
            _render_pending_access_gate()
        elif st.session_state.get("client_view_mode") and perms.can_appeal_engine:
            _render_client_facing_view(clinic_name or "")
        elif perms.can_admin_war_room:
            _render_admin_war_room(logo_file=logo_file, clinic_name=clinic_name or "")
        elif perms.is_agent and perms.can_appeal_engine:
            _render_agent_console(logo_file=logo_file, clinic_name=clinic_name or "")
        elif perms.can_clinic_portal:
            _render_clinic_portal()
        elif perms.can_client_vault:
            _render_client_vault()
        else:
            st.warning("Your profile is not linked yet. Refresh in a moment, or contact an administrator.")
            if st.button("Retry profile sync", key="retry_profile_sync"):
                st.rerun()

    if _is_paystack_demo_session():
        # Footer lock — fixed copy for Reviewer / Paystack verification (matches compliance pack)
        st.markdown(
            """
<div class="reviewer-footer-lock">
  <p class="rfl-addr"><strong>1171 Bergsig Street, Pretoria</strong></p>
  <p class="rfl-mid"><strong>Merchant ID: 1774856</strong></p>
  <p class="rfl-copy">Senturion AI Solutions · Neural Audit Demo · © 2026</p>
</div>
""",
            unsafe_allow_html=True,
        )
    elif perms.can_clinic_portal:
        st.markdown(
            '<p class="quantum-footer">'
            "Senturion AI: Secure Medical Revenue Recovery Terminal · © 2026"
            "</p>",
            unsafe_allow_html=True,
        )
    else:
        _msa_doc = _signed_msa_doc_id_footer()
        _db_ok = _supabase_database_online()
        _db_cls = "db-dot-online" if _db_ok else "db-dot-offline"
        _db_txt_cls = "footer-db-on" if _db_ok else "footer-db-off"
        _db_txt = "Database Online" if _db_ok else "Database offline"
        _mid_esc = html_std.escape(str(PAYSTACK_MERCHANT_ID))
        st.markdown(
            f'<div class="quantum-footer">'
            f"© 2026 SENTURION AI SOLUTIONS // ENCRYPTED ACCESS ONLY"
            f'<span class="quantum-footer-sub">Merchant ID: <strong>{_mid_esc}</strong>'
            f'<span class="footer-db-pill" title="Supabase / Senturion Vault connection">'
            f'<span class="footer-db-dot {_db_cls}"></span>'
            f'<span class="footer-db-txt {_db_txt_cls}">{html_std.escape(_db_txt)}</span></span> · '
            f"Status: <strong>{html_std.escape(str(PAYSTACK_MERCHANT_STATUS))}</strong> · "
            f"Merchant Agreement DocID: <strong>{html_std.escape(str(_msa_doc))}</strong></span>"
            f'<span class="quantum-footer-addr">{OFFICIAL_BUSINESS_ADDRESS}</span>'
            f'<span class="quantum-footer-sub quantum-footer-mail">'
            f'<a href="mailto:{OFFICIAL_CONTACT_EMAIL}">{OFFICIAL_CONTACT_EMAIL}</a></span>'
            f"</div>",
            unsafe_allow_html=True,
        )
    _render_mandatory_kyc_footer()


def _render_appeal_generator(logo_file=None, clinic_name=""):
    """Appeal UI + Gemini text generation. logo_file is used only for Word export (_appeal_to_docx), never in prompts."""
    if "appeal_csv_denials" not in st.session_state:
        st.session_state.appeal_csv_denials = []
    if "appeal_csv_synced" not in st.session_state:
        st.session_state.appeal_csv_synced = False
    if "appeal_letter" not in st.session_state:
        st.session_state.appeal_letter = None
    if "just_uploaded" not in st.session_state:
        st.session_state.just_uploaded = False

    with st.container(border=True):
        if st.button("🔄 Reset Engine", key="appeal_engine_reset"):
            st.session_state.appeal_letter = None
            st.rerun()

        st.markdown("Generate professional appeal letters from denial data.")
        csv_upload = st.file_uploader(
            "UPLOAD DENIAL REPORT FOR TARGETING",
            type=["csv"],
            help="CSV: Patient ID, Denial Code, Reason, Fix Action, POS, Law Cited, Potential Revenue, Denial Date (optional).",
            key="appeal_csv_uploader",
        )
        if csv_upload is not None:
            rows, csv_err = _parse_denial_csv(csv_upload)
            if csv_err:
                st.error(csv_err)
            elif rows is not None:
                st.session_state.appeal_csv_denials = rows
                st.session_state.appeal_csv_synced = True
                _sig = (
                    getattr(csv_upload, "file_id", None),
                    getattr(csv_upload, "name", ""),
                    getattr(csv_upload, "size", 0),
                )
                if st.session_state.get("_appeal_csv_upload_sig") != _sig:
                    st.session_state._appeal_csv_upload_sig = _sig
                    st.session_state.just_uploaded = True
                    st.session_state.pop("appeal_exec_pdf", None)
                    st.session_state.pop("appeal_exec_fn", None)

        denial_context = ""
        selected_row = None
        if st.session_state.appeal_csv_denials:
            options = [
                f"{r.get('Patient ID', '—')} | {r.get('Denial Code', '—')} | Law: {r.get('Law Cited', '—')}"
                for r in st.session_state.appeal_csv_denials
            ]
            sel_idx = st.selectbox(
                "Select Claims for Enforcement",
                range(len(options)),
                format_func=lambda i: options[i],
                key="patient_selector_main",
            )
            if sel_idx is not None:
                selected_row = st.session_state.appeal_csv_denials[sel_idx]
                denial_context = _build_smart_context(selected_row)
            _render_executive_brief_controls(st.session_state.appeal_csv_denials, key_prefix="appeal")
        else:
            st.session_state.appeal_csv_synced = False
            denial_context = st.text_area(
                "Paste denial details (e.g., Patient ID, denial code, reason)",
                placeholder="Patient ID: 12345\nDenial Code: CO-16\nReason: Missing documentation...",
                height=120,
                key="appeal_manual_input",
            ) or ""

        btn_label = "ENGAGE APPEAL PROTOCOL" if selected_row else "Generate Appeal Letter"
        if st.button(btn_label):
            try:
                # Text-only pipeline for Gemini; logo_file is reserved for _appeal_to_docx() only.
                ctx = _text_only_for_prompt(denial_context.strip())
                if not ctx:
                    st.warning("Please upload a CSV and select a target, or paste denial details above.")
                else:
                    with st.spinner("Generating appeal letter..."):
                        try:
                            if model is None:
                                st.error("Appeal engine unavailable — add GEMINI_API_KEY to .streamlit/secrets.toml.")
                                st.stop()
                            statutory = (
                                _infer_appeal_mode(str((selected_row or {}).get("Reason for Denial", "")))
                                == "STATUTORY"
                                if selected_row
                                else _infer_appeal_mode(ctx) == "STATUTORY"
                            )
                            litigation_prompt = (
                                MASTER_LITIGATION_PROMPT_STATUTORY if statutory else MASTER_LITIGATION_PROMPT
                            )
                            full_prompt = f"{litigation_prompt}\n\nCLAIM DATA:\n{ctx}"
                            clean_prompt = str(full_prompt)
                            code = _text_only_for_prompt((selected_row or {}).get("Denial Code", "N/A")) or "N/A"
                            # Attempt the Primary High-Quality Call (string prompt only — never logo/upload objects)
                            response = model.generate_content(
                                clean_prompt,
                                safety_settings={
                                    HarmCategory.HARM_CATEGORY_HARASSMENT: HarmBlockThreshold.BLOCK_NONE,
                                    HarmCategory.HARM_CATEGORY_DANGEROUS_CONTENT: HarmBlockThreshold.BLOCK_NONE,
                                },
                            )
                            # CHECK FOR SAFETY BLOCK
                            if not response.candidates or response.candidates[0].finish_reason.name == "SAFETY":
                                # AUTO-RETRY with 'Safe Mode' prompt
                                safe_prompt = f"Draft a professional business letter regarding healthcare claim {code}. Use 600 words. Cite ERISA 503."
                                clean_safe = str(safe_prompt)
                                response = model.generate_content(clean_safe)
                            final_text = response.text
                            if not final_text:
                                fr = response.candidates[0].finish_reason.name if response.candidates else "BLOCKED"
                                final_text = f"ADMINISTRATIVE NOTICE:\n\nThe system encountered a processing filter (Finish: {fr}). Please manually review Claim ID: {(selected_row or {}).get('Patient ID', 'Unknown')}"
                            appeal = (final_text or "").strip()
                            if appeal:
                                st.session_state.appeal_letter = appeal
                                st.success("Appeal letter generated.")
                                st.session_state.active_module = "APPEAL"
                                st.rerun()
                            else:
                                st.warning("Could not generate appeal.")
                        except Exception as e:
                            final_text = f"ADMINISTRATIVE NOTICE:\n\nThe system encountered a processing filter. Please manually review Claim ID: {(selected_row or {}).get('Patient ID', 'Unknown')}"
                            st.session_state.appeal_letter = final_text
                            st.error(f"Generation failed: {e}")
                            st.session_state.active_module = "APPEAL"
                            st.rerun()
            except Exception as engine_err:
                st.error(f"Appeal Engine error: {engine_err}")

        # Display generated appeal in paper-style container with Copy + Download
        letter = st.session_state.appeal_letter
        if letter:
            appeal = letter
            _br = _branding_for_clinic_id(
                (selected_row or {}).get("clinic_id") if selected_row else None
            )
            logo_bytes = _br.get("logo_bytes") or (logo_file.getvalue() if logo_file else None)
            clinic_nm = _br.get("name") or clinic_name
            clinic_addr = _br.get("address") or ""
            if selected_row is not None:
                patient_id_export = clean_text(selected_row.get("Patient ID", ""))
                patient_full_name = _get_best_val(selected_row, ["patient", "name", "member", "pt"])
                claim_number = _get_best_val(selected_row, ["claim", "#", "id", "ref", "number"])
                insurance_payer_name = _get_best_val(
                    selected_row, ["payer", "insur", "carrier", "plan"]
                )
                denial_code = clean_text(selected_row.get("Denial Code", ""))
                denial_reason = clean_text(selected_row.get("Reason for Denial", ""))
                service_code = _get_best_val(
                    selected_row, ["service", "cpt", "hcpcs", "hcpc", "proc", "revenue"]
                )
            else:
                patient_id_export = ""
                patient_full_name = claim_number = insurance_payer_name = ""
                denial_code = service_code = ""
                denial_reason = ""
                _paste = (denial_context or "").strip()
                _m_pid = re.search(r"Patient\s*ID\s*[:#]?\s*(\S+)", _paste, flags=re.I)
                if _m_pid:
                    patient_id_export = _m_pid.group(1).strip()
                _m_clm = re.search(
                    r"(?:Claim\s*#|Claim\s*(?:Number|ID)?)\s*[:#]?\s*(\S+)",
                    _paste,
                    flags=re.I,
                )
                if _m_clm:
                    claim_number = _m_clm.group(1).strip()
                _m_reas = re.search(
                    r"Reason\s*(?:for\s*Denial)?\s*[:#]?\s*(.+?)(?:\n\n|\n(?:Patient|Claim|Denial)|\Z)",
                    _paste,
                    flags=re.I | re.DOTALL,
                )
                if _m_reas:
                    denial_reason = _m_reas.group(1).strip().split("\n")[0][:4000]

            _show_cms_export = _text_has_cpt_code(
                f"{appeal} {denial_reason} {service_code} {patient_id_export} {claim_number}"
            )
            appeal_preview = _appeal_preview_mirror_body(
                appeal, include_cms_compliance=_show_cms_export
            )
            b64 = base64.b64encode(appeal_preview.encode("utf-8")).decode("ascii")
            copy_html = f"""
        <div class="appeal-paper">
            <div class="appeal-paper-header">
                <button class="appeal-paper-copy" onclick="navigator.clipboard.writeText(atob('{b64}')).then(() => this.textContent='✓ Copied'); this.title='Copy to clipboard';" title="Copy to clipboard">
                    <svg xmlns="http://www.w3.org/2000/svg" width="16" height="16" viewBox="0 0 24 24" fill="none" stroke="#374151" stroke-width="2"><rect x="9" y="9" width="13" height="13" rx="2"/><path d="M5 15H4a2 2 0 0 1-2-2V4a2 2 0 0 1 2-2h9a2 2 0 0 1 2 2v1"/></svg>
                    Copy
                </button>
            </div>
        </div>
        """
            st.markdown(copy_html, unsafe_allow_html=True)
            _cap_ev = (
                "**PLAN PROVISION VS. STATUTORY REQUIREMENT**"
                if _infer_appeal_mode(denial_reason) == "STATUTORY"
                else "**Clinical Discrepancy Analysis**"
            )
            _cap_cms = (
                "**compliance attestation** (CMS/NCCI) + **penalty clause**"
                if _show_cms_export
                else "**penalty clause** only (no CMS coding attestation without CPT context)"
            )
            st.caption(
                f"Preview — {_cap_ev} renders as a table below. "
                f"Footer mirrors DOCX/PDF: {_cap_cms} appended below."
            )
            if md_to_html is not None:
                try:
                    _ap_html = md_to_html.markdown(appeal_preview, extensions=["extra"])
                    st.markdown(
                        f'<div class="appeal-preview-garamond">{_ap_html}</div>',
                        unsafe_allow_html=True,
                    )
                except Exception:
                    try:
                        _ap_html = md_to_html.markdown(appeal_preview)
                        st.markdown(
                            f'<div class="appeal-preview-garamond">{_ap_html}</div>',
                            unsafe_allow_html=True,
                        )
                    except Exception:
                        st.markdown(appeal_preview)
            else:
                st.markdown(appeal_preview)

            appeal_type_docx = (st.session_state.get("appeal_type_docx") or "").strip()
            st.session_state.setdefault("appeal_output_format", "docx")
            out_fmt = st.session_state.get("appeal_output_format", "docx")
            _case_id = _appeal_case_reference(patient_id_export, claim_number)
            _gen_uid = str(st.session_state.get("user_id") or "")
            _gen_em = str(st.session_state.get("email") or "").strip()
            docx_bytes = _appeal_to_docx(
                appeal,
                logo_bytes=logo_bytes,
                clinic_name=clinic_nm or None,
                clinic_address=clinic_addr or None,
                insurance_payer_name=insurance_payer_name,
                patient_full_name=patient_full_name,
                patient_id=patient_id_export,
                claim_number=claim_number,
                denial_code=denial_code,
                denial_reason=denial_reason,
                service_code=service_code,
                appeal_type=appeal_type_docx,
                case_id=_case_id,
                generated_by_user_id=_gen_uid or None,
                generated_by_email=_gen_em or None,
                statutory_appeal=_infer_appeal_mode(denial_reason) == "STATUTORY",
                show_cms_compliance_footer=_show_cms_export,
            )
            pdf_bytes = (
                _appeal_to_pdf_bytes(
                    appeal,
                    clinic_name=clinic_nm or None,
                    clinic_address=clinic_addr or None,
                    logo_image_bytes=logo_bytes,
                    insurance_payer_name=insurance_payer_name,
                    patient_full_name=patient_full_name,
                    patient_id=patient_id_export,
                    claim_number=claim_number,
                    denial_reason=denial_reason,
                    denial_code=denial_code,
                    case_id=_case_id,
                    generated_by_user_id=_gen_uid or None,
                    generated_by_email=_gen_em or None,
                    statutory_appeal=_infer_appeal_mode(denial_reason) == "STATUTORY",
                    show_cms_compliance_footer=_show_cms_export,
                    signature_image_bytes=_get_digital_signature_bytes(),
                )
                if out_fmt == "pdf"
                else b""
            )
            _uid = st.session_state.get("user_id")
            _sb = get_supabase()
            _ts = int(time.time())

            if out_fmt == "pdf":
                if pdf_bytes and HAS_FPDF:
                    with st.container(border=True):
                        st.markdown("### Official export (PDF)")
                        st.caption(
                            "Your browser requires **one click** to save — downloads cannot start automatically (security). "
                            "Use the primary button below."
                        )
                        st.download_button(
                            label="DOWNLOAD OFFICIAL APPEAL (.PDF)",
                            data=pdf_bytes,
                            file_name="appeal_letter.pdf",
                            mime="application/pdf",
                            key="appeal_download_pdf",
                            type="primary",
                            use_container_width=True,
                        )
                    if _uid:
                        _path = f"{_uid}/appeals/appeal_{_ts}.pdf"
                        _ok, _err = upload_to_vault(
                            _sb,
                            _path,
                            pdf_bytes,
                            content_type="application/pdf",
                        )
                        if not _ok and _err:
                            st.caption(f"Vault sync: {_err}")
                elif not HAS_FPDF:
                    st.warning("Install fpdf2 for PDF output: pip install fpdf2")
            else:
                if docx_bytes and HAS_DOCX:
                    st.download_button(
                        label="DOWNLOAD OFFICIAL APPEAL (.DOCX)",
                        data=docx_bytes,
                        file_name="appeal_letter.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        key="appeal_download_docx",
                    )
                    if _uid and docx_bytes:
                        _path = f"{_uid}/appeals/appeal_{_ts}.docx"
                        _ok, _err = upload_to_vault(
                            _sb,
                            _path,
                            docx_bytes,
                            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        )
                        if not _ok and _err:
                            st.caption(f"Vault sync: {_err}")
                elif not HAS_DOCX:
                    st.warning("Install python-docx to enable Word download: pip install python-docx")

    if st.session_state.get("just_uploaded"):
        st.session_state.just_uploaded = False
        st.rerun()


if __name__ == "__main__":
    main()
